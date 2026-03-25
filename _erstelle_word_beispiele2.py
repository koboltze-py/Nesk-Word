# -*- coding: utf-8 -*-
"""
Erstellt 5 moderne Word-Designs mit echten Dienstplandaten (25.03.2026).
Zielordner: Desktop/bei
"""
import os, sys
from pathlib import Path
from datetime import datetime
from collections import defaultdict

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Zielordner ────────────────────────────────────────────────────────────────
_OD = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
ZIELORDNER = os.path.join(_OD, "Desktop", "bei") if os.path.exists(_OD) else r"C:\Temp\bei"
os.makedirs(ZIELORDNER, exist_ok=True)

LOGO_PFAD = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Daten", "Email", "Logo.jpg")
EXCEL_PFAD = (
    r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
    r"\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26"
    r"\04_Tagesdienstpläne\03_März\25.03.2026.xlsx"
)

# ── Echtdaten laden ───────────────────────────────────────────────────────────
def lade_dienstplan():
    try:
        from functions.dienstplan_parser import DienstplanParser
        r = DienstplanParser(EXCEL_PFAD, alle_anzeigen=True).parse()
        if r.get("success") in (True, "True"):
            return r
    except Exception as e:
        print(f"  [Parser-Warnung] {e}")
    return None

# ── Verbrauchsmaterial ────────────────────────────────────────────────────────
MEDIZIN_MATERIAL = [
    ("Einmalhandschuhe Nitril (M/L)",   "Karton",     "8",   "3"),
    ("Verbandpäckchen groß/klein",      "Stück",      "24",  "10"),
    ("Mullbinden 6cm / 8cm / 10cm",     "Rollen",     "60",  "20"),
    ("Wundkompressen steril 10x10",     "Päckchen",   "40",  "15"),
    ("Pflaster-Sortiment",              "Päckchen",   "10",  "4"),
    ("Hände-Desinfektionsmittel 1L",    "Flaschen",   "6",   "2"),
    ("Flächen-Desinfektion 1L",         "Flaschen",   "4",   "2"),
    ("Rettungsdecken (Gold/Silber)",    "Stück",      "20",  "8"),
    ("Einmal-Beatmungsmaske",           "Stück",      "10",  "4"),
    ("Venenverweilkanülen G18/G20",     "Stück",      "30",  "10"),
    ("Einmalspritzen 5ml/10ml/20ml",    "Stück",      "50",  "20"),
    ("Infusionssets + NaCl 0,9%/500ml", "Sets",       "10",  "4"),
    ("EKG-Elektroden Einmal",           "Päckchen",   "8",   "3"),
    ("Sauerstoffmasken Einmal",         "Stück",      "15",  "5"),
    ("Atemschutzmaske FFP2",            "Stück",      "50",  "20"),
    ("AED-Elektroden (Einmal)",         "Paar",       "4",   "2"),
    ("Blutzucker-Teststreifen",         "Päckchen",   "5",   "2"),
    ("Urinbeutel steril",               "Stück",      "12",  "4"),
    ("Blasenkatheter Ch14/Ch16",        "Stück",      "6",   "2"),
    ("Druckverband CELOX Hämostase",    "Stück",      "4",   "2"),
]

# ── Farb-Paletten ─────────────────────────────────────────────────────────────
def rgb(r,g,b): return RGBColor(r,g,b)

# Palette 1 – Command Center (dunkelblau/cyan)
P1 = dict(
    bg1=rgb(0x0D,0x1B,0x2A), bg2=rgb(0x1B,0x26,0x3B), bg3=rgb(0x41,0x5A,0x77),
    acc=rgb(0x00,0xB4,0xD8), acc2=rgb(0x90,0xE0,0xEF), txt=rgb(0xFF,0xFF,0xFF),
    sub=rgb(0xCA,0xE9,0xFF), gold=rgb(0xFF,0xD1,0x66),
    grn=rgb(0x2A,0xC0,0x7E), org=rgb(0xFF,0x99,0x00), rot=rgb(0xFF,0x33,0x33),
    grau=rgb(0xCC,0xCC,0xCC), hell=rgb(0xF4,0xF8,0xFF),
)
# Palette 2 – Swiss Clean (weiß/schwarz/rot)
P2 = dict(
    bg1=rgb(0xFF,0xFF,0xFF), bg2=rgb(0xF5,0xF5,0xF5), bg3=rgb(0xE8,0xE8,0xE8),
    acc=rgb(0xBE,0x00,0x00), acc2=rgb(0xE8,0x30,0x30), txt=rgb(0x1A,0x1A,0x1A),
    sub=rgb(0x55,0x55,0x55), gold=rgb(0x1A,0x1A,0x1A),
    grn=rgb(0x10,0x7E,0x3E), org=rgb(0xE6,0x7E,0x00), rot=rgb(0xBE,0x00,0x00),
    grau=rgb(0x88,0x88,0x88), hell=rgb(0xF5,0xF5,0xF5),
)
# Palette 3 – Midnight Navy (dunkelblau/gold)
P3 = dict(
    bg1=rgb(0x0A,0x0F,0x2E), bg2=rgb(0x14,0x1E,0x4A), bg3=rgb(0x1E,0x2D,0x6B),
    acc=rgb(0xF5,0xC5,0x18), acc2=rgb(0xFF,0xE0,0x7A), txt=rgb(0xFF,0xFF,0xFF),
    sub=rgb(0xD4,0xD8,0xFF), gold=rgb(0xF5,0xC5,0x18),
    grn=rgb(0x00,0xE0,0x7F), org=rgb(0xFF,0xA5,0x00), rot=rgb(0xFF,0x44,0x44),
    grau=rgb(0xAA,0xAA,0xCC), hell=rgb(0xF0,0xF2,0xFF),
)
# Palette 4 – Airport Signal Board (hellgrau/teal)
P4 = dict(
    bg1=rgb(0x1C,0x1C,0x1E), bg2=rgb(0x2C,0x2C,0x2E), bg3=rgb(0x3A,0x3A,0x3C),
    acc=rgb(0x30,0xD1,0x58), acc2=rgb(0x34,0xC7,0x59), txt=rgb(0xFF,0xFF,0xFF),
    sub=rgb(0xAA,0xFF,0xDC), gold=rgb(0xFF,0xD6,0x00),
    grn=rgb(0x30,0xD1,0x58), org=rgb(0xFF,0x9F,0x0A), rot=rgb(0xFF,0x45,0x3A),
    grau=rgb(0xAA,0xAA,0xAA), hell=rgb(0xF0,0xF8,0xF4),
)
# Palette 5 – Kräftiges DRK-Rot / Clean-White
P5 = dict(
    bg1=rgb(0xBE,0x00,0x00), bg2=rgb(0x96,0x00,0x00), bg3=rgb(0xDC,0x14,0x14),
    acc=rgb(0xFF,0xFF,0xFF), acc2=rgb(0xF0,0xF0,0xF0), txt=rgb(0xFF,0xFF,0xFF),
    sub=rgb(0xFF,0xCC,0xCC), gold=rgb(0xFF,0xF0,0x90),
    grn=rgb(0x00,0xCC,0x66), org=rgb(0xFF,0xAA,0x00), rot=rgb(0xFF,0xFF,0x00),
    grau=rgb(0xCC,0xCC,0xCC), hell=rgb(0xF9,0xF9,0xF9),
)

def _s(rgb): return str(rgb)  # RGBColor → RRGGBB-String

def _shd(cell, color: RGBColor):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), _s(color))
    p.append(s)

def _border_bottom(para, color="888888", sz="6"):
    pPr = para._p.get_or_add_pPr(); b = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), color)
    b.append(bot); pPr.append(b)

def _cell_txt(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]; p.alignment = align
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = color
    return r

def _col_width(tbl, widths_cm):
    """Spaltenbreiten in cm setzen."""
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW_list = tcPr.findall(qn('w:tcW'))
                for old in tcW_list:
                    tcPr.remove(old)
                tcW = OxmlElement('w:tcW')
                twips = int(widths_cm[i] * 567)
                tcW.set(qn('w:w'), str(twips))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.insert(0, tcW)

def _no_space_before(para):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '40')
    pPr.append(sp)

def _bulmor_status(im_einsatz, p):
    if im_einsatz <= 2:   return p['rot'],  "KRITISCH"
    elif im_einsatz == 3: return p['org'],  "EINGESCHRÄNKT"
    else:                 return p['grn'],  "VOLLSTÄNDIG"

def _add_logo_header(doc, palette, titel_text, unt_text=""):
    """Logo + Org-Text als Kopfzeile."""
    section = doc.sections[0]
    header  = section.header
    ht = header.add_table(rows=1, cols=2, width=Inches(6.5))
    ht.autofit = False
    logo_path = Path(LOGO_PFAD)
    if logo_path.exists():
        lp = ht.rows[0].cells[0].paragraphs[0]
        lp.add_run().add_picture(str(logo_path), width=Inches(1.1))
    else:
        ht.rows[0].cells[0].paragraphs[0].add_run("DRK")
    rp = ht.rows[0].cells[1].paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = rp.add_run("Deutsches Rotes Kreuz Kreisverband Köln e.V.\n")
    r1.font.size = Pt(9); r1.font.bold = True; r1.font.color.rgb = palette['acc'] if _s(palette['acc']) != _s(rgb(0xFF,0xFF,0xFF)) else rgb(0xBE,0,0)
    r2 = rp.add_run("Unfallhilfsstelle · Erste-Hilfe-Station Flughafen Köln/Bonn")
    r2.font.size = Pt(8); r2.font.color.rgb = palette['sub'] if _s(palette['bg1']) == _s(rgb(0xFF,0xFF,0xFF)) else rgb(0x55,0x55,0x55)
    tp = header.add_paragraph()
    _border_bottom(tp, _s(palette['acc']) if len(_s(palette['acc']))==6 else "BE0000", "8")

def _add_footer(doc):
    f = doc.sections[0].footer
    fp = f.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("☎ +49 2203 40-2323   |   ✉ flughafen@drk-koeln.de   |   Stationsleitung: Lars Peters")
    r.font.size = Pt(8); r.font.color.rgb = rgb(0x88,0x88,0x88)

def _margins(doc, l=1.8, r=1.8, t=2.8, b=2.0):
    for s in doc.sections:
        s.left_margin = Cm(l); s.right_margin = Cm(r)
        s.top_margin  = Cm(t); s.bottom_margin = Cm(b)

# ── Dienstplan gruppieren ─────────────────────────────────────────────────────
def gruppiere_nach_zeit(personen, ist_dispo=False):
    """Gibt sortiertes dict {zeitschlüssel: [namen]} zurück."""
    gruppen = defaultdict(list)
    for p in personen:
        if p.get('ist_krank') in (True, 'True'): continue
        start = (p.get('start_zeit') or '')[:5]
        end   = (p.get('end_zeit')   or '')[:5]
        if ist_dispo:
            if start and ':' in start: start = f"{int(start.split(':')[0]):02d}:00"
            if end   and ':' in end:   end   = f"{int(end.split(':')[0]):02d}:00"
        key = f"{start} – {end}"
        gruppen[key].append(p.get('anzeigename',''))
    return dict(sorted(gruppen.items()))

def alle_betreuer_schichten(data):
    """Alle Betreuer sortiert nach Schichtbeginn."""
    betreuer = [p for p in data.get('betreuer',[]) if p.get('ist_krank') not in (True,'True')]
    gruppen = defaultdict(list)
    for p in betreuer:
        s = (p.get('start_zeit') or '')[:5]
        e = (p.get('end_zeit')   or '')[:5]
        key = f"{s} – {e}"
        gruppen[key].append(p.get('anzeigename',''))
    return dict(sorted(gruppen.items()))

def bulmorfahrer_liste(data):
    alle = data.get('betreuer',[]) + data.get('dispo',[])
    return [p.get('anzeigename','') for p in alle if p.get('ist_bulmorfahrer') in (True,'True')]

def kranke_liste(data):
    return [p.get('anzeigename','') for p in data.get('kranke',[]) if p.get('ist_krank') in (True,'True')]


# ═══════════════════════════════════════════════════════════════════════════════
# VARIANTE 1 – COMMAND CENTER (dunkelblau / cyan)
# ═══════════════════════════════════════════════════════════════════════════════
def variante1_command_center(data, im_einsatz=4, einsaetze=28, patienten=5, pax=42500):
    P = P1
    doc = Document(); _margins(doc); _add_logo_header(doc, P, ""); _add_footer(doc)

    # Titel-Banner
    banner = doc.add_table(rows=1, cols=1); banner.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = banner.rows[0].cells[0]; _shd(c, P['bg1'])
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STÄRKEMELDUNG UND EINSÄTZE")
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = P['acc']
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("DRK Köln · Erste-Hilfe-Station Flughafen Köln/Bonn · 25.03.2026 · 07:45 Uhr")
    r2.font.size = Pt(9); r2.font.color.rgb = P['sub']

    doc.add_paragraph()

    # ── Kennzahlen-Scorecard ──────────────────────────────────────────────────
    sc = doc.add_table(rows=2, cols=4); sc.style = 'Table Grid'
    kz = [("Einsätze", str(einsaetze), P['acc']),
          ("Pat. auf Station", str(patienten), P['gold']),
          ("Passagiere (PAX)", f"{pax:,}".replace(",","."), P['acc2']),
          ("Personal gesamt", str(_gesamt_personal(data)), P['grn'])]
    for i,(lbl,val,f) in enumerate(kz):
        hc=sc.rows[0].cells[i]; vc=sc.rows[1].cells[i]
        _shd(hc,P['bg2']); _shd(vc,P['bg3'])
        _cell_txt(hc,lbl,bold=True,size=9,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(vc,val,bold=True,size=24,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Bulmor-Status ─────────────────────────────────────────────────────────
    _section_header(doc, "BULMOR – FAHRZEUGSTATUS (5 Fahrzeuge)", P['acc'], P['bg1'])
    bul = doc.add_table(rows=2, cols=5); bul.style = 'Table Grid'
    farbe, label = _bulmor_status(im_einsatz, P)
    for i in range(5):
        aktiv = (i+1) <= im_einsatz
        c1 = bul.rows[0].cells[i]; c2 = bul.rows[1].cells[i]
        _shd(c1, farbe if aktiv else P['bg2'])
        _shd(c2, P['bg3'])
        sym = "◉" if aktiv else "○"
        _cell_txt(c1, f"{sym}  B{i+1}", bold=True, size=13,
                  color=P['txt'] if aktiv else P['grau'],
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        status_txt = "IM EINSATZ" if aktiv else "RESERVE"
        _cell_txt(c2, status_txt, size=8,
                  color=farbe if aktiv else P['grau'],
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"Status: {im_einsatz} von 5 Bulmor im Einsatz  —  {label}")
    sr.font.size=Pt(12); sr.font.bold=True; sr.font.color.rgb=farbe
    bdf = bulmorfahrer_liste(data)
    bfp = doc.add_paragraph()
    bfp.add_run(f"Bulmor-Fahrer heute: {', '.join(bdf)}").font.size = Pt(9)
    bfp.runs[0].font.color.rgb = P['acc2']

    doc.add_paragraph()

    # ── Dispo ─────────────────────────────────────────────────────────────────
    _section_header(doc, "DISPOSITION", P['acc'], P['bg1'])
    _schicht_tabelle(doc, data.get('dispo',[]), P, ist_dispo=True)

    doc.add_paragraph()

    # ── Betreuer ──────────────────────────────────────────────────────────────
    _section_header(doc, "BEHINDERTENBETREUER / STATION", P['acc'], P['bg1'])
    _schicht_tabelle(doc, data.get('betreuer',[]), P, ist_dispo=False)

    doc.add_paragraph()

    # ── Krankmeldungen ────────────────────────────────────────────────────────
    kranke = kranke_liste(data)
    if kranke:
        _section_header(doc, "KRANKMELDUNGEN", P['rot'], P['bg2'])
        kp = doc.add_paragraph()
        kp.add_run("Heute nicht im Dienst (K): " + ", ".join(kranke)).font.color.rgb = P['rot']

    doc.add_paragraph()

    # ── Medizinisches Verbrauchsmaterial ─────────────────────────────────────
    _section_header(doc, "MEDIZINISCHES VERBRAUCHSMATERIAL", P['acc'], P['bg1'])
    _material_tabelle(doc, P)

    pfad = os.path.join(ZIELORDNER, "D1_CommandCenter_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")

# ═══════════════════════════════════════════════════════════════════════════════
# VARIANTE 2 – SWISS CLEAN (weiß / schwarz / DRK-Rot)
# ═══════════════════════════════════════════════════════════════════════════════
def variante2_swiss_clean(data, im_einsatz=5, einsaetze=28, patienten=5, pax=42500):
    P = P2
    doc = Document(); _margins(doc, l=2.0, r=2.0); _add_logo_header(doc, P, ""); _add_footer(doc)

    # Minimalist-Titel
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = tp.add_run("Stärkemeldung und Einsätze")
    tr.font.size = Pt(26); tr.font.bold = True; tr.font.color.rgb = P['txt']
    _border_bottom(tp, "BE0000", "12")

    mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mr = mp.add_run("25. März 2026  ·  Erste-Hilfe-Station Flughafen Köln/Bonn")
    mr.font.size = Pt(10); mr.font.color.rgb = P['sub']

    doc.add_paragraph()

    # ── Vier Kennzahlen-Boxen ─────────────────────────────────────────────────
    sc = doc.add_table(rows=3, cols=4); sc.style = 'Table Grid'
    kz = [("EINSÄTZE", str(einsaetze), P['acc']),
          ("PAT. AUF STATION", str(patienten), P['acc']),
          ("PASSAGIERE", f"{pax:,}".replace(",","."), P['grau']),
          ("PERSONAL", str(_gesamt_personal(data)), P['acc'])]
    for i,(lbl,val,f) in enumerate(kz):
        hc=sc.rows[0].cells[i]; vc=sc.rows[1].cells[i]; fc=sc.rows[2].cells[i]
        _shd(hc,P['acc']); _shd(vc,rgb(0xFF,0xFF,0xFF)); _shd(fc,P['bg3'])
        _cell_txt(hc,lbl,bold=True,size=8,color=rgb(0xFF,0xFF,0xFF),align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(vc,val,bold=True,size=28,color=P['acc'],align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(fc,"",size=4,color=P['bg3'],align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Bulmor ────────────────────────────────────────────────────────────────
    bh = doc.add_paragraph()
    br = bh.add_run("BULMOR-FAHRZEUGSTATUS")
    br.font.size=Pt(11); br.font.bold=True; br.font.color.rgb=P['acc']
    _border_bottom(bh,"BE0000","6")

    bul = doc.add_table(rows=1, cols=5); bul.style = 'Table Grid'
    farbe, label = _bulmor_status(im_einsatz, P)
    for i in range(5):
        aktiv = (i+1) <= im_einsatz
        c1 = bul.rows[0].cells[i]
        _shd(c1, farbe if aktiv else RGB_HELL)
        sym = "●" if aktiv else "○"
        col = rgb(0xFF,0xFF,0xFF) if aktiv else P['grau']
        _cell_txt(c1, f"{sym}\nBulmor {i+1}", bold=True, size=10, color=col, align=WD_ALIGN_PARAGRAPH.CENTER)

    sp2 = doc.add_paragraph(); sp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sr2 = sp2.add_run(f"  {im_einsatz} / 5 im Einsatz  —  {label}")
    sr2.font.size=Pt(11); sr2.font.bold=True; sr2.font.color.rgb=farbe

    doc.add_paragraph()

    # ── Disposition & Betreuer nebeneinander (Tabelle)  ───────────────────────
    bh2 = doc.add_paragraph()
    bh2.add_run("EINSATZKRÄFTE").font.size = Pt(11)
    bh2.runs[0].font.bold = True; bh2.runs[0].font.color.rgb = P['acc']
    _border_bottom(bh2,"BE0000","6")

    # Dispo
    dp = doc.add_paragraph()
    dp.add_run("  Disposition:").font.bold = True
    dp.runs[0].font.size = Pt(10); dp.runs[0].font.color.rgb = P['txt']
    _schicht_tabelle(doc, data.get('dispo',[]), P, ist_dispo=True, mini=True)

    doc.add_paragraph()
    bp2 = doc.add_paragraph()
    bp2.add_run("  Behindertenbetreuer:").font.bold = True
    bp2.runs[0].font.size = Pt(10); bp2.runs[0].font.color.rgb = P['txt']
    _schicht_tabelle(doc, data.get('betreuer',[]), P, ist_dispo=False, mini=True)

    kranke = kranke_liste(data)
    if kranke:
        doc.add_paragraph()
        kp = doc.add_paragraph()
        kr = kp.add_run("⚕ Krank: " + ", ".join(kranke))
        kr.font.size=Pt(10); kr.font.color.rgb=P['acc']

    doc.add_paragraph()

    bh3 = doc.add_paragraph()
    bh3.add_run("MEDIZINISCHES VERBRAUCHSMATERIAL").font.size = Pt(11)
    bh3.runs[0].font.bold = True; bh3.runs[0].font.color.rgb = P['acc']
    _border_bottom(bh3,"BE0000","6")
    _material_tabelle(doc, P)

    pfad = os.path.join(ZIELORDNER, "D2_SwissClean_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")

# ═══════════════════════════════════════════════════════════════════════════════
# VARIANTE 3 – MIDNIGHT NAVY (dunkelblau / gold)
# ═══════════════════════════════════════════════════════════════════════════════
def variante3_midnight_navy(data, im_einsatz=3, einsaetze=28, patienten=5, pax=42500):
    P = P3
    doc = Document(); _margins(doc); _add_logo_header(doc, P, ""); _add_footer(doc)

    # Titel
    banner = doc.add_table(rows=2, cols=1)
    c1 = banner.rows[0].cells[0]; c2 = banner.rows[1].cells[0]
    _shd(c1, P['bg1']); _shd(c2, P['bg2'])
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ges = p1.add_run("Stärkemeldung und Einsätze")
    r_ges.font.size=Pt(24); r_ges.font.bold=True; r_ges.font.color.rgb=P['gold']
    p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("DRK Kreisverband Köln e.V.  ·  Erste-Hilfe-Station Flughafen Köln/Bonn  ·  25.03.2026")
    r2.font.size=Pt(9); r2.font.color.rgb=P['sub']

    doc.add_paragraph()

    # ── Scorecards ────────────────────────────────────────────────────────────
    sc = doc.add_table(rows=2, cols=4); sc.style = 'Table Grid'
    farbe_s, label_s = _bulmor_status(im_einsatz, P)
    kz=[("EINSÄTZE",str(einsaetze),P['gold']),
        ("PAT. AUF STATION",str(patienten),P['grn']),
        ("BULMOR",f"{im_einsatz}/5",farbe_s),
        ("PERSONAL",str(_gesamt_personal(data)),P['acc2'])]
    for i,(lbl,val,f) in enumerate(kz):
        hc=sc.rows[0].cells[i]; vc=sc.rows[1].cells[i]
        _shd(hc,P['bg3']); _shd(vc,P['bg2'])
        _cell_txt(hc,lbl,bold=True,size=8,color=P['gold'],align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(vc,val,bold=True,size=26,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Bulmor visuell ────────────────────────────────────────────────────────
    _section_header(doc, "BULMOR-FLOTTE  (5 Spezialfahrzeuge)", P['gold'], P['bg1'])
    bul = doc.add_table(rows=2, cols=5); bul.style = 'Table Grid'
    for i in range(5):
        aktiv = (i+1) <= im_einsatz
        c1 = bul.rows[0].cells[i]; c2 = bul.rows[1].cells[i]
        _shd(c1, farbe_s if aktiv else P['bg3'])
        _shd(c2, P['bg2'])
        sym = "▶" if aktiv else "▷"
        _cell_txt(c1, f"{sym} B{i+1}", bold=True, size=12, color=P['txt'], align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(c2, "EINSATZ" if aktiv else "BEREIT", size=8, color=farbe_s if aktiv else P['grau'],
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    sp3 = doc.add_paragraph(); sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sp3.add_run(f"◈  {im_einsatz} von 5 im Einsatz  —  {label_s}  ◈")
    r3.font.size=Pt(12); r3.font.bold=True; r3.font.color.rgb=farbe_s

    doc.add_paragraph()

    # ── Schichten – große Detail-Tabelle ─────────────────────────────────────
    _section_header(doc, "DIENSTPLAN – VOLLSTÄNDIGE BESETZUNG", P['gold'], P['bg1'])
    _schicht_tabelle_detailliert(doc, data, P)

    doc.add_paragraph()

    # ── Verbrauchsmaterial ────────────────────────────────────────────────────
    _section_header(doc, "MEDIZINPRODUKTE / VERBRAUCHSMATERIAL", P['gold'], P['bg1'])
    _material_tabelle(doc, P)

    pfad = os.path.join(ZIELORDNER, "D3_MidnightNavy_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")

# ═══════════════════════════════════════════════════════════════════════════════
# VARIANTE 4 – AIRPORT SIGNAL BOARD (dunkelgrau / grün)
# ═══════════════════════════════════════════════════════════════════════════════
def variante4_airport_board(data, im_einsatz=2, einsaetze=28, patienten=5, pax=42500):
    P = P4
    doc = Document(); _margins(doc); _add_logo_header(doc, P, ""); _add_footer(doc)

    # Titel im Airport-Stil
    banner = doc.add_table(rows=1, cols=1)
    bc = banner.rows[0].cells[0]; _shd(bc, P['bg1'])
    bp = bc.paragraphs[0]; bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("✈  STÄRKEMELDUNG UND EINSÄTZE  ✈")
    br.font.size=Pt(18); br.font.bold=True; br.font.color.rgb=P['acc']
    bp2 = bc.add_paragraph(); bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = bp2.add_run("ERSTE-HILFE-STATION FLUGHAFEN KÖLN/BONN  ·  DRK KREISVERBAND KÖLN  ·  25.03.2026")
    r4.font.size=Pt(8); r4.font.color.rgb=P['sub']

    doc.add_paragraph()

    # ── Große Anzeigetafel-Tabelle ────────────────────────────────────────────
    info = doc.add_table(rows=2, cols=5); info.style = 'Table Grid'
    farbe_b, label_b = _bulmor_status(im_einsatz, P)
    kz2=[("EINSÄTZE",str(einsaetze),P['acc']),
         ("PATIENTEN",str(patienten),P['gold']),
         ("PAX",f"{pax:,}".replace(",","."),P['grau']),
         ("PERSONAL",str(_gesamt_personal(data)),P['acc']),
         ("BULMOR",f"{im_einsatz}/5",farbe_b)]
    for i,(lbl,val,f) in enumerate(kz2):
        hc=info.rows[0].cells[i]; vc=info.rows[1].cells[i]
        _shd(hc,P['bg2']); _shd(vc,P['bg3'])
        _cell_txt(hc,lbl,bold=True,size=8,color=P['acc'],align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(vc,val,bold=True,size=22,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Bulmor-Panel ──────────────────────────────────────────────────────────
    _section_header(doc, "BULMOR FAHRZEUGSTATUS", P['acc'], P['bg1'])
    bul = doc.add_table(rows=1, cols=5); bul.style = 'Table Grid'
    for i in range(5):
        aktiv = (i+1) <= im_einsatz
        c1 = bul.rows[0].cells[i]
        _shd(c1, farbe_b if aktiv else P['bg2'])
        sym = "◉" if aktiv else "○"
        txt = f"{sym}\nBulmor {i+1}\n{'EINSATZ' if aktiv else 'BEREIT'}"
        _cell_txt(c1, txt, bold=aktiv, size=10,
                  color=P['txt'] if aktiv else P['grau'],
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    sp4 = doc.add_paragraph(); sp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = sp4.add_run(f"STATUS: {im_einsatz}/5 BULMOR IM EINSATZ  —  {label_b}")
    r5.font.size=Pt(12); r5.font.bold=True; r5.font.color.rgb=farbe_b

    bdf2 = bulmorfahrer_liste(data)
    bfp2 = doc.add_paragraph()
    bfp2.add_run(f"Bulmor-Fahrer: {', '.join(bdf2)}").font.size=Pt(9)
    bfp2.runs[0].font.color.rgb = P['sub']

    doc.add_paragraph()

    # ── Abflugbrett-Stil Schichten ────────────────────────────────────────────
    _section_header(doc, "SCHICHTBESETZUNG", P['acc'], P['bg1'])
    alle = alle_betreuer_schichten(data)
    dispo_g = gruppiere_nach_zeit(data.get('dispo',[]), ist_dispo=True)

    st = doc.add_table(rows=1 + len(dispo_g) + len(alle), cols=3)
    st.style = 'Table Grid'
    # Kopf
    for ci, txt in enumerate(["SCHICHT", "MITARBEITER", "ANZ."]):
        hc = st.rows[0].cells[ci]
        _shd(hc, P['bg1'])
        _cell_txt(hc, txt, bold=True, size=10, color=P['acc'], align=WD_ALIGN_PARAGRAPH.CENTER)

    ri = 1
    for key, namen in dispo_g.items():
        row = st.rows[ri]
        _shd(row.cells[0], rgb(0x14,0x3D,0x26))
        _shd(row.cells[1], P['bg3'])
        _shd(row.cells[2], P['bg3'])
        _cell_txt(row.cells[0], f"DISPO {key}", bold=True, size=9, color=P['acc'], align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(row.cells[1], " / ".join(namen), size=9, color=P['txt'])
        _cell_txt(row.cells[2], str(len(namen)), bold=True, size=10, color=P['acc'], align=WD_ALIGN_PARAGRAPH.CENTER)
        ri += 1
    for key, namen in alle.items():
        row = st.rows[ri]
        bg = P['bg2'] if ri%2==0 else P['bg3']
        _shd(row.cells[0],bg); _shd(row.cells[1],bg); _shd(row.cells[2],bg)
        _cell_txt(row.cells[0], key, bold=True, size=9, color=P['acc2'], align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(row.cells[1], " / ".join(namen), size=9, color=P['txt'])
        _cell_txt(row.cells[2], str(len(namen)), bold=True,size=10,color=P['gold'],align=WD_ALIGN_PARAGRAPH.CENTER)
        ri += 1

    kranke = kranke_liste(data)
    if kranke:
        doc.add_paragraph()
        kp2 = doc.add_paragraph()
        r6 = kp2.add_run("⚠ KRANK: " + ", ".join(kranke))
        r6.font.size=Pt(10); r6.font.bold=True; r6.font.color.rgb=P['rot']

    doc.add_paragraph()
    _section_header(doc, "MEDIZINPRODUKTE – VERBRAUCHSMATERIAL", P['acc'], P['bg1'])
    _material_tabelle(doc, P)

    pfad = os.path.join(ZIELORDNER, "D4_AirportBoard_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")

# ═══════════════════════════════════════════════════════════════════════════════
# VARIANTE 5 – DRK ROT / CLEAN WHITE (vollständige Besetzung alle Schichten)
# ═══════════════════════════════════════════════════════════════════════════════
def variante5_drk_vollbesetzung(data, im_einsatz=5, einsaetze=28, patienten=5, pax=42500):
    P = P5
    doc = Document(); _margins(doc); _add_logo_header(doc, P, ""); _add_footer(doc)

    # Header-Banner
    banner = doc.add_table(rows=1, cols=1)
    bc = banner.rows[0].cells[0]; _shd(bc, P['bg1'])
    bp = bc.paragraphs[0]; bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("Stärkemeldung und Einsätze")
    br.font.size=Pt(24); br.font.bold=True; br.font.color.rgb=P['txt']
    bp2 = bc.add_paragraph(); bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = bp2.add_run("VOLLSTÄNDIGE SCHICHTBESETZUNG  ·  25.03.2026  ·  DRK Köln")
    r2.font.size=Pt(10); r2.font.color.rgb=P['sub']

    doc.add_paragraph()

    # ── Kennzahlen (hell) ────────────────────────────────────────────────────
    sc = doc.add_table(rows=2, cols=4); sc.style = 'Table Grid'
    kz=[("EINSÄTZE",str(einsaetze),rgb(0x1F,0x49,0x7D)),
        ("PAT. AUF STATION",str(patienten),rgb(0x10,0x7E,0x3E)),
        ("PASSAGIERE",f"{pax:,}".replace(",","."),rgb(0x55,0x55,0x55)),
        ("PERSONAL",str(_gesamt_personal(data)),rgb(0x1F,0x49,0x7D))]
    for i,(lbl,val,f) in enumerate(kz):
        hc=sc.rows[0].cells[i]; vc=sc.rows[1].cells[i]
        _shd(hc,P['bg1']); _shd(vc,rgb(0xFF,0xFF,0xFF))
        _cell_txt(hc,lbl,bold=True,size=8,color=P['txt'],align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(vc,val,bold=True,size=26,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Bulmor (alle grün = Vollbesetzung) ────────────────────────────────────
    bh = doc.add_paragraph()
    br3 = bh.add_run("BULMOR-FAHRZEUGSTATUS  (5 / 5 verfügbar)")
    br3.font.size=Pt(12); br3.font.bold=True; br3.font.color.rgb=P['bg1']
    _border_bottom(bh, "BE0000", "8")
    bul = doc.add_table(rows=2, cols=5); bul.style = 'Table Grid'
    farbe_v, label_v = _bulmor_status(im_einsatz, P)
    for i in range(5):
        aktiv = (i+1) <= im_einsatz
        c1=bul.rows[0].cells[i]; c2=bul.rows[1].cells[i]
        col_bg = P['grn'] if aktiv else rgb(0xCC,0xCC,0xCC)
        _shd(c1, col_bg); _shd(c2, rgb(0xF5,0xF5,0xF5))
        sym = "✔" if aktiv else "○"
        _cell_txt(c1, f"{sym}  Bulmor {i+1}", bold=True, size=11,
                  color=rgb(0xFF,0xFF,0xFF) if aktiv else rgb(0x88,0x88,0x88),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(c2, "IM EINSATZ" if aktiv else "RESERVE", size=8,
                  color=P['grn'] if aktiv else rgb(0x88,0x88,0x88),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    sp5 = doc.add_paragraph(); sp5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r7 = sp5.add_run(f"Alle {im_einsatz} Bulmor im Einsatz  —  {label_v}  🚑")
    r7.font.size=Pt(13); r7.font.bold=True; r7.font.color.rgb=P['grn']

    bdf3 = bulmorfahrer_liste(data)
    doc.add_paragraph().add_run(f"Bulmor-Fahrer: {', '.join(bdf3)}").font.size=Pt(9)

    doc.add_paragraph()

    # ── Komplette Besetzungstabelle ───────────────────────────────────────────
    bh2 = doc.add_paragraph()
    bh2.add_run("VOLLSTÄNDIGE SCHICHT-BESETZUNG").font.size=Pt(12)
    bh2.runs[0].font.bold=True; bh2.runs[0].font.color.rgb=P['bg1']
    _border_bottom(bh2,"BE0000","8")

    _schicht_tabelle_detailliert(doc, data, P)

    kranke = kranke_liste(data)
    if kranke:
        doc.add_paragraph()
        kp3 = doc.add_paragraph()
        r8 = kp3.add_run("Krankmeldungen: " + ", ".join(kranke))
        r8.font.size=Pt(10); r8.font.bold=True; r8.font.color.rgb=P['rot']

    doc.add_paragraph()

    # ── Material ──────────────────────────────────────────────────────────────
    bh3 = doc.add_paragraph()
    bh3.add_run("MEDIZINPRODUKTE / VERBRAUCHSMATERIAL").font.size=Pt(12)
    bh3.runs[0].font.bold=True; bh3.runs[0].font.color.rgb=P['bg1']
    _border_bottom(bh3,"BE0000","8")
    _material_tabelle(doc, P)

    pfad = os.path.join(ZIELORDNER, "D5_DRK_Vollbesetzung_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")

# ═══════════════════════════════════════════════════════════════════════════════
# VARIANTE 6 – KOMPAKTER TAGESBERICHT (alle Sektionen auf 1 Seite)
# ═══════════════════════════════════════════════════════════════════════════════
def variante6_tagesbericht_kompakt(data, im_einsatz=3, einsaetze=28, patienten=5, pax=42500):
    """Alles auf möglichst wenig Raum, für schnellen Überblick."""
    P = P2  # Swiss Clean Basis
    AKZENT = rgb(0x1F,0x49,0x7D)
    doc = Document(); _margins(doc, l=1.5, r=1.5, t=2.5, b=1.8)
    _add_logo_header(doc, P, ""); _add_footer(doc)

    # Titelzeile
    tt = doc.add_table(rows=1, cols=2); tt.style = 'Table Grid'
    lc = tt.rows[0].cells[0]; rc = tt.rows[0].cells[1]
    _shd(lc, AKZENT); _shd(rc, P['bg2'])
    lp = lc.paragraphs[0]
    lr = lp.add_run("Stärkemeldung und Einsätze")
    lr.font.size=Pt(16); lr.font.bold=True; lr.font.color.rgb=rgb(0xFF,0xFF,0xFF)
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(f"25. März 2026\n07:45 Uhr")
    rr.font.size=Pt(11); rr.font.bold=True; rr.font.color.rgb=AKZENT

    doc.add_paragraph()

    # ── Kompakte Kennzahlen-Reihe ─────────────────────────────────────────────
    kn = doc.add_table(rows=1, cols=5); kn.style = 'Table Grid'
    farbe_k, label_k = _bulmor_status(im_einsatz, P)
    kzk=[("Einsätze",str(einsaetze),P['acc']),
         ("Patienten",str(patienten),P['grn']),
         ("PAX",f"{pax:,}".replace(",","."),rgb(0x55,0x55,0x55)),
         ("Personal",str(_gesamt_personal(data)),AKZENT),
         (f"Bulmor",f"{im_einsatz}/5",farbe_k)]
    for i,(lbl,val,f) in enumerate(kzk):
        c = kn.rows[0].cells[i]
        _shd(c, P['bg3'])
        _cell_txt(c, f"{lbl}\n{val}", bold=True, size=11, color=f, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Bulmor-Zeile kompakt ──────────────────────────────────────────────────
    bul = doc.add_table(rows=1, cols=5); bul.style = 'Table Grid'
    for i in range(5):
        aktiv = (i+1) <= im_einsatz
        c = bul.rows[0].cells[i]
        _shd(c, farbe_k if aktiv else P['bg3'])
        sym = "◉" if aktiv else "○"
        _cell_txt(c, f"{sym} B{i+1}", bold=True, size=11,
                  color=rgb(0xFF,0xFF,0xFF) if aktiv else P['grau'],
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    bfp = doc.add_paragraph()
    bfp.add_run(f"  {label_k} | Fahrer: {', '.join(bulmorfahrer_liste(data))}").font.size=Pt(9)
    bfp.runs[0].font.color.rgb = farbe_k

    doc.add_paragraph()

    # ── Alle Schichten kompakt ────────────────────────────────────────────────
    lp2 = doc.add_paragraph()
    lp2.add_run("Diensteinteilung").font.bold=True
    lp2.runs[0].font.size=Pt(11); lp2.runs[0].font.color.rgb=AKZENT
    _border_bottom(lp2,"1F497D","6")

    alle_p = data.get('dispo',[]) + data.get('betreuer',[])
    alle_p = [p for p in alle_p if p.get('ist_krank') not in (True,'True')]
    gruppen = defaultdict(list)
    for p in alle_p:
        s=(p.get('start_zeit') or '')[:5]; e=(p.get('end_zeit') or '')[:5]
        ist_d = p.get('ist_dispo') in (True,'True')
        k = f"{'★ DISPO ' if ist_d else ''}{s} – {e}"
        gruppen[k].append(p.get('anzeigename',''))
    dienstplan_tbl = doc.add_table(rows=1+len(gruppen), cols=3); dienstplan_tbl.style='Table Grid'
    for ci, h in enumerate(["Schicht","Mitarbeiter","Anz."]):
        hc = dienstplan_tbl.rows[0].cells[ci]
        _shd(hc, AKZENT)
        _cell_txt(hc, h, bold=True, size=9, color=rgb(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, (key, namen) in enumerate(sorted(gruppen.items()),1):
        row = dienstplan_tbl.rows[ri]
        bg = P['bg2'] if ri%2==0 else rgb(0xFF,0xFF,0xFF)
        _shd(row.cells[0],bg); _shd(row.cells[1],bg); _shd(row.cells[2],bg)
        _cell_txt(row.cells[0], key, bold="DISPO" in key, size=9, color=AKZENT if "DISPO" in key else P['txt'])
        _cell_txt(row.cells[1], " / ".join(namen), size=9, color=P['txt'])
        _cell_txt(row.cells[2], str(len(namen)), bold=True, size=10, color=P['acc'], align=WD_ALIGN_PARAGRAPH.CENTER)

    kranke = kranke_liste(data)
    if kranke:
        kp4 = doc.add_paragraph()
        r9 = kp4.add_run("K: " + ", ".join(kranke) + " (krank)")
        r9.font.size=Pt(9); r9.font.italic=True; r9.font.color.rgb=P['acc']

    doc.add_paragraph()

    lp3 = doc.add_paragraph()
    lp3.add_run("Verbrauchsmaterial (Auswahl)").font.bold=True
    lp3.runs[0].font.size=Pt(11); lp3.runs[0].font.color.rgb=AKZENT
    _border_bottom(lp3,"1F497D","6")
    _material_tabelle(doc, P)

    pfad = os.path.join(ZIELORDNER, "D6_TagesBericht_Kompakt_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")

# ═══════════════════════════════════════════════════════════════════════════════
# HILFSFUNKTIONEN
# ═══════════════════════════════════════════════════════════════════════════════
RGB_HELL = rgb(0xF5,0xF5,0xF5)

def _gesamt_personal(data):
    alle = [p for p in data.get('betreuer',[]) + data.get('dispo',[])
            if p.get('ist_krank') not in (True,'True')]
    return len(alle)

def _section_header(doc, text, fg: RGBColor, bg: RGBColor):
    tbl = doc.add_table(rows=1, cols=1)
    c = tbl.rows[0].cells[0]; _shd(c, bg)
    p = c.paragraphs[0]
    r = p.add_run(f"  {text}")
    r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=fg

def _schicht_tabelle(doc, personen, P, ist_dispo=False, mini=False):
    gruppen = gruppiere_nach_zeit(personen, ist_dispo=ist_dispo)
    if not gruppen:
        doc.add_paragraph().add_run("(keine Einträge)").font.size=Pt(9)
        return
    tbl = doc.add_table(rows=1+len(gruppen), cols=3 if not mini else 2)
    tbl.style = 'Table Grid'
    hdrs = ["Schicht","Mitarbeiter","Anz."] if not mini else ["Schicht","Mitarbeiter"]
    for ci, h in enumerate(hdrs):
        hc = tbl.rows[0].cells[ci]
        _shd(hc, P.get('bg2', RGB_HELL))
        _cell_txt(hc, h, bold=True, size=9, color=P.get('acc',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri,(key,namen) in enumerate(gruppen.items(),1):
        row = tbl.rows[ri]
        bg = RGB_HELL if ri%2==1 else rgb(0xFF,0xFF,0xFF)
        for ci in range(len(hdrs)): _shd(row.cells[ci], bg)
        _cell_txt(row.cells[0], key, bold=True, size=9, color=P.get('txt',rgb(0,0,0)))
        _cell_txt(row.cells[1], " / ".join(namen), size=9, color=P.get('txt',rgb(0,0,0)))
        if not mini:
            _cell_txt(row.cells[2], str(len(namen)), bold=True, size=10,
                      color=P.get('acc',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)

def _schicht_tabelle_detailliert(doc, data, P):
    """Alle Betreuer + Dispo, je mit Kategorie, Bulmor-Flag."""
    alle = []
    for p in data.get('dispo',[]):
        if p.get('ist_krank') not in (True,'True'): alle.append(('DISPO', p))
    for p in data.get('betreuer',[]):
        if p.get('ist_krank') not in (True,'True'): alle.append(('BET', p))
    alle.sort(key=lambda x: (x[1].get('start_zeit') or 'ZZZZ'))

    if not alle:
        doc.add_paragraph().add_run("(keine Einträge)").font.size=Pt(9)
        return

    tbl = doc.add_table(rows=1+len(alle), cols=5)
    tbl.style = 'Table Grid'
    for ci, h in enumerate(["Rolle","Name","Schicht","Kat.","Bulmor"]):
        c = tbl.rows[0].cells[ci]
        _shd(c, P.get('bg2',RGB_HELL))
        _cell_txt(c, h, bold=True, size=9, color=P.get('acc',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, (rolle, p) in enumerate(alle, 1):
        row = tbl.rows[ri]
        bg = RGB_HELL if ri%2==1 else rgb(0xFF,0xFF,0xFF)
        for ci in range(5): _shd(row.cells[ci], bg)
        _cell_txt(row.cells[0], rolle, bold=True, size=8,
                  color=P.get('acc',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(row.cells[1], p.get('anzeigename',''), size=9, color=P.get('txt',rgb(0,0,0)))
        s = (p.get('start_zeit') or '')[:5]; e = (p.get('end_zeit') or '')[:5]
        _cell_txt(row.cells[2], f"{s} – {e}", size=9, color=P.get('txt',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(row.cells[3], p.get('dienst_kategorie',''), bold=True, size=9,
                  color=P.get('acc',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)
        ist_bul = p.get('ist_bulmorfahrer') in (True, 'True')
        _cell_txt(row.cells[4], "✔" if ist_bul else "", bold=True, size=11,
                  color=P.get('grn',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)

def _material_tabelle(doc, P):
    tbl = doc.add_table(rows=1+len(MEDIZIN_MATERIAL), cols=4)
    tbl.style = 'Table Grid'
    for ci, h in enumerate(["Material","Einheit","Soll-Bestand","Mind.-Bestand"]):
        c = tbl.rows[0].cells[ci]
        _shd(c, P.get('bg2',RGB_HELL))
        _cell_txt(c, h, bold=True, size=9, color=P.get('acc',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, (mat, eh, soll, mind) in enumerate(MEDIZIN_MATERIAL, 1):
        row = tbl.rows[ri]
        bg = RGB_HELL if ri%2==1 else rgb(0xFF,0xFF,0xFF)
        for ci in range(4): _shd(row.cells[ci], bg)
        _cell_txt(row.cells[0], mat, size=9, color=P.get('txt',rgb(0,0,0)))
        _cell_txt(row.cells[1], eh, size=9, color=P.get('txt',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(row.cells[2], soll, bold=True, size=9, color=P.get('txt',rgb(0,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_txt(row.cells[3], mind, bold=True, size=9,
                  color=P.get('rot',rgb(0xBE,0,0)), align=WD_ALIGN_PARAGRAPH.CENTER)

# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"Lade Dienstplan aus Excel ...")
    data = lade_dienstplan()
    if not data:
        print("  Fallback auf Beispieldaten ...")
        data = {"betreuer": [], "dispo": [], "kranke": []}
    else:
        bet = len(data.get('betreuer',[]))
        dis = len(data.get('dispo',[]))
        kra = len(data.get('kranke',[]))
        print(f"  Betreuer: {bet}  |  Dispo: {dis}  |  Krank: {kra}")

    print(f"\nZielordner: {ZIELORDNER}\n")
    print("Erstelle Designs ...")
    variante1_command_center(data,  im_einsatz=4, einsaetze=28, patienten=5, pax=42500)
    variante2_swiss_clean(data,     im_einsatz=5, einsaetze=28, patienten=5, pax=42500)
    variante3_midnight_navy(data,   im_einsatz=3, einsaetze=28, patienten=5, pax=42500)
    variante4_airport_board(data,   im_einsatz=2, einsaetze=28, patienten=5, pax=42500)
    variante5_drk_vollbesetzung(data,im_einsatz=5,einsaetze=28, patienten=5, pax=42500)
    variante6_tagesbericht_kompakt(data,im_einsatz=3,einsaetze=28,patienten=5,pax=42500)
    print(f"\n✓ Alle 6 Beispiele gespeichert in:\n  {ZIELORDNER}")
