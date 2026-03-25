# -*- coding: utf-8 -*-
"""
P9-Split-Stil – 6 Farbvarianten (PDF) + 2 Word-Versionen
Alle Mitarbeiter werden vollständig aufgelistet.
"""
import os, sys
from pathlib import Path
from collections import defaultdict

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# ── PDF-Importe ────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor, white, black
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

# ── Word-Importe ───────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W, H = A4

# ── Pfade ─────────────────────────────────────────────────────────────────────
_OD = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
ZIEL = os.path.join(_OD, "Desktop", "bei") if os.path.exists(_OD) else r"C:\Temp\bei"
os.makedirs(ZIEL, exist_ok=True)
LOGO = Path(os.path.dirname(os.path.abspath(__file__))) / "Daten" / "Email" / "Logo.jpg"
EXCEL = (
    _OD + r"\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26"
    r"\04_Tagesdienstpläne\03_März\25.03.2026.xlsx"
)
DATUM = "25.03.2026"; UHRZEIT = "07:45 Uhr"
STATION = "Erste-Hilfe-Station · Flughafen Köln/Bonn"

def h(hex_): return HexColor(f"#{hex_.lstrip('#')}")

ROT_WARN = h("FF3333"); GRN_OK = h("10A050"); ORG_WARN = h("E07800")

def _bul_col(n):
    if n <= 2:   return ROT_WARN, "KRITISCH"
    elif n == 3: return ORG_WARN, "EINGESCHRÄNKT"
    else:        return GRN_OK,   "VOLLSTÄNDIG"

# ── Daten laden ────────────────────────────────────────────────────────────────
def lade():
    try:
        from functions.dienstplan_parser import DienstplanParser
        r = DienstplanParser(EXCEL, alle_anzeigen=True).parse()
        if r.get("success") in (True, "True"): return r
    except Exception as e:
        print(f"  [Parser] {e}")
    return {"betreuer": [], "dispo": [], "kranke": []}

def _alle_betreuer(d):
    return [p.get('anzeigename','').strip()
            for p in d.get('betreuer', [])
            if p.get('ist_krank') not in (True, 'True') and p.get('anzeigename','').strip()]

def _alle_dispo(d):
    return [p.get('anzeigename','').strip()
            for p in d.get('dispo', [])
            if p.get('ist_krank') not in (True, 'True') and p.get('anzeigename','').strip()]

def _kranke(d):
    return [p.get('anzeigename','').strip()
            for p in d.get('betreuer', []) + d.get('dispo', [])
            if p.get('ist_krank') in (True, 'True') and p.get('anzeigename','').strip()]

def _bulfhr(d):
    return [p.get('anzeigename','').strip()
            for p in d.get('betreuer', []) + d.get('dispo', [])
            if p.get('ist_bulmorfahrer') in (True, 'True') and p.get('anzeigename','').strip()]

def _gruppen(d, typ='dispo'):
    gru = defaultdict(list)
    for p in d.get(typ, []):
        if p.get('ist_krank') in (True, 'True'): continue
        s = (p.get('start_zeit') or '')[:5]; e = (p.get('end_zeit') or '')[:5]
        gru[f"{s}–{e}"].append(p.get('anzeigename','').strip())
    return dict(sorted(gru.items()))

# ── PDF Zeichenhilfen ──────────────────────────────────────────────────────────
def rect(cv, x, yt, w_, ht, color):
    cv.setFillColor(color); cv.rect(x, H - yt - ht, w_, ht, fill=1, stroke=0)

def ln(cv, x1, yt1, x2, yt2, color, lw=0.8):
    cv.setStrokeColor(color); cv.setLineWidth(lw); cv.line(x1, H-yt1, x2, H-yt2)

def t(cv, x, yt, text, font="Helvetica", size=9, color=black, align="left"):
    cv.setFont(font, size); cv.setFillColor(color)
    if align == "center": cv.drawCentredString(x, H - yt, str(text))
    elif align == "right": cv.drawRightString(x, H - yt, str(text))
    else: cv.drawString(x, H - yt, str(text))

def logo_draw(cv, x, yt, w_=55, ht=52):
    if LOGO.exists():
        try: cv.drawImage(ImageReader(str(LOGO)), x, H-yt-ht, w_, ht,
                          mask='auto', preserveAspectRatio=True)
        except: pass

def kreis(cv, cx, cy_from_top, r, fill, stroke_c=None, lw=0):
    cv.setFillColor(fill)
    if stroke_c:
        cv.setStrokeColor(stroke_c); cv.setLineWidth(lw)
        cv.circle(cx, H - cy_from_top, r, fill=1, stroke=1)
    else:
        cv.circle(cx, H - cy_from_top, r, fill=1, stroke=0)


# ═══════════════════════════════════════════════════════════════════════════════
# KERNFUNKTION: geteilte Seite (Split-Panel) – links Farbe, rechts Weiß
# Alle Mitarbeiter vollständig aufgelistet
# ═══════════════════════════════════════════════════════════════════════════════
def draw_split(cv, links_bg, links_dunkel, links_hell, akzent, akzent2, titel_farbe,
               data, bul=5, einz=28, pat=5, pax=42500,
               scheme_name=""):
    """
    links_bg    : Hauptfarbe linkes Panel
    links_dunkel: etwas dunklere Version (Tiles)
    links_hell  : sehr helle Version (Subtitle-Text)
    akzent      : Akzentfarbe (Kreise aktiv, Trennlinien)
    akzent2     : zweiter Akzent (Überschriften rechts)
    titel_farbe : Farbe für den großen Titel links
    """
    fc_bul, lbl_bul = _bul_col(bul)
    betreuer   = _alle_betreuer(data)
    dispo_lst  = _alle_dispo(data)
    kranke_lst = _kranke(data)
    bulfhr_lst = _bulfhr(data)
    dispo_grp  = _gruppen(data, 'dispo')
    WEISS = white; DUNKEL = h("111111"); GRAU = h("555555")

    rect(cv, 0, 0, W, H, WEISS)

    # ── Linkes Panel (42% der Breite) ─────────────────────────────────────────
    LW = W * 0.42
    rect(cv, 0, 0, LW, H, links_bg)

    # Weicher Übergang-Akzent (schmales dunkleres Rechteck am rechten Rand des Panels)
    rect(cv, LW - 4, 0, 4, H, links_dunkel)

    # Logo
    logo_draw(cv, 16, 14, w_=55, ht=52)

    # Titel
    t(cv, LW/2, 78, "Stärke-", "Helvetica-Bold", 22, titel_farbe, "center")
    t(cv, LW/2, 99, "meldung", "Helvetica-Bold", 22, titel_farbe, "center")
    t(cv, LW/2, 118, "& Einsätze", "Helvetica-Bold", 15, akzent, "center")
    ln(cv, 18, 126, LW-18, 126, h("FFFFFF"), 1.2)
    t(cv, LW/2, 140, DATUM, "Helvetica-Bold", 10, links_hell, "center")
    t(cv, LW/2, 153, UHRZEIT, "Helvetica", 8.5, h("BBBBBB"), "center")

    # Kennzahlen-Kacheln (links)
    kacheln = [
        ("Einsätze", str(einz), akzent),
        ("Patienten", str(pat), h("AAFFCC")),
        ("PAX", f"{pax:,}".replace(",", "."), links_hell),
        ("Personal", str(len(betreuer)+len(dispo_lst)), links_hell),
    ]
    for i, (lbl, val, tc) in enumerate(kacheln):
        ky = 162 + i * 56
        rect(cv, 14, ky, LW - 28, 48, links_dunkel)
        ln(cv, 14, ky, 14+4, ky, akzent, 3)
        t(cv, 26, ky + 14, lbl, "Helvetica", 7.5, links_hell)
        t(cv, 26, ky + 40, val, "Helvetica-Bold", 22, tc)

    # Bulmor-Status
    b_yt = 162 + 4 * 56 + 8
    t(cv, LW/2, b_yt, "BULMOR", "Helvetica-Bold", 9.5, WEISS, "center")
    rect(cv, 14, b_yt + 5, LW - 28, 24, links_dunkel)
    fc2, lbl2 = _bul_col(bul)
    t(cv, LW/2, b_yt + 22, f"{bul}/5  ·  {lbl2}", "Helvetica-Bold", 9, fc2, "center")
    # 5 Kreise
    for i in range(5):
        aktiv = (i + 1) <= bul
        cx_ = 18 + i * (LW - 36) / 4
        ky_k = b_yt + 44
        kreis(cv, cx_, ky_k, 13, fc2 if aktiv else links_dunkel,
              WEISS if aktiv else None, 1.2)
        t(cv, cx_, ky_k + 5, f"B{i+1}", "Helvetica-Bold", 7,
          WEISS if aktiv else h("777777"), "center")

    # Bulmor-Fahrer
    if bulfhr_lst:
        fyt = b_yt + 64
        t(cv, LW/2, fyt, "Fahrer:", "Helvetica-Bold", 7.5, links_hell, "center")
        for j, n in enumerate(bulfhr_lst[:5]):
            t(cv, LW/2, fyt + 11 + j * 12, n, "Helvetica", 7.5, h("BBBBBB"), "center")

    # Krankmeldungen (links unten)
    if kranke_lst:
        ln(cv, 14, H - 60, LW - 14, H - 60, h("FF4444"), 0.8)
        t(cv, LW/2, H - 45, "KRANK", "Helvetica-Bold", 7.5, ROT_WARN, "center")
        for j, kname in enumerate(kranke_lst[:4]):
            t(cv, LW/2, H - 33 + j * 11, kname, "Helvetica", 7.5, h("FFAAAA"), "center")

    # ── Rechtes Panel (58% der Breite, weiß) ──────────────────────────────────
    RX = LW + 14  # Startpunkt X rechts
    RW = W - RX - 14  # Breite rechts

    # Überschrift rechts
    rect(cv, LW, 0, W - LW, 72, links_bg)
    t(cv, RX, 22, STATION, "Helvetica-Bold", 9, links_hell)
    t(cv, RX, 36, f"DRK Kreisverband Köln e.V.", "Helvetica", 8, h("BBBBBB"))
    t(cv, W - 16, 22, DATUM, "Helvetica-Bold", 10, akzent, "right")
    t(cv, W - 16, 36, UHRZEIT, "Helvetica", 8.5, links_hell, "right")

    # ── ALLE BETREUER ──────────────────────────────────────────────────────────
    ay = 82
    rect(cv, LW, ay, W - LW, 18, akzent2)
    t(cv, RX, ay + 13, f"ALLE BETREUER ({len(betreuer)} Personen)", "Helvetica-Bold", 9, WEISS)

    ay += 22
    # 2-Spaltige Namensliste
    half = (len(betreuer) + 1) // 2
    col_w = RW / 2
    for i, name in enumerate(betreuer):
        row = i % half; col = i // half
        bx_ = RX + col * col_w; by_ = ay + row * 13.5
        if by_ > H - 120: break
        bg = h("F5F5F5") if row % 2 == 0 else WEISS
        rect(cv, bx_, by_, col_w - 2, 12.5, bg)
        t(cv, bx_ + 4, by_ + 9.5, name, "Helvetica", 7.5, DUNKEL)
    ay += max(half, 1) * 13.5 + 6

    # ── ALLE DISPO ────────────────────────────────────────────────────────────
    if ay < H - 100:
        rect(cv, LW, ay, W - LW, 18, links_dunkel)
        t(cv, RX, ay + 13, f"DISPOSITION ({len(dispo_lst)} Personen)", "Helvetica-Bold", 9, WEISS)
        ay += 22
        for i, (zeit, namen) in enumerate(dispo_grp.items()):
            if ay > H - 44: break
            bg = h("EEF3FF") if i % 2 == 0 else WEISS
            rect(cv, LW, ay, W - LW, 21, bg)
            t(cv, RX, ay + 9, zeit, "Helvetica-Bold", 8, akzent2)
            t(cv, RX, ay + 18, ", ".join(namen), "Helvetica", 7.5, DUNKEL)
            ay += 22

    # Footer
    rect(cv, 0, H - 18, W, 18, links_dunkel)
    ln(cv, 0, H - 18, W, H - 18, akzent, 1.5)
    t(cv, W / 2, H - 6, "DRK Köln  ·  +49 2203 40-2323  ·  flughafen@drk-koeln.de",
      "Helvetica", 7.5, WEISS, "center")


# ═══════════════════════════════════════════════════════════════════════════════
# 6 FARBSCHEMATA
# ═══════════════════════════════════════════════════════════════════════════════

SCHEMES = [
    # (Name, bg, dunkel, hell, akzent, akzent2, titel)
    ("S1_Smaragd",   "1B6B45", "0F4D32", "D4F5E2", "F5C518", "1B6B45", "FFFFFF"),
    ("S2_Burgund",   "7B1A3A", "590F28", "FFD6E0", "FFB347", "7B1A3A", "FFFFFF"),
    ("S3_OzeanBlau", "1A3460", "0F1F3C", "C8DAFF", "00C8FF", "1A3460", "FFFFFF"),
    ("S4_Schiefer",  "2D4A6B", "1A2D45", "D0E8FF", "E8603A", "2D4A6B", "FFFFFF"),
    ("S5_Violett",   "4A1E7E", "32145A", "EDD6FF", "00DDB8", "4A1E7E", "FFFFFF"),
    ("S6_Anthrazit", "2C3440", "1A2028", "D8E8F0", "A8D040", "2C3440", "FFFFFF"),
]


def erstelle_alle_pdfs(data, bul=5, einz=28, pat=5, pax=42500):
    for (name, bg, dunkel, hell, akz, akz2, titel) in SCHEMES:
        outpath = os.path.join(ZIEL, f"{name}_25032026.pdf")
        cv_ = canvas.Canvas(outpath, pagesize=A4)
        draw_split(cv_,
                   links_bg=h(bg), links_dunkel=h(dunkel), links_hell=h(hell),
                   akzent=h(akz), akzent2=h(akz2), titel_farbe=h(titel),
                   data=data, bul=bul, einz=einz, pat=pat, pax=pax,
                   scheme_name=name)
        cv_.save()
        print(f"[OK] {name}_25032026.pdf")


# ═══════════════════════════════════════════════════════════════════════════════
# WORD-VERSION: Split-Layout mit python-docx
# ═══════════════════════════════════════════════════════════════════════════════

def _rgb(hex_str):
    hx = hex_str.lstrip("#")
    return RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#').upper())
    tcPr.append(shd)

def _no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def _par(cell, text, bold=False, size=9, color="000000", align="left", space_after=0, space_before=0):
    p = cell.add_paragraph()
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    return p

def _mini_tbl(container_cell, rows, col_widths_cm):
    """Erstellt eine verschachtelte Tabelle innerhalb einer Zelle"""
    tbl = container_cell.add_table(rows=len(rows), cols=len(col_widths_cm))
    tbl.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, (text, bold, size, fg, bg) in enumerate(row_data):
            cell = tbl.cell(i, j)
            if bg: _set_cell_bg(cell, bg)
            _no_border(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.bold = bold; run.font.size = Pt(size)
            run.font.color.rgb = _rgb(fg)
    # Spaltenbreiten
    for i in range(len(rows)):
        for j, w in enumerate(col_widths_cm):
            tbl.cell(i, j).width = Cm(w)
    return tbl

def erstelle_word(data, scheme_idx=0, bul=5, einz=28, pat=5, pax=42500):
    (name, bg, dunkel, hell, akz, akz2, titel) = SCHEMES[scheme_idx]
    betreuer   = _alle_betreuer(data)
    dispo_lst  = _alle_dispo(data)
    kranke_lst = _kranke(data)
    bulfhr_lst = _bulfhr(data)
    dispo_grp  = _gruppen(data, 'dispo')
    fc_bul, lbl_bul = _bul_col(bul)

    doc = Document()
    # Seitenränder minimal
    for sec in doc.sections:
        sec.page_width  = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = sec.bottom_margin = Cm(0.8)
        sec.left_margin = sec.right_margin = Cm(0.6)

    # ── Haupttabelle: 2 Spalten (links farbig, rechts weiß) ─────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    LEFT_W  = Cm(7.8)
    RIGHT_W = Cm(12.2)

    left_cell  = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)
    left_cell.width  = LEFT_W
    right_cell.width = RIGHT_W
    _no_border(left_cell); _no_border(right_cell)
    _set_cell_bg(left_cell, bg)
    _set_cell_bg(right_cell, "FFFFFF")
    left_cell.vertical_alignment  = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ── LINKE ZELLE ──────────────────────────────────────────────────────────
    # Logo + Org
    try:
        if LOGO.exists():
            logo_p = left_cell.add_paragraph()
            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_p.paragraph_format.space_before = Pt(4)
            logo_run = logo_p.add_run()
            logo_run.add_picture(str(LOGO), width=Cm(3.0))
    except: pass

    _par(left_cell, "Deutsches Rotes Kreuz", bold=True, size=10, color=titel, align="center", space_before=2)
    _par(left_cell, "Kreisverband Köln e.V.", size=8, color=hell, align="center")
    _par(left_cell, STATION, size=7.5, color="AAAAAA", align="center", space_after=4)

    # Trennlinie als leerer Absatz mit Unterrandlinie
    sep = left_cell.add_paragraph()
    sep.paragraph_format.space_after = Pt(2)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), akz.lstrip('#').upper())
    pBdr.append(bottom); pPr.append(pBdr)

    _par(left_cell, f"📅  {DATUM}", bold=True, size=10, color=titel, align="center", space_before=2)
    _par(left_cell, f"🕖  {UHRZEIT}", size=9, color=hell, align="center", space_after=4)

    # Kennzahlen-Kacheln
    kacheln = [
        ("✦ Einsätze",  str(einz),  akz, dunkel),
        ("✦ Patienten", str(pat),   hell, dunkel),
        ("✦ PAX",       f"{pax:,}".replace(",","."), hell, dunkel),
        ("✦ Personal",  str(len(betreuer)+len(dispo_lst)), hell, dunkel),
    ]
    for lbl, val, vc, bc in kacheln:
        p = left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(3)
        run1 = p.add_run(f"{lbl}  ")
        run1.font.size = Pt(7.5); run1.font.color.rgb = _rgb(hell)
        run2 = p.add_run(val)
        run2.bold = True; run2.font.size = Pt(16); run2.font.color.rgb = _rgb(vc)

    # Bulmor
    sep2 = left_cell.add_paragraph()
    sep2.paragraph_format.space_before = Pt(4); sep2.paragraph_format.space_after = Pt(1)
    pPr2 = sep2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top2 = OxmlElement('w:top'); top2.set(qn('w:val'), 'single'); top2.set(qn('w:sz'), '4')
    top2.set(qn('w:color'), akz.lstrip('#').upper()); top2.set(qn('w:space'), '1')
    pBdr2.append(top2); pPr2.append(pBdr2)

    _par(left_cell, "BULMOR – FAHRZEUGSTATUS", bold=True, size=8, color=akz, align="center", space_before=2)
    fc_h = "FF4444" if bul <= 2 else ("E07800" if bul == 3 else "10A050")
    _par(left_cell, f"{'● '*bul}{'○ '*(5-bul)}",   size=12, color=fc_h, align="center")
    _par(left_cell, f"{bul}/5 im Einsatz  ·  {lbl_bul}", bold=True, size=8.5, color=fc_h, align="center")

    # Bulmor-Fahrer
    if bulfhr_lst:
        _par(left_cell, "Fahrer:", bold=True, size=7.5, color=hell, align="center", space_before=4)
        for fname in bulfhr_lst:
            _par(left_cell, fname, size=7.5, color="BBBBBB", align="center")

    # Kranke
    if kranke_lst:
        _par(left_cell, "── KRANK ──", bold=True, size=7.5, color="FF4444", align="center", space_before=6)
        for kname in kranke_lst:
            _par(left_cell, kname, size=7.5, color="FF8888", align="center")

    # Footer links
    _par(left_cell, "+49 2203 40-2323", size=7, color="888888", align="center", space_before=8)

    # ── RECHTE ZELLE ─────────────────────────────────────────────────────────
    # Header-Streifen (Absatz mit Hintergrund simulieren – nächste Tabelle)
    hdr_tbl = right_cell.add_table(rows=1, cols=1)
    hdr_tbl.style = 'Table Grid'
    hdr_c = hdr_tbl.cell(0, 0); _no_border(hdr_c); _set_cell_bg(hdr_c, akz2)
    hdr_tbl.cell(0, 0).width = RIGHT_W
    _par(hdr_c, f"ALLE BETREUER  ({len(betreuer)} Personen)", bold=True, size=9, color="FFFFFF")

    # Betreuer-Namen in 2 Spalten
    half = (len(betreuer) + 1) // 2
    names_tbl = right_cell.add_table(rows=half, cols=2)
    names_tbl.style = 'Table Grid'
    COL_W = Cm(6.1)
    for i, name_ in enumerate(betreuer):
        row = i % half; col = i // half
        c = names_tbl.cell(row, col)
        _no_border(c); c.width = COL_W
        _set_cell_bg(c, "F5F5F5" if row % 2 == 0 else "FFFFFF")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        run = p.add_run(name_); run.font.size = Pt(8); run.font.color.rgb = _rgb("222222")

    # Dispo-Tabelle
    dispo_hdr = right_cell.add_table(rows=1, cols=1)
    dispo_hdr.style = 'Table Grid'
    dh_c = dispo_hdr.cell(0, 0); _no_border(dh_c); _set_cell_bg(dh_c, dunkel)
    dispo_hdr.cell(0, 0).width = RIGHT_W
    _par(dh_c, f"DISPOSITION  ({len(dispo_lst)} Personen)", bold=True, size=9, color="FFFFFF", space_before=4)

    for i, (zeit, namen) in enumerate(dispo_grp.items()):
        dgrp_tbl = right_cell.add_table(rows=1, cols=2)
        dgrp_tbl.style = 'Table Grid'
        zc = dgrp_tbl.cell(0, 0); nc = dgrp_tbl.cell(0, 1)
        _no_border(zc); _no_border(nc)
        bg_r = "EEF3FF" if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(zc, bg_r); _set_cell_bg(nc, bg_r)
        zc.width = Cm(2.5); nc.width = Cm(9.7)
        p1 = zc.paragraphs[0]; r1 = p1.add_run(zeit)
        r1.bold = True; r1.font.size = Pt(8); r1.font.color.rgb = _rgb(akz2)
        p2 = nc.paragraphs[0]; r2 = p2.add_run(", ".join(namen))
        r2.font.size = Pt(8); r2.font.color.rgb = _rgb("222222")

    # Abstand
    right_cell.add_paragraph().paragraph_format.space_after = Pt(4)

    # Footer rechts
    foot_tbl = right_cell.add_table(rows=1, cols=1)
    foot_tbl.style = 'Table Grid'
    ft_c = foot_tbl.cell(0, 0); _no_border(ft_c); _set_cell_bg(ft_c, dunkel)
    foot_tbl.cell(0, 0).width = RIGHT_W
    _par(ft_c, "DRK Köln  ·  +49 2203 40-2323  ·  flughafen@drk-koeln.de",
         size=7.5, color="FFFFFF", align="center", space_before=2)

    outpath = os.path.join(ZIEL, f"W_{name}_25032026.docx")
    doc.save(outpath)
    print(f"[OK] W_{name}_25032026.docx")


# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Lade Dienstplan-Daten ...")
    data = lade()
    btr = _alle_betreuer(data); dsp = _alle_dispo(data)
    print(f"  Betreuer: {len(btr)}, Dispo: {len(dsp)}, Krank: {len(_kranke(data))}")
    print(f"\nZielordner: {ZIEL}\n")

    print("── 6 PDF-Varianten (P9-Split-Stil) ─────────────────────────────")
    erstelle_alle_pdfs(data, bul=5, einz=28, pat=5, pax=42500)

    print("\n── 2 Word-Versionen (P9-Split-Stil) ────────────────────────────")
    erstelle_word(data, scheme_idx=0, bul=5, einz=28, pat=5, pax=42500)  # Smaragd
    erstelle_word(data, scheme_idx=2, bul=5, einz=28, pat=5, pax=42500)  # OzeanBlau

    print(f"\n✓ Alle Dateien gespeichert in:\n  {ZIEL}")
