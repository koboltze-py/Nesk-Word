# -*- coding: utf-8 -*-
"""
F2 / F3 – Stärkemeldung Demo-Generator
F2 = alles auf einer DIN-A4-Seite (kompakt)
F3 = zwei Seiten (Dispo + Schichtleiter auf Seite 2)

Änderungen:
  - Schichtleiter TAG und NACHT getrennt
  - Tag/Nacht-Farben auch bei Betreuern
  - Bulmor: einzelne grüne/rote Punkte pro Fahrzeug, Gesamtstatus grün/orange/rot
  - Keine Bulmor-Fahrerliste
  - Telefon direkt unter Station, keine Uhrzeit
"""
import os, sys
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Pfade ─────────────────────────────────────────────────────────────────────
_OD   = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
ZIEL  = os.path.join(_OD, "Desktop", "bei", "Neuer Ordner (8)")
os.makedirs(ZIEL, exist_ok=True)
LOGO  = Path(os.path.dirname(os.path.abspath(__file__))) / "Daten" / "Email" / "Logo.jpg"
EXCEL = (_OD + r"\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26"
         r"\04_Tagesdienstpläne\03_März\25.03.2026.xlsx")
DATUM   = "25.03.2026"
STATION = "Erste-Hilfe-Station · Flughafen Köln/Bonn"
TEL     = "+49 2203 40-2323"
MAIL    = "erste-hilfe-station-flughafen@drk-koeln.de"

# ── Farbschema OzeanBlau ──────────────────────────────────────────────────────
BG  = "1A3460"    # Linkes Panel-Hauptfarbe
DK  = "0F1F3C"    # Dunkler Ton
HE  = "C8DAFF"    # Heller Ton (Texte links)
AZ  = "00C8FF"    # Tagdienst-Akzent
AZ2 = "00A878"    # Nachtdienst-Akzent

LW = Cm(7.4)   # Breite linkes Panel
RW = Cm(12.6)  # Breite rechtes Panel
BW = Cm(1.43)  # Breite pro Bulmor-Zelle (5 × 1.43 ≈ 7.15 cm)
FW = LW + RW   # Volle Breite Seite 2  (7.4 + 12.6 = 20.0 cm)

# ── Personen, die immer ausgeschlossen werden ───────────────────────────────────
AUSSCHLIESSEN = {"lars peters"}


# ══════════════════════════════════════════════════════════════════════════════
# VBA-EINBETTUNG – speichert .docx als .docm mit Bulmor-Makro
# ══════════════════════════════════════════════════════════════════════════════
def _inject_bulmor_vba(docx_path, docm_path):
    """
    Öffnet docx unsichtbar in Word, bettet das Bulmor-Makro ein,
    speichert als .docm (FileFormat=13) und schließt Word.

    Voraussetzung: Word → Datei → Optionen → Trust Center → Einstellungen
      → Makroeinstellungen → 'Zugriff auf das VBA-Projektobjektmodell vertrauen'
    """
    import win32com.client, pythoncom

    VBA = """
Sub BulmorAktualisieren()
    Dim n As Integer, eingabe As String
    eingabe = InputBox("Wie viele Bulmor-Fahrzeuge sind in Dienst? (0 bis 5)", "Bulmor-Status", "5")
    If eingabe = "" Then Exit Sub
    If Not IsNumeric(eingabe) Then MsgBox "Bitte Zahl 0-5 eingeben.", vbExclamation: Exit Sub
    n = CInt(eingabe)
    If n < 0 Or n > 5 Then MsgBox "Zahl muss zwischen 0 und 5 liegen.", vbExclamation: Exit Sub

    ' Tabelle 3 = Bulmor-Kacheln (1 Zeile, 5 Spalten)
    Dim tbl As Table
    Set tbl = ActiveDocument.Tables(3)
    Dim i As Integer
    For i = 1 To 5
        If i <= n Then
            tbl.Cell(1, i).Shading.BackgroundPatternColor = RGB(16, 160, 80)
            If tbl.Cell(1, i).Range.Paragraphs.Count >= 3 Then
                Dim r1 As Range: Set r1 = tbl.Cell(1, i).Range.Paragraphs(3).Range
                r1.MoveEnd wdCharacter, -1: r1.Text = "Dienst"
                r1.Font.Color = RGB(204, 255, 204): r1.Font.Size = 6
            End If
        Else
            tbl.Cell(1, i).Shading.BackgroundPatternColor = RGB(204, 34, 0)
            If tbl.Cell(1, i).Range.Paragraphs.Count >= 3 Then
                Dim r2 As Range: Set r2 = tbl.Cell(1, i).Range.Paragraphs(3).Range
                r2.MoveEnd wdCharacter, -1: r2.Text = "Aus"
                r2.Font.Color = RGB(255, 204, 204): r2.Font.Size = 6
            End If
        End If
    Next i

    ' Tabelle 4 = Gesamtstatus-Balken (1x1)
    Dim st As Table: Set st = ActiveDocument.Tables(4)
    Dim sc As Cell: Set sc = st.Cell(1, 1)
    Dim bgC As Long, fgC As Long, txt As String
    If n <= 2 Then
        bgC = RGB(58, 0, 0): fgC = RGB(255, 51, 51)
        txt = "Gesamt: " & n & "/5  " & Chr(9679) & "  ROT - KRITISCH"
    ElseIf n = 3 Then
        bgC = RGB(58, 34, 0): fgC = RGB(255, 140, 0)
        txt = "Gesamt: " & n & "/5  " & Chr(9679) & "  ORANGE - EINGESCHRAENKT"
    Else
        bgC = RGB(0, 58, 24): fgC = RGB(16, 160, 80)
        txt = "Gesamt: " & n & "/5  " & Chr(9679) & "  GRUEN - VOLLSTAENDIG"
    End If
    sc.Shading.BackgroundPatternColor = bgC
    Dim rng As Range: Set rng = sc.Range
    rng.MoveEnd wdCharacter, -1: rng.Text = txt
    rng.Font.Bold = True: rng.Font.Size = 8
    rng.Font.Color = fgC: rng.ParagraphFormat.Alignment = wdAlignParagraphCenter

    MsgBox "Bulmor-Status aktualisiert: " & n & " von 5 Fahrzeugen in Dienst.", vbInformation
End Sub
"""
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = None
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        vba_mod = doc.VBProject.VBComponents.Add(1)  # 1 = vbext_ct_StdModule
        vba_mod.Name = "BulmorMakro"
        vba_mod.CodeModule.AddFromString(VBA)
        doc.SaveAs2(os.path.abspath(docm_path), FileFormat=13)  # 13 = wdFormatXMLDocumentMacroEnabled
        doc.Close(False)
    except Exception as e:
        print(f"  [VBA] Fehler beim Einbetten: {e}")
        print(f"  [VBA] Hinweis: Word → Optionen → Trust Center → Einstellungen →")
        print(f"          'Zugriff auf VBA-Projektobjektmodell vertrauen' aktivieren")
        if doc:
            try: doc.Close(False)
            except: pass
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

# ── Daten laden ────────────────────────────────────────────────────────────────
def lade():
    try:
        from functions.dienstplan_parser import DienstplanParser
        r = DienstplanParser(EXCEL, alle_anzeigen=True).parse()
        if r.get("success") in (True, "True"): return r
    except Exception as e:
        print(f"  [Parser] {e}")
    return {"betreuer": [], "dispo": [], "kranke": []}

def _ist_tag(p):
    try: return int((p.get('start_zeit') or '00:00')[:2]) < 14
    except: return True

def _zeit(p):
    s = (p.get('start_zeit') or '')[:5]; e = (p.get('end_zeit') or '')[:5]
    return f"{s}–{e}" if s and e else ''

def _schicht(d, typ):
    """Gibt (DT-Liste, DN-Liste) zurück: [(name, zeit), ...], sortiert nach Startzeit"""
    dt, dn = [], []
    for p in d.get(typ, []):
        if p.get('ist_krank') in (True, 'True'): continue
        n = p.get('anzeigename', '').strip()
        if not n or n.lower() in AUSSCHLIESSEN: continue
        start = (p.get('start_zeit') or '00:00')[:5]
        z = _zeit(p)
        (dt if _ist_tag(p) else dn).append((start, n, z))
    dt.sort()  # frühster Beginn zuerst, bei Gleichstand alphabetisch
    dn.sort()
    return [(n, z) for _, n, z in dt], [(n, z) for _, n, z in dn]

def _kranke_liste(d):
    return [(p.get('anzeigename', '').strip(), _zeit(p))
            for p in d.get('betreuer', []) + d.get('dispo', [])
            if p.get('ist_krank') in (True, 'True') and p.get('anzeigename', '').strip()]

def _schichtleiter(d, tag=True):
    """Gibt (name, zeit) für Schichtleiter Tag oder Nacht zurück"""
    marked, fallback = [], []
    for p in d.get('dispo', []):
        if p.get('ist_krank') in (True, 'True'): continue
        if _ist_tag(p) != tag: continue
        n = p.get('anzeigename', '').strip()
        if not n or n.lower() in AUSSCHLIESSEN: continue
        if p.get('ist_schichtleiter') in (True, 'True'):
            marked.append((n, _zeit(p)))
        else:
            fallback.append((n, _zeit(p)))
    lst = marked or fallback
    return lst[0] if lst else ('—', '')


# ── Word-Helfer ────────────────────────────────────────────────────────────────
def _rgb(hx):
    hx = hx.lstrip('#')
    return RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

def _no_b(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = OxmlElement('w:tcBorders')
    for s in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{s}'); b.set(qn('w:val'), 'none'); tb.append(b)
    tcPr.append(tb)

def _no_wrap(cell):
    """Verhindert Zeilenumbruch in einer Tabellenzelle (w:noWrap)."""
    tcPr = cell._tc.get_or_add_tcPr()
    nw = OxmlElement('w:noWrap')
    nw.set(qn('w:val'), '1')
    tcPr.append(nw)

def _bg(cell, hx):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hx.lstrip('#').upper()); tcPr.append(shd)


def _zero_empty_paras(cell):
    """
    Alle Leerabsaetze in der Zelle auf exakt 1pt zwingen.
    Word fügt zwischen jede Tabelle und am Ende der Zelle automatisch
    einen Leerabsatz ein (Normal-Stil: 8-21pt). Das verhindert hier Seitenuberlauf.
    """
    for p in cell.paragraphs:
        if not p.text.strip():
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            pPr = p._p.get_or_add_pPr()
            # style-Link entfernen damit Normal-Stil nicht overridet
            for pStyle in pPr.findall(qn('w:pStyle')):
                pPr.remove(pStyle)
            sp = pPr.find(qn('w:spacing'))
            if sp is None:
                sp = OxmlElement('w:spacing'); pPr.append(sp)
            sp.set(qn('w:before'),    '0')
            sp.set(qn('w:after'),     '0')
            sp.set(qn('w:line'),      '20')   # 1pt exakt
            sp.set(qn('w:lineRule'), 'exact')

def _par(cont, text, bold=False, sz=9, fg="000000", align="left", sb=0, sa=0):
    p = cont.add_paragraph()
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(sz)
    r.font.color.rgb = _rgb(fg); return p

def _shdr(cont, text, hx, fg="FFFFFF", w=None, sz=8.5, sb=1, sa=1, align="left"):
    """Section-Header als 1×1 Tabelle"""
    t = cont.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.cell(0, 0); _no_b(c); _bg(c, hx)
    if w: c.width = w
    p = c.add_paragraph()
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(str(text)); r.bold = True; r.font.size = Pt(sz)
    r.font.color.rgb = _rgb(fg)

def _trennL(cont, hx, oben=False):
    """Dünne Trennlinie oben oder unten an einem Absatz"""
    sep = cont.add_paragraph()
    sep.paragraph_format.space_before = Pt(2); sep.paragraph_format.space_after = Pt(1)
    pPr = sep._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    e = OxmlElement('w:top' if oben else 'w:bottom')
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '6')
    e.set(qn('w:space'), '1'); e.set(qn('w:color'), hx.upper())
    pBdr.append(e); pPr.append(pBdr)

def _pagebreak_doc(doc):
    """Seitenumbruch auf Dokumentebene (zwischen Tabellen)"""
    p = doc.add_paragraph()
    r = p.add_run(); br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page'); r._r.append(br)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _pagebreak_in_cell(cell):
    """Seitenumbruch als letzten Absatz innerhalb einer Tabellenzelle."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(); br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page'); r._r.append(br)


# ══════════════════════════════════════════════════════════════════════════════
# BULMOR-BLOCK  –  einzelne grüne/rote Punkte + Gesamtstatus
# ══════════════════════════════════════════════════════════════════════════════
def _bulmor_block(cont, bul, panel_w):
    _shdr(cont, "BULMOR – FAHRZEUGSTATUS", "143060", w=panel_w)

    # 5 farbige Zellen: grün = in Dienst, rot = außer Dienst
    t5 = cont.add_table(rows=1, cols=5); t5.style = 'Table Grid'
    for i in range(5):
        aktiv = (i + 1) <= bul
        c = t5.cell(0, i); _no_b(c)
        _bg(c, "10A050" if aktiv else "CC2200")
        c.width = BW

        # Großer Punkt
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(5); p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run("●"); r1.font.size = Pt(22); r1.font.color.rgb = _rgb("FFFFFF")

        # Fahrzeug-Label
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(f"B{i+1}"); r2.bold = True; r2.font.size = Pt(7.5)
        r2.font.color.rgb = _rgb("FFFFFF")

        # Status-Symbol darunter
        p3 = c.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(3)
        r3 = p3.add_run("Dienst" if aktiv else "Aus"); r3.font.size = Pt(6)
        r3.font.color.rgb = _rgb("CCFFCC" if aktiv else "FFCCCC")

    # Gesamtstatus-Balken (grün/orange/rot je Anzahl)
    if bul <= 2:   sc, sbg, sl = "FF3333", "3A0000", "ROT – KRITISCH"
    elif bul == 3: sc, sbg, sl = "FF8C00", "3A2200", "ORANGE – EINGESCHRÄNKT"
    else:          sc, sbg, sl = "10A050", "003A18", "GRÜN – VOLLSTÄNDIG"

    st = cont.add_table(rows=1, cols=1); st.style = 'Table Grid'
    stc = st.cell(0, 0); _no_b(stc); _bg(stc, sbg); stc.width = panel_w
    p_s = stc.paragraphs[0]; p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s.paragraph_format.space_before = Pt(4); p_s.paragraph_format.space_after = Pt(4)
    r_s = p_s.add_run(f"Gesamt: {bul}/5  ●  {sl}")
    r_s.bold = True; r_s.font.size = Pt(10); r_s.font.color.rgb = _rgb(sc)


# ══════════════════════════════════════════════════════════════════════════════
# NAMENTABELLEN
# ══════════════════════════════════════════════════════════════════════════════
def _set_row_height(row, height_pt):
    """Setzt exakte Zeilenhöhe via XML (trHeight), in Punkt."""
    from docx.oxml import OxmlElement
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    # alte trHeight entfernen falls vorhanden
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_pt * 20)))  # twips: 1pt = 20 twips
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)


def _calc_betreuer_fill(n_dt, n_dn, avail_pt=540):
    """Berechnet exakte Zeilenhöhe (row_h_pt) und Schriftgröße."""
    n_rows = max(1, (n_dt + 1) // 2) + max(1, (n_dn + 1) // 2)
    row_h    = avail_pt / n_rows
    font_sz  = max(7.0, min(18.0, row_h / 3.0))
    return round(row_h, 1), round(font_sz, 1)


def _2spalten(cont, lst, bga, bgb, row_pad=2, font_sz=8, row_h_pt=None, name_w=None, zeit_w=None):
    """Namensliste 2-spaltig mit Arbeitszeit, farbige Zeilen.
    row_h_pt: wenn gesetzt, wird die Zeilenhöhe per XML exakt erzwungen."""
    if not lst: return
    nw = name_w or Cm(4.0); zw = zeit_w or Cm(2.3)
    half = (len(lst) + 1) // 2
    tbl = cont.add_table(rows=half, cols=4); tbl.style = 'Table Grid'
    for i, (name, zeit) in enumerate(lst):
        row_idx = i % half; col = (i // half) * 2
        nc = tbl.cell(row_idx, col); zc = tbl.cell(row_idx, col + 1)
        _no_b(nc); _no_b(zc)
        bg = bga if row_idx % 2 == 0 else bgb
        _bg(nc, bg); _bg(zc, bg)
        nc.width = nw; zc.width = zw
        _no_wrap(zc)
        nc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        zc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pn = nc.paragraphs[0]
        pn.paragraph_format.space_before = Pt(row_pad); pn.paragraph_format.space_after = Pt(row_pad)
        rn = pn.add_run(name); rn.font.size = Pt(font_sz); rn.font.color.rgb = _rgb("111111")
        pz = zc.paragraphs[0]
        pz.paragraph_format.space_before = Pt(row_pad); pz.paragraph_format.space_after = Pt(row_pad)
        rz = pz.add_run(zeit); rz.bold = True; rz.font.size = Pt(max(7, font_sz - 1))
        rz.font.color.rgb = _rgb("666666")

    if row_h_pt:
        seen = set()
        for i in range(len(lst)):
            row_idx = i % half
            if row_idx not in seen:
                _set_row_height(tbl.rows[row_idx], row_h_pt)
                seen.add(row_idx)

def _1spalte(cont, lst, bga, bgb, fg_zeit, name_w=None, zeit_w=None):
    """Namensliste 1-spaltig – als einheitliche Tabelle (gleicher Stil wie _2spalten)"""
    if not lst: return
    nw = name_w or Cm(9.6); zw = zeit_w or Cm(3.0)
    tbl = cont.add_table(rows=len(lst), cols=2); tbl.style = 'Table Grid'
    for i, (name, zeit) in enumerate(lst):
        nc = tbl.cell(i, 0); zc = tbl.cell(i, 1)
        _no_b(nc); _no_b(zc)
        bg = bga if i % 2 == 0 else bgb
        _bg(nc, bg); _bg(zc, bg)
        nc.width = nw; zc.width = zw
        nc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        zc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _no_wrap(zc)
        pn = nc.paragraphs[0]
        pn.paragraph_format.space_before = Pt(2); pn.paragraph_format.space_after = Pt(2)
        rn = pn.add_run(name); rn.bold = True; rn.font.size = Pt(10)
        rn.font.color.rgb = _rgb("111111")
        pz = zc.paragraphs[0]; pz.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pz.paragraph_format.space_before = Pt(2); pz.paragraph_format.space_after = Pt(2)
        rz = pz.add_run(zeit); rz.bold = True; rz.font.size = Pt(9)
        rz.font.color.rgb = _rgb(fg_zeit)


# ══════════════════════════════════════════════════════════════════════════════
# SCHICHTLEITER-BLOCK  –  DT und DN mit Badges
# ══════════════════════════════════════════════════════════════════════════════
def _add_sl(cont, sl_dt, sl_dn, w=None):
    w = w or RW; zw = Cm(3.0); nw = w - zw
    _shdr(cont, "SCHICHTLEITER", "0F2848", w=w)
    for hx, hx_bg, label, (name, zeit) in [
        (AZ,  "E8F6FF", "Tagdienst",   sl_dt),
        (AZ2, "E5F8F2", "Nachtdienst", sl_dn),
    ]:
        _shdr(cont, label, hx, "FFFFFF", w, sz=7.5, sb=0, sa=0)
        t = cont.add_table(rows=1, cols=2); t.style = 'Table Grid'
        nc = t.cell(0, 0); zc = t.cell(0, 1)
        _no_b(nc); _no_b(zc)
        _bg(nc, hx_bg); _bg(zc, hx_bg)
        nc.width = nw; zc.width = zw
        nc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        zc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _no_wrap(zc)
        pn = nc.paragraphs[0]
        pn.paragraph_format.space_before = Pt(4); pn.paragraph_format.space_after = Pt(4)
        rn = pn.add_run(name if name else "—")
        rn.bold = True; rn.font.size = Pt(13); rn.font.color.rgb = _rgb("0A2040")
        pz = zc.paragraphs[0]; pz.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pz.paragraph_format.space_before = Pt(6); pz.paragraph_format.space_after = Pt(4)
        rz = pz.add_run(zeit if zeit else "")
        rz.bold = True; rz.font.size = Pt(10); rz.font.color.rgb = _rgb(hx)


# ══════════════════════════════════════════════════════════════════════════════
# DISPO-BLOCK
# ══════════════════════════════════════════════════════════════════════════════
def _add_dispo(cont, di_dt, di_dn, w=None):
    w = w or RW; zw = Cm(3.0); nw = w - zw
    _shdr(cont, "DISPOSITION", DK, w=w)
    if di_dt:
        _shdr(cont, "Tagdienst", AZ, "FFFFFF", w, sz=7.5, sb=0, sa=0)
        _1spalte(cont, di_dt, bga="E8F6FF", bgb="F4FAFF", fg_zeit=AZ, name_w=nw, zeit_w=zw)
    if di_dn:
        _shdr(cont, "Nachtdienst", AZ2, "FFFFFF", w, sz=7.5, sb=0, sa=0)
        _1spalte(cont, di_dn, bga="E5F7F2", bgb="F2FBF8", fg_zeit=AZ2, name_w=nw, zeit_w=zw)


# ══════════════════════════════════════════════════════════════════════════════
# LINKES PANEL aufbauen
# ══════════════════════════════════════════════════════════════════════════════
def _build_left(lc, bul, einz, pat, pax, mini=False):
    """
    mini=True  → kompakte Version für Seite 2 (F3)
    mini=False → volle Version für Seite 1
    """
    # Logo
    try:
        if LOGO.exists():
            lp = lc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp.paragraph_format.space_before = Pt(3)
            lp.add_run().add_picture(str(LOGO), width=Cm(2.6 if mini else 2.8))
    except: pass

    _par(lc, "Deutsches Rotes Kreuz", bold=True, sz=9 if mini else 12, fg="FFFFFF", align="center", sb=2)
    _par(lc, "Kreisverband Köln e.V.", sz=8 if mini else 9, fg=HE, align="center")
    _par(lc, STATION, sz=7 if mini else 8, fg="CCDDFF", align="center")
    _par(lc, TEL,     bold=True, sz=9 if mini else 10, fg="FFFFFF", align="center", sb=2, sa=0)
    _par(lc, MAIL,    sz=6.5 if mini else 7.5, fg="AABBDD", align="center", sa=3)

    # Trennlinie
    sep = lc.add_paragraph(); sep.paragraph_format.space_after = Pt(1)
    pPr = sep._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), AZ.upper())
    pBdr.append(bot); pPr.append(pBdr)

    _par(lc, f"Datum:  {DATUM}", bold=True, sz=9, fg="FFFFFF", align="center", sb=3, sa=4)

    if not mini:
        # Kennzahlen
        for lbl, val, vc in [
            ("✦ Rettungsdiensteinsätze",  str(einz),                          AZ),
            ("✦ Patienten auf Station",   str(pat),                           "AAFFCC"),
            ("✦ PAX",                     f"{pax:,}".replace(",", "."),       HE),
        ]:
            pl = lc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pl.paragraph_format.space_before = Pt(6); pl.paragraph_format.space_after = Pt(1)
            rl = pl.add_run(lbl); rl.font.size = Pt(10); rl.font.color.rgb = _rgb(HE)
            p = lc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
            r2 = p.add_run(val); r2.bold = True; r2.font.size = Pt(18)
            r2.font.color.rgb = _rgb(vc)
    else:
        # Mini: Kennzahlen kompakter
        for lbl, val, vc in [("Einsätze", str(einz), AZ)]:
            p = lc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(f"{lbl}: "); r1.font.size = Pt(8); r1.font.color.rgb = _rgb(HE)
            r2 = p.add_run(val); r2.bold = True; r2.font.size = Pt(14)
            r2.font.color.rgb = _rgb(vc)

    # Bulmor-Block
    _bulmor_block(lc, bul, LW)


# ══════════════════════════════════════════════════════════════════════════════
# RECHTES PANEL – Header
# ══════════════════════════════════════════════════════════════════════════════
def _build_right_header(rc, seite2=False):
    # 3-zeiliger farbiger Org-Header
    ht = rc.add_table(rows=3, cols=1); ht.style = 'Table Grid'
    r0 = ht.cell(0, 0); _no_b(r0); _bg(r0, DK); r0.width = RW
    _par(r0, STATION, bold=True, sz=9, fg=HE, sb=2)
    r1 = ht.cell(1, 0); _no_b(r1); _bg(r1, DK); r1.width = RW
    _par(r1, DATUM, sz=8, fg="AAAAAA")
    r2 = ht.cell(2, 0); _no_b(r2); _bg(r2, "143060"); r2.width = RW
    suffix = "  ·  Seite 2" if seite2 else ""
    _par(r2, f"{DATUM}{suffix}", bold=True, sz=9, fg=AZ, sa=2)

    # Tag/Nacht Legende
    lg = rc.add_table(rows=1, cols=2); lg.style = 'Table Grid'
    for j, (hx, lbl) in enumerate([
        (AZ,  "■ TAGDIENST"),
        (AZ2, "■ NACHTDIENST"),
    ]):
        c = lg.cell(0, j); _no_b(c); _bg(c, hx); c.width = Cm(6.3)
        p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(lbl); r.bold = True; r.font.size = Pt(7); r.font.color.rgb = _rgb("FFFFFF")

    rc.add_paragraph().paragraph_format.space_after = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
# RECHTES PANEL – Betreuer mit Tag/Nacht-Farben
# ══════════════════════════════════════════════════════════════════════════════
def _build_right_betreuer(rc, bt_dt, bt_dn, fill=False):
    if fill:
        row_h, fsz = _calc_betreuer_fill(len(bt_dt), len(bt_dn))
        pad = max(1.0, (row_h - fsz * 1.3) / 2)
    else:
        row_h, fsz, pad = None, 8, 1
    if bt_dt:
        _shdr(rc, "BETREUER \u2013 TAGDIENST", DK, w=RW)
        _2spalten(rc, bt_dt, bga="E8F6FF", bgb="F4FAFF", row_pad=pad, font_sz=fsz, row_h_pt=row_h)

    if bt_dn:
        _shdr(rc, "BETREUER \u2013 NACHTDIENST", "1E2F4A", w=RW)
        _2spalten(rc, bt_dn, bga="E5F7F2", bgb="F2FBF8", row_pad=pad, font_sz=fsz, row_h_pt=row_h)


# ══════════════════════════════════════════════════════════════════════════════
# RECHTES PANEL – Dispo + SL + Footer
# ══════════════════════════════════════════════════════════════════════════════
def _build_right_dispo_sl_footer(rc, di_dt, di_dn, sl_dt, sl_dn):
    _add_dispo(rc, di_dt, di_dn)
    _add_sl(rc, sl_dt, sl_dn)

    # Footer
    rc.add_paragraph().paragraph_format.space_before = Pt(4)
    _shdr(rc, f"DRK Köln  ·  {DATUM}", DK, "FFFFFF", RW, sz=7.5, sb=3, sa=3, align="center")


# ══════════════════════════════════════════════════════════════════════════════
# SEITE 2 – VARIANTE A: Nur Schichtleiter (volle Breite)
# ══════════════════════════════════════════════════════════════════════════════
def _build_page2_sl_only(doc, sl_dt, sl_dn):
    """Seite 2: Nur Schichtleiter-Block, volle Seitenbreite, kein Sidebar"""
    ht = doc.add_table(rows=2, cols=1); ht.style = 'Table Grid'
    r0 = ht.cell(0, 0); _no_b(r0); _bg(r0, DK); r0.width = FW
    _par(r0, f"SCHICHTLEITER  ·  {DATUM}  ·  Seite 2", bold=True, sz=10, fg=HE, sb=3, sa=3)
    r1 = ht.cell(1, 0); _no_b(r1); _bg(r1, "143060"); r1.width = FW
    _par(r1, f"{STATION}  ·  {TEL}", sz=8, fg=AZ, sb=2, sa=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _add_sl(doc, sl_dt, sl_dn, w=FW)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)
    _shdr(doc, f"DRK Köln  ·  {DATUM}", DK, "FFFFFF", FW, sz=7.5, sb=3, sa=3, align="center")


# ══════════════════════════════════════════════════════════════════════════════
# SEITE 2 – VARIANTE B: Dispo + Schichtleiter (volle Breite)
# ══════════════════════════════════════════════════════════════════════════════
def _build_page2(doc, di_dt, di_dn, sl_dt, sl_dn):
    # Kopfzeile
    ht = doc.add_table(rows=2, cols=1); ht.style = 'Table Grid'
    r0 = ht.cell(0, 0); _no_b(r0); _bg(r0, DK); r0.width = FW
    _par(r0, f"DISPOSITION & SCHICHTLEITER  ·  {DATUM}", bold=True, sz=10, fg=HE, sb=3, sa=3)
    r1 = ht.cell(1, 0); _no_b(r1); _bg(r1, "143060"); r1.width = FW
    _par(r1, f"Seite 2  ·  {DATUM}", sz=8, fg=AZ, sb=2, sa=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    # Dispo + SL vollbreite
    _add_dispo(doc, di_dt, di_dn, w=FW)
    _add_sl(doc, sl_dt, sl_dn, w=FW)
    # Footer
    doc.add_paragraph().paragraph_format.space_before = Pt(4)
    _shdr(doc, f"DRK Köln  ·  {DATUM}", DK, "FFFFFF", FW, sz=7.5, sb=3, sa=3, align="center")


# ══════════════════════════════════════════════════════════════════════════════
# SEITE 2 – VARIANTE C: Betreuer (volle Breite, füllt Seite)
# ══════════════════════════════════════════════════════════════════════════════
def _build_page2_betreuer(doc, bt_dt, bt_dn):
    ht = doc.add_table(rows=2, cols=1); ht.style = 'Table Grid'
    r0 = ht.cell(0, 0); _no_b(r0); _bg(r0, DK); r0.width = FW
    _par(r0, f"BETREUER  ·  {DATUM}  ·  Seite 2", bold=True, sz=10, fg=HE, sb=3, sa=3)
    r1 = ht.cell(1, 0); _no_b(r1); _bg(r1, "143060"); r1.width = FW
    _par(r1, STATION, sz=8, fg=AZ, sb=2, sa=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    nw = Cm(7.4); zw = Cm(2.6)
    row_h, fsz = _calc_betreuer_fill(len(bt_dt), len(bt_dn), avail_pt=600)
    pad = max(1.0, (row_h - fsz * 1.3) / 2)
    if bt_dt:
        _shdr(doc, "BETREUER \u2013 TAGDIENST", DK, w=FW)
        _2spalten(doc, bt_dt, bga="E8F6FF", bgb="F4FAFF", row_pad=pad, font_sz=fsz, row_h_pt=row_h,
                  name_w=nw, zeit_w=zw)
    if bt_dn:
        _shdr(doc, "BETREUER \u2013 NACHTDIENST", "1E2F4A", w=FW)
        _2spalten(doc, bt_dn, bga="E5F7F2", bgb="F2FBF8", row_pad=pad, font_sz=fsz, row_h_pt=row_h,
                  name_w=nw, zeit_w=zw)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)
    _shdr(doc, f"DRK Köln  ·  {DATUM}", DK, "FFFFFF", FW, sz=7.5, sb=3, sa=3, align="center")


# ══════════════════════════════════════════════════════════════════════════════
# NEUE HAUPT-TABELLE erzeugen (Links-Rechts-Panel)
# ══════════════════════════════════════════════════════════════════════════════
def _new_main_table(doc):
    main = doc.add_table(rows=1, cols=2); main.style = 'Table Grid'
    lc = main.cell(0, 0); rc = main.cell(0, 1)
    lc.width = LW; rc.width = RW
    _no_b(lc); _no_b(rc)
    _bg(lc, BG); _bg(rc, "FFFFFF")
    lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # Mindesthoehe: A4 - 2x Rand = 29.7cm - 2x0.7cm = 28.3cm = 802pt
    tr = main.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(28.3 / 2.54 * 72 * 20)))  # cm -> pt -> twips
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)
    return lc, rc


# ══════════════════════════════════════════════════════════════════════════════
# DOKUMENT ERSTELLEN – Hauptfunktion
# ══════════════════════════════════════════════════════════════════════════════
def erstelle(filename, data, bul=5, einz=28, pat=5, pax=42500,
             layout="bt_auf_s2", max_bt_dt=None, max_bt_dn=None, fill_betreuer=False,
             sl_demo_dt=None, sl_demo_dn=None):
    """
    layout:
      "bt_auf_s2"       → Dispo+SL Seite 1, Betreuer Seite 2  (Standard)
      "1seite"          → Dispo+SL oben, Betreuer unten – alles Seite 1
      "dispo_sl_auf_s2" → Betreuer Seite 1, Dispo+SL Seite 2 (alt)
      "sl_auf_s2"       → Dispo+Betreuer Seite 1, Schichtleiter Seite 2 (alt)
    """
    bt_dt, bt_dn = _schicht(data, 'betreuer')
    # Auf max_bt_dt/dn kürzen ODER bis zur gewünschten Anzahl mit Demo-Namen auffüllen
    def _pad(lst, ziel, schicht):
        if ziel is None: return lst
        result = list(lst)
        i = 1
        while len(result) < ziel:
            result.append((f"Demo-Person {i:02d}", "07:00–19:00" if schicht == "DT" else "19:00–07:00"))
            i += 1
        return result[:ziel]
    bt_dt = _pad(bt_dt, max_bt_dt, "DT")
    bt_dn = _pad(bt_dn, max_bt_dn, "DN")
    di_dt, di_dn = _schicht(data, 'dispo')
    sl_dt = sl_demo_dt if sl_demo_dt is not None else _schichtleiter(data, tag=True)
    sl_dn = sl_demo_dn if sl_demo_dn is not None else _schichtleiter(data, tag=False)

    doc = Document()
    for sec in doc.sections:
        sec.page_width  = Cm(21.0); sec.page_height = Cm(29.7)
        sec.top_margin  = sec.bottom_margin = Cm(0.7)
        sec.left_margin = sec.right_margin  = Cm(0.5)

    # ── SEITE 1 ───────────────────────────────────────────────────────────────
    lc, rc = _new_main_table(doc)
    _build_left(lc, bul, einz, pat, pax, mini=False)
    _build_right_header(rc, seite2=False)

    if layout == "bt_auf_s2":
        # Dispo+SL auf Seite 1, Betreuer auf Seite 2
        _add_dispo(rc, di_dt, di_dn)
        _add_sl(rc, sl_dt, sl_dn)
        _shdr(rc, "▶  Betreuer  →  Seite 2", "1A3A5A", AZ, RW, sz=8.5, sb=3, sa=3, align="center")
        _pagebreak_in_cell(rc)
        _build_page2_betreuer(doc, bt_dt, bt_dn)

    elif layout == "1seite":
        # Dispo+SL oben, dann Betreuer, alles auf Seite 1
        _add_dispo(rc, di_dt, di_dn)
        _add_sl(rc, sl_dt, sl_dn)
        _build_right_betreuer(rc, bt_dt, bt_dn, fill=fill_betreuer)
        rc.add_paragraph().paragraph_format.space_before = Pt(4)
        _shdr(rc, f"DRK Köln  ·  {DATUM}", DK, "FFFFFF", RW, sz=7.5, sb=3, sa=3, align="center")
        if fill_betreuer:
            _zero_empty_paras(rc)

    elif layout == "sl_auf_s2":
        # Dispo+Betreuer Seite 1, Schichtleiter Seite 2
        _add_dispo(rc, di_dt, di_dn)
        _build_right_betreuer(rc, bt_dt, bt_dn, fill=fill_betreuer)
        _shdr(rc, "▶  Schichtleiter  →  Seite 2", "1A3A5A", AZ, RW, sz=8.5, sb=3, sa=3, align="center")
        if fill_betreuer:
            _zero_empty_paras(rc)
        _pagebreak_in_cell(rc)
        _build_page2_sl_only(doc, sl_dt, sl_dn)

    elif layout == "dispo_sl_auf_s2":
        # Betreuer Seite 1, Dispo+SL Seite 2
        _build_right_betreuer(rc, bt_dt, bt_dn, fill=fill_betreuer)
        _shdr(rc, "▶  Disposition & Schichtleiter  →  Seite 2", "1A3A5A", AZ, RW, sz=8.5, sb=3, sa=3, align="center")
        if fill_betreuer:
            _zero_empty_paras(rc)
        _pagebreak_in_cell(rc)
        _build_page2(doc, di_dt, di_dn, sl_dt, sl_dn)

    out = os.path.join(ZIEL, filename)
    if filename.endswith('.docm'):
        import shutil
        tmp = out[:-1] + 'x'  # .docm → temp .docx
        doc.save(tmp)
        _inject_bulmor_vba(tmp, out)
        if not os.path.exists(out):  # Fallback: VBA-Einbettung gescheitert → .docx als .docm kopieren
            shutil.copy2(tmp, out)
            print(f"  [WARN] Ohne VBA gespeichert – Trust Center aktivieren!")
        try: os.remove(tmp)
        except: pass
    else:
        doc.save(out)
    print(f"  [OK] {filename}")


# ══════════════════════════════════════════════════════════════════════════════
# DEMO-KONFIGURATION  –  alle Betreuer-Größen × alle Bulmor-Zustände (0–5)
# ══════════════════════════════════════════════════════════════════════════════
#  dateiname                              bul  max_dt  max_dn
DEMOS = [
    # ── Klein (5 DT + 4 DN) ───────────────────────────────────────────────
    ("Klein_Bul0_Rot.docm",                 0,    5,    4),
    ("Klein_Bul1_Rot.docm",                 1,    5,    4),
    ("Klein_Bul2_Rot.docm",                 2,    5,    4),
    ("Klein_Bul3_Orange.docm",              3,    5,    4),
    ("Klein_Bul4_Gruen.docm",               4,    5,    4),
    ("Klein_Bul5_Gruen.docm",               5,    5,    4),
    # ── Mittel (10 DT + 8 DN) ─────────────────────────────────────────────
    ("Mittel_Bul0_Rot.docm",                0,   10,    8),
    ("Mittel_Bul1_Rot.docm",                1,   10,    8),
    ("Mittel_Bul2_Rot.docm",                2,   10,    8),
    ("Mittel_Bul3_Orange.docm",             3,   10,    8),
    ("Mittel_Bul4_Gruen.docm",              4,   10,    8),
    ("Mittel_Bul5_Gruen.docm",              5,   10,    8),
    # ── Voll (reale Daten) ────────────────────────────────────────────────
    ("Voll_Bul0_Rot.docm",                  0, None, None),
    ("Voll_Bul1_Rot.docm",                  1, None, None),
    ("Voll_Bul2_Rot.docm",                  2, None, None),
    ("Voll_Bul3_Orange.docm",               3, None, None),
    ("Voll_Bul4_Gruen.docm",                4, None, None),
    ("Voll_Bul5_Gruen.docm",                5, None, None),
    # ── Gross (25 DT + 20 DN) ─────────────────────────────────────────────
    ("Gross_Bul0_Rot.docm",                 0,   25,   20),
    ("Gross_Bul1_Rot.docm",                 1,   25,   20),
    ("Gross_Bul2_Rot.docm",                 2,   25,   20),
    ("Gross_Bul3_Orange.docm",              3,   25,   20),
    ("Gross_Bul4_Gruen.docm",               4,   25,   20),
    ("Gross_Bul5_Gruen.docm",               5,   25,   20),
    # ── Maximum (30 DT + 30 DN) ───────────────────────────────────────────
    ("Max30_Bul0_Rot.docm",                 0,   30,   30),
    ("Max30_Bul1_Rot.docm",                 1,   30,   30),
    ("Max30_Bul2_Rot.docm",                 2,   30,   30),
    ("Max30_Bul3_Orange.docm",              3,   30,   30),
    ("Max30_Bul4_Gruen.docm",               4,   30,   30),
    ("Max30_Bul5_Gruen.docm",               5,   30,   30),
]


# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Lade Dienstplan-Daten ...")
    data = lade()
    bt_dt, bt_dn = _schicht(data, 'betreuer')
    di_dt, di_dn = _schicht(data, 'dispo')
    sl_dt = _schichtleiter(data, tag=True)
    sl_dn = _schichtleiter(data, tag=False)
    print(f"  Betreuer  DT:{len(bt_dt)}  DN:{len(bt_dn)}")
    print(f"  Dispo     DT:{len(di_dt)}  DN:{len(di_dn)}")
    print(f"  SL Tag:   {sl_dt[0]}  {sl_dt[1]}")
    print(f"  SL Nacht: {sl_dn[0]}  {sl_dn[1]}")
    print(f"\nZielordner: {ZIEL}\n")

    print(f"Erstelle {len(DEMOS)} Demo-Dateien ...")
    for fname, bul, max_dt, max_dn in DEMOS:
        erstelle(fname, data, bul=bul, einz=28, pat=5, pax=42500,
                 layout="bt_auf_s2",
                 max_bt_dt=max_dt, max_bt_dn=max_dn,
                 fill_betreuer=True)

    print(f"\n✓ {len(DEMOS)} Dateien gespeichert in:\n  {ZIEL}")
