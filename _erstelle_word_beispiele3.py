# -*- coding: utf-8 -*-
"""
Word-Beispiele Batch 3 – 6 neue moderne Designs
Recherche-Basis: Coolors Trending, Visme Award-Winning Schemes, Material Design
"""
import os, sys
from pathlib import Path
from collections import defaultdict

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Pfade ─────────────────────────────────────────────────────────────────────
_OD = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
ZIELORDNER = os.path.join(_OD, "Desktop", "bei") if os.path.exists(_OD) else r"C:\Temp\bei"
os.makedirs(ZIELORDNER, exist_ok=True)
LOGO_PFAD = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Daten", "Email", "Logo.jpg")
EXCEL_PFAD = (
    _OD + r"\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26"
    r"\04_Tagesdienstpläne\03_März\25.03.2026.xlsx"
)

# ── Echtdaten laden ───────────────────────────────────────────────────────────
def lade_dienstplan():
    try:
        from functions.dienstplan_parser import DienstplanParser
        r = DienstplanParser(EXCEL_PFAD, alle_anzeigen=True).parse()
        if r.get("success") in (True, "True"): return r
    except Exception as e:
        print(f"  [Parser] {e}")
    return {"betreuer": [], "dispo": [], "kranke": []}

# ── Verbrauchsmaterial ────────────────────────────────────────────────────────
MAT = [
    ("Einmalhandschuhe Nitril M",   "Karton", "8",  "3"),
    ("Einmalhandschuhe Nitril L",   "Karton", "8",  "3"),
    ("Verbandpäckchen groß",        "Stück",  "20", "8"),
    ("Mullbinden 6/8/10 cm",        "Rollen", "60", "20"),
    ("Wundkompressen 10×10 steril", "Päck.", "40", "15"),
    ("Pflaster-Sortiment",          "Päck.", "10",  "4"),
    ("Hände-Desinfektion 1L",       "Fl.",   "6",   "2"),
    ("Flächen-Desinfektion 1L",     "Fl.",   "4",   "2"),
    ("Rettungsdecken Gold/Silber",  "Stück", "20",  "8"),
    ("Einmal-Beatmungsmaske",       "Stück", "10",  "4"),
    ("Venenverweilkanüle G18/G20",  "Stück", "30", "10"),
    ("Einmalspritze 5/10/20 ml",    "Stück", "50", "20"),
    ("Infusionsset + NaCl 500ml",   "Sets",  "10",  "4"),
    ("EKG-Elektroden Einmal",       "Päck.", "8",   "3"),
    ("Sauerstoffmaske Einmal",      "Stück", "15",  "5"),
    ("Atemschutzmaske FFP2",        "Stück", "50", "20"),
    ("AED-Elektroden Einmal",       "Paar",  "4",   "2"),
    ("Blutzucker-Teststreifen",     "Päck.", "5",   "2"),
    ("Urinbeutel steril",           "Stück", "12",  "4"),
    ("CELOX Hämostase-Verband",     "Stück", "4",   "2"),
]

# ── Basis-Tools ───────────────────────────────────────────────────────────────
def c(r,g,b): return RGBColor(r,g,b)
def _s(rgb): return str(rgb)
WEISS = c(0xFF,0xFF,0xFF); SCHWARZ = c(0x1A,0x1A,0x1A)
ROT_A = c(0xBE,0x00,0x00); GRN_A = c(0x10,0x7E,0x3E); ORG_A = c(0xE6,0x7E,0x00)
ROT_B = c(0xFF,0x33,0x33); GRN_B = c(0x2A,0xC0,0x7E); ORG_B = c(0xFF,0x99,0x00)

def _shd(cell, color):
    tc=cell._tc; p=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),_s(color))
    p.append(s)

def _bdr(para, col_hex="888888", sz="6"):
    pPr=para._p.get_or_add_pPr(); b=OxmlElement('w:pBdr')
    bot=OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),sz)
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),col_hex)
    b.append(bot); pPr.append(b)

def _run(para, text, bold=False, size=10, color=None, italic=False):
    r=para.add_run(text); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic
    if color: r.font.color.rgb=color
    return r

def _cell(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p=cell.paragraphs[0]; p.alignment=align
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold
    if color: r.font.color.rgb=color

def _cell_add(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p=cell.add_paragraph(); p.alignment=align
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold
    if color: r.font.color.rgb=color

def _margins(doc, l=1.8, r=1.8, t=2.8, b=2.0):
    for s in doc.sections:
        s.left_margin=Cm(l); s.right_margin=Cm(r)
        s.top_margin=Cm(t); s.bottom_margin=Cm(b)

def _logo_header(doc, acc_color=None, sub_color=None):
    section=doc.sections[0]; header=section.header
    ht=header.add_table(rows=1,cols=2,width=Inches(6.5)); ht.autofit=False
    logo_path=Path(LOGO_PFAD)
    if logo_path.exists():
        lp=ht.rows[0].cells[0].paragraphs[0]; lp.add_run().add_picture(str(logo_path),width=Inches(1.0))
    rp=ht.rows[0].cells[1].paragraphs[0]; rp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r1=rp.add_run("Deutsches Rotes Kreuz Kreisverband Köln e.V.\n")
    r1.font.size=Pt(9); r1.font.bold=True
    r1.font.color.rgb = acc_color if acc_color else ROT_A
    r2=rp.add_run("Unfallhilfsstelle · EH-Station Flughafen Köln/Bonn")
    r2.font.size=Pt(8)
    r2.font.color.rgb = sub_color if sub_color else c(0x55,0x55,0x55)
    tp=header.add_paragraph(); _bdr(tp,"BE0000","8")

def _footer(doc):
    f=doc.sections[0].footer; fp=f.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run("☎ +49 2203 40-2323   |   ✉ flughafen@drk-koeln.de   |   Stationsleitung: Lars Peters")
    r.font.size=Pt(8); r.font.color.rgb=c(0x77,0x77,0x77)

def _bulmor_col(n, grn, org, rot):
    if n<=2: return rot, "KRITISCH"
    elif n==3: return org, "EINGESCHRÄNKT"
    else: return grn, "VOLLSTÄNDIG"

def _personal(data):
    return len([p for p in data.get('betreuer',[])+data.get('dispo',[])
                if p.get('ist_krank') not in (True,'True')])

def _kranke(data):
    return [p.get('anzeigename','') for p in data.get('kranke',[])
            if p.get('ist_krank') in (True,'True')]

def _bulmor_fahrer(data):
    alle=data.get('betreuer',[])+data.get('dispo',[])
    return [p.get('anzeigename','') for p in alle if p.get('ist_bulmorfahrer') in (True,'True')]

def _gruppen_dispo(data):
    gru=defaultdict(list)
    for p in data.get('dispo',[]):
        if p.get('ist_krank') in (True,'True'): continue
        s=(p.get('start_zeit') or '')[:5]; e=(p.get('end_zeit') or '')[:5]
        s2=f"{int(s.split(':')[0]):02d}:00" if s and ':' in s else s
        e2=f"{int(e.split(':')[0]):02d}:00" if e and ':' in e else e
        gru[f"{s2} – {e2}"].append(p.get('anzeigename',''))
    return dict(sorted(gru.items()))

def _gruppen_bet(data):
    gru=defaultdict(list)
    for p in data.get('betreuer',[]):
        if p.get('ist_krank') in (True,'True'): continue
        s=(p.get('start_zeit') or '')[:5]; e=(p.get('end_zeit') or '')[:5]
        gru[f"{s} – {e}"].append(p.get('anzeigename',''))
    return dict(sorted(gru.items()))

def _alle_grouped(data):
    gru=defaultdict(list)
    for p in data.get('dispo',[])+data.get('betreuer',[]):
        if p.get('ist_krank') in (True,'True'): continue
        ist_d=p.get('ist_dispo') in (True,'True')
        s=(p.get('start_zeit') or '')[:5]; e=(p.get('end_zeit') or '')[:5]
        prefix="[D] " if ist_d else ""
        gru[f"{prefix}{s} – {e}"].append(p.get('anzeigename',''))
    return dict(sorted(gru.items()))

def _alle_personen_sortiert(data):
    alle=[]
    for p in data.get('dispo',[]):
        if p.get('ist_krank') not in (True,'True'): alle.append(('DISPO',p))
    for p in data.get('betreuer',[]):
        if p.get('ist_krank') not in (True,'True'): alle.append(('BET.',p))
    alle.sort(key=lambda x:(x[1].get('start_zeit') or 'ZZZZ'))
    return alle


# ═══════════════════════════════════════════════════════════════════════════════
# D7 – FIERY OCEAN  (Coolors Trending: 780000/C1121F/FDF0D5/003049/669BBC)
# Dunkelrot · Cremeweiß · Navyblau · Stahlblau
# ═══════════════════════════════════════════════════════════════════════════════
def design7_fiery_ocean(data, bul=4, einz=28, pat=5, pax=42500):
    NAVY    = c(0x00,0x30,0x49)
    STAHL   = c(0x66,0x9B,0xBC)
    DUNKELR = c(0x78,0x00,0x00)
    HELLR   = c(0xC1,0x12,0x1F)
    CREME   = c(0xFD,0xF0,0xD5)
    GRN     = c(0x2B,0x9E,0x6A)
    ORG     = c(0xE6,0x80,0x00)
    ROT     = c(0xC1,0x12,0x1F)

    doc=Document(); _margins(doc); _logo_header(doc, HELLR, NAVY); _footer(doc)

    # Titel
    tbl=doc.add_table(rows=1,cols=1)
    tc=tbl.rows[0].cells[0]; _shd(tc,DUNKELR)
    tp=tc.paragraphs[0]; tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(tp,"Stärkemeldung und Einsätze",bold=True,size=22,color=CREME)
    sp=tc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(sp,"DRK Köln  ·  Flughafen Köln/Bonn  ·  25.03.2026  ·  07:45 Uhr",size=9,color=STAHL)
    doc.add_paragraph()

    # Scorecard
    sc=doc.add_table(rows=2,cols=4); sc.style='Table Grid'
    farbe,label=_bulmor_col(bul,GRN,ORG,ROT)
    kz=[("EINSÄTZE",str(einz),HELLR),("PATIENTEN",str(pat),NAVY),
        ("PAX",f"{pax:,}".replace(",","."),STAHL),("PERSONAL",str(_personal(data)),NAVY)]
    for i,(lbl,val,f) in enumerate(kz):
        h=sc.rows[0].cells[i]; v=sc.rows[1].cells[i]
        _shd(h,NAVY); _shd(v,CREME)
        _cell(h,lbl,bold=True,size=8,color=CREME,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(v,val,bold=True,size=26,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Bulmor – Ozean-Wellen-Stil
    bh=doc.add_paragraph()
    _run(bh,"BULMOR – FAHRZEUGSTATUS",bold=True,size=12,color=DUNKELR)
    _bdr(bh,"780000","8")
    bul_tbl=doc.add_table(rows=2,cols=5); bul_tbl.style='Table Grid'
    fc,_=_bulmor_col(bul,GRN,ORG,ROT)
    for i in range(5):
        aktiv=(i+1)<=bul
        c1=bul_tbl.rows[0].cells[i]; c2=bul_tbl.rows[1].cells[i]
        _shd(c1,fc if aktiv else c(0xCC,0xCC,0xCC))
        _shd(c2,CREME)
        sym="◉" if aktiv else "○"
        _cell(c1,f"{sym}  B{i+1}",bold=True,size=12,
              color=WEISS if aktiv else c(0x99,0x99,0x99),align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(c2,"EINSATZ" if aktiv else "BEREIT",size=8,
              color=fc if aktiv else c(0x99,0x99,0x99),align=WD_ALIGN_PARAGRAPH.CENTER)
    sp2=doc.add_paragraph(); sp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(sp2,f"{bul}/5 im Einsatz  —  {label}",bold=True,size=12,color=fc)
    fp2=doc.add_paragraph()
    _run(fp2,f"  Bulmor-Fahrer: {', '.join(_bulmor_fahrer(data))}",size=9,color=NAVY)
    doc.add_paragraph()

    # Dispo + Betreuer
    bh2=doc.add_paragraph(); _run(bh2,"DISPOSITION",bold=True,size=11,color=DUNKELR)
    _bdr(bh2,"780000","6")
    _schicht_tbl(doc,_gruppen_dispo(data),NAVY,CREME,HELLR)
    doc.add_paragraph()
    bh3=doc.add_paragraph(); _run(bh3,"BEHINDERTENBETREUER",bold=True,size=11,color=DUNKELR)
    _bdr(bh3,"780000","6")
    _schicht_tbl(doc,_gruppen_bet(data),NAVY,CREME,STAHL)

    # Krank
    kl=_kranke(data)
    if kl:
        doc.add_paragraph()
        kp=doc.add_paragraph()
        _run(kp,f"⚕ Krankmeldungen: {', '.join(kl)}",size=10,color=ROT)

    doc.add_paragraph()
    bh4=doc.add_paragraph(); _run(bh4,"MEDIZINPRODUKTE / VERBRAUCHSMATERIAL",bold=True,size=11,color=DUNKELR)
    _bdr(bh4,"780000","6")
    _mat_tbl(doc,NAVY,CREME,HELLR)

    pfad=os.path.join(ZIELORDNER,"D7_FieryOcean_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")


# ═══════════════════════════════════════════════════════════════════════════════
# D8 – GUNMETAL FUTURIST (Charcoal/Platin/Elektrisch-Blau/Peach)
# Inspiriert by: "Sleek and Futuristic" – MediaMonks Award Winner
# ═══════════════════════════════════════════════════════════════════════════════
def design8_gunmetal_futurist(data, bul=2, einz=28, pat=5, pax=42500):
    CHARCOAL = c(0x2B,0x2D,0x42)
    PLATIN   = c(0xED,0xED,0xED)
    ELBLAU   = c(0x00,0x72,0xFF)
    PEACH    = c(0xFF,0x77,0x6A)
    SILBER   = c(0x8D,0x99,0xAE)
    GRN      = c(0x06,0xD6,0xA0)
    ORG      = c(0xFF,0xA0,0x2E)
    ROT      = c(0xEF,0x23,0x3C)

    doc=Document(); _margins(doc,l=1.6,r=1.6); _logo_header(doc,ELBLAU,SILBER); _footer(doc)

    # Titel-Panel
    tbl=doc.add_table(rows=1,cols=2)
    lc=tbl.rows[0].cells[0]; rc=tbl.rows[0].cells[1]
    _shd(lc,CHARCOAL); _shd(rc,c(0x3D,0x40,0x5E))
    lp=lc.paragraphs[0]
    _run(lp,"STÄRKEMELDUNG\nUND EINSÄTZE",bold=True,size=16,color=ELBLAU)
    rp=rc.paragraphs[0]; rp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    _run(rp,"25.03.2026\n07:45 Uhr\nFlughafen Köln/Bonn",size=10,color=PLATIN)
    doc.add_paragraph()

    # Metrics-Row – nur Zahlen groß
    mr=doc.add_table(rows=1,cols=4); mr.style='Table Grid'
    fc,label=_bulmor_col(bul,GRN,ORG,ROT)
    kz=[("Einsätze",str(einz),ELBLAU),("Patienten",str(pat),GRN),
        ("PAX",f"{pax:,}".replace(",","."),SILBER),("Personal",str(_personal(data)),PEACH)]
    for i,(lbl,val,f) in enumerate(kz):
        c_=mr.rows[0].cells[i]; _shd(c_,CHARCOAL)
        p_=c_.paragraphs[0]; p_.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p_,f"{val}\n",bold=True,size=28,color=f)
        _run(p_,lbl,size=8,color=SILBER)
    doc.add_paragraph()

    # Bulmor – Fortschrittsbalken-Optik
    bh=doc.add_paragraph()
    _run(bh,"  BULMOR FLOTTE  ",bold=True,size=11,color=CHARCOAL)
    bh.runs[0].font.highlight_color = None
    _bdr(bh,_s(ELBLAU),"6")
    bt=doc.add_table(rows=3,cols=5); bt.style='Table Grid'
    for i in range(5):
        aktiv=(i+1)<=bul
        r0=bt.rows[0].cells[i]; r1=bt.rows[1].cells[i]; r2=bt.rows[2].cells[i]
        _shd(r0,fc if aktiv else c(0x44,0x44,0x55))
        _shd(r1,CHARCOAL); _shd(r2,c(0x22,0x22,0x33))
        _cell(r0,f"B{i+1}",bold=True,size=11,color=WEISS,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(r1,"▓▓▓ 100%" if aktiv else "▒▒▒  0%",size=8,
              color=fc if aktiv else c(0x55,0x55,0x66),align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(r2,"EINSATZ" if aktiv else "RESERVE",size=7,
              color=fc if aktiv else SILBER,align=WD_ALIGN_PARAGRAPH.CENTER)
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(sp,f"Status: {bul}/5  —  {label}",bold=True,size=12,color=fc)
    fp2=doc.add_paragraph()
    _run(fp2,f"  Fahrer: {', '.join(_bulmor_fahrer(data))}",size=9,color=SILBER)
    doc.add_paragraph()

    # Alle Schichten in einer kompakten Tabelle
    bh2=doc.add_paragraph()
    _run(bh2,"  SCHICHTBESETZUNG",bold=True,size=11,color=ELBLAU)
    _bdr(bh2,_s(ELBLAU),"6")
    alle=_alle_personen_sortiert(data)
    at=doc.add_table(rows=1+len(alle),cols=5); at.style='Table Grid'
    for ci,h in enumerate(["Rolle","Name","Schicht","Kat.","Bulmor-Fhr."]):
        hc=at.rows[0].cells[ci]; _shd(hc,CHARCOAL)
        _cell(hc,h,bold=True,size=8,color=ELBLAU,align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri,(rolle,p) in enumerate(alle,1):
        row=at.rows[ri]
        bg=c(0x22,0x22,0x33) if ri%2==1 else CHARCOAL
        for ci in range(5): _shd(row.cells[ci],bg)
        ist_bul=p.get('ist_bulmorfahrer') in (True,'True')
        s=(p.get('start_zeit') or '')[:5]; e=(p.get('end_zeit') or '')[:5]
        _cell(row.cells[0],rolle,bold=True,size=8,color=ELBLAU if rolle=='DISPO' else SILBER,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[1],p.get('anzeigename',''),size=9,color=PLATIN)
        _cell(row.cells[2],f"{s}–{e}",size=9,color=SILBER,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[3],p.get('dienst_kategorie',''),bold=True,size=8,color=PEACH,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[4],"✔ Ja" if ist_bul else "",bold=ist_bul,size=9,color=GRN,align=WD_ALIGN_PARAGRAPH.CENTER)

    kl=_kranke(data)
    if kl:
        doc.add_paragraph()
        kp=doc.add_paragraph()
        _run(kp,f"Krank: {', '.join(kl)}",size=10,color=ROT)

    doc.add_paragraph()
    bh3=doc.add_paragraph(); _run(bh3,"  MEDIZINPRODUKTE",bold=True,size=11,color=ELBLAU)
    _bdr(bh3,_s(ELBLAU),"6")
    _mat_tbl(doc,CHARCOAL,c(0x22,0x22,0x33),ELBLAU,txt=PLATIN)

    pfad=os.path.join(ZIELORDNER,"D8_GunmetalFuturist_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")


# ═══════════════════════════════════════════════════════════════════════════════
# D9 – TEAL MEDICAL (Teal/Mintgrün/Weiß – sauber, klinisch, modern)
# Inspiriert by: "Clean and Modern" Umwelt A/S (Myrtle Green + Keppel)
# ═══════════════════════════════════════════════════════════════════════════════
def design9_teal_medical(data, bul=5, einz=28, pat=5, pax=42500):
    TEAL     = c(0x00,0x8B,0x8B)
    TEAL_D   = c(0x00,0x5F,0x5F)
    MINT     = c(0xCC,0xF5,0xF0)
    MINT_D   = c(0xA0,0xE8,0xDC)
    GRAU     = c(0x44,0x55,0x55)
    HELLG    = c(0xF4,0xFC,0xFB)
    GRN      = c(0x00,0xAA,0x44)
    ORG      = c(0xE0,0x80,0x00)
    ROT      = c(0xCC,0x22,0x22)

    doc=Document(); _margins(doc,l=2.0,r=2.0); _logo_header(doc,TEAL,GRAU); _footer(doc)

    # Klinischer Titel
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _run(tp,"Stärkemeldung und Einsätze",bold=True,size=24,color=TEAL_D)
    _bdr(tp,_s(TEAL),"10")
    dp=doc.add_paragraph()
    _run(dp,"Erste-Hilfe-Station Flughafen Köln/Bonn  ·  DRK Köln  ·  Mittwoch, 25. März 2026",size=10,color=GRAU)
    doc.add_paragraph()

    # Kennzahlen-Grid (2×2)
    sc=doc.add_table(rows=2,cols=4); sc.style='Table Grid'
    fc,label=_bulmor_col(bul,GRN,ORG,ROT)
    kz=[("Einsätze gesamt",str(einz),TEAL_D),
        ("Patienten auf Station",str(pat),c(0x22,0x88,0x55)),
        ("Passagiere (PAX)",f"{pax:,}".replace(",","."),GRAU),
        ("Bulmor-Status",f"{bul}/5",fc)]
    for i,(lbl,val,f) in enumerate(kz):
        h=sc.rows[0].cells[i]; v=sc.rows[1].cells[i]
        _shd(h,TEAL); _shd(v,HELLG)
        _cell(h,lbl,bold=True,size=8,color=WEISS,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(v,val,bold=True,size=24,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Bulmor – Medizinische Check-Optik
    bh=doc.add_paragraph()
    _run(bh,"Bulmor – Fahrzeugstatus",bold=True,size=12,color=TEAL_D)
    _bdr(bh,_s(TEAL),"8")
    bt=doc.add_table(rows=2,cols=5); bt.style='Table Grid'
    for i in range(5):
        aktiv=(i+1)<=bul
        r0=bt.rows[0].cells[i]; r1=bt.rows[1].cells[i]
        _shd(r0,fc if aktiv else c(0xDD,0xDD,0xDD))
        _shd(r1,MINT if aktiv else HELLG)
        _cell(r0,f"{'✔' if aktiv else '–'}  Bulmor {i+1}",bold=True,size=10,
              color=WEISS,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(r1,"Im Einsatz" if aktiv else "Reserve",size=9,
              color=fc if aktiv else GRAU,align=WD_ALIGN_PARAGRAPH.CENTER)
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(sp,f"  {bul} von 5 Fahrzeugen im Einsatz  —  {label}",bold=True,size=11,color=fc)
    fp2=doc.add_paragraph()
    _run(fp2,f"  Fahrer: {', '.join(_bulmor_fahrer(data))}",size=9,color=TEAL)
    doc.add_paragraph()

    # Dispo
    bh2=doc.add_paragraph(); _run(bh2,"Disposition",bold=True,size=12,color=TEAL_D)
    _bdr(bh2,_s(TEAL),"6")
    _schicht_tbl(doc,_gruppen_dispo(data),TEAL,MINT,TEAL_D)
    doc.add_paragraph()

    # Betreuer
    bh3=doc.add_paragraph(); _run(bh3,"Behindertenbetreuer",bold=True,size=12,color=TEAL_D)
    _bdr(bh3,_s(TEAL),"6")
    _schicht_tbl(doc,_gruppen_bet(data),TEAL,HELLG,TEAL_D)

    kl=_kranke(data)
    if kl:
        doc.add_paragraph()
        kp=doc.add_paragraph()
        _run(kp,f"Krankmeldung: {', '.join(kl)}",size=10,color=ROT,italic=True)

    doc.add_paragraph()
    bh4=doc.add_paragraph(); _run(bh4,"Verbrauchsmaterial (Medizinprodukte)",bold=True,size=12,color=TEAL_D)
    _bdr(bh4,_s(TEAL),"6")
    _mat_tbl(doc,TEAL,MINT,TEAL_D)

    pfad=os.path.join(ZIELORDNER,"D9_TealMedical_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")


# ═══════════════════════════════════════════════════════════════════════════════
# D10 – ROYAL PRESTIGE (Royalblau/Gold/Cyan – Regierungsreport-Stil)
# Inspiriert by: "Audacious and Unique" – royal blue + gold + vivid cyan
# ═══════════════════════════════════════════════════════════════════════════════
def design10_royal_prestige(data, bul=3, einz=28, pat=5, pax=42500):
    ROYAL    = c(0x1A,0x36,0x7C)
    ROYAL_D  = c(0x0C,0x1F,0x50)
    GOLD     = c(0xD4,0xAF,0x37)
    GOLD_H   = c(0xFF,0xE0,0x7A)
    CYAN_V   = c(0x00,0xC9,0xFF)
    HELLB    = c(0xE8,0xF0,0xFF)
    WEISSB   = c(0xF8,0xFA,0xFF)
    GRN      = c(0x00,0xCC,0x88)
    ORG      = c(0xFF,0xA5,0x00)
    ROT      = c(0xFF,0x33,0x44)

    doc=Document(); _margins(doc); _logo_header(doc,GOLD,c(0x44,0x55,0x88)); _footer(doc)

    # Royale Titelzeile mit Gold-Akzent
    tbl=doc.add_table(rows=1,cols=1)
    tc=tbl.rows[0].cells[0]; _shd(tc,ROYAL_D)
    tp=tc.paragraphs[0]; tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(tp,"✦  Stärkemeldung und Einsätze  ✦",bold=True,size=22,color=GOLD)
    sp=tc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(sp,"DRK Kreisverband Köln e.V.  ◈  Flughafen Köln/Bonn  ◈  25. März 2026",size=9,color=CYAN_V)
    doc.add_paragraph()

    # Kennzahlen in Gold-Rahmen
    sc=doc.add_table(rows=2,cols=4); sc.style='Table Grid'
    fc,label=_bulmor_col(bul,GRN,ORG,ROT)
    kz=[("EINSÄTZE",str(einz),CYAN_V),("PAT. AUF STATION",str(pat),GRN),
        ("PAX (PASSAGIERE)",f"{pax:,}".replace(",","."),GOLD),("PERSONAL",str(_personal(data)),CYAN_V)]
    for i,(lbl,val,f) in enumerate(kz):
        h=sc.rows[0].cells[i]; v=sc.rows[1].cells[i]
        _shd(h,ROYAL); _shd(v,HELLB)
        _cell(h,lbl,bold=True,size=8,color=GOLD,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(v,val,bold=True,size=26,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Bulmor – Wappenschild-Stil
    bh=doc.add_paragraph()
    _run(bh,"◈  BULMOR – EINSATZFAHRZEUGE  ◈",bold=True,size=12,color=ROYAL)
    _bdr(bh,_s(GOLD),"8")
    bt=doc.add_table(rows=2,cols=5); bt.style='Table Grid'
    for i in range(5):
        aktiv=(i+1)<=bul
        r0=bt.rows[0].cells[i]; r1=bt.rows[1].cells[i]
        _shd(r0,fc if aktiv else c(0x99,0xAA,0xBB))
        _shd(r1,HELLB)
        sym="▶" if aktiv else "▷"
        _cell(r0,f"{sym}  Bulmor {i+1}",bold=True,size=11,
              color=WEISS,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(r1,"Im Einsatz" if aktiv else "Bereit",size=9,
              color=fc if aktiv else c(0x77,0x88,0x99),align=WD_ALIGN_PARAGRAPH.CENTER)
    sp2=doc.add_paragraph(); sp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(sp2,f"◈  {bul} von 5 Bulmor im Einsatz  —  {label}  ◈",bold=True,size=12,color=fc)
    fp2=doc.add_paragraph()
    _run(fp2,f"   Bulmor-Fahrer: {', '.join(_bulmor_fahrer(data))}",size=9,color=ROYAL)
    doc.add_paragraph()

    # Disposition
    bh2=doc.add_paragraph()
    _run(bh2,"DISPOSITION",bold=True,size=11,color=ROYAL_D); _bdr(bh2,_s(GOLD),"6")
    _schicht_tbl(doc,_gruppen_dispo(data),ROYAL,HELLB,GOLD)
    doc.add_paragraph()

    # Betreuer
    bh3=doc.add_paragraph()
    _run(bh3,"BEHINDERTENBETREUER",bold=True,size=11,color=ROYAL_D); _bdr(bh3,_s(GOLD),"6")
    _schicht_tbl(doc,_gruppen_bet(data),ROYAL,WEISSB,CYAN_V)

    kl=_kranke(data)
    if kl:
        doc.add_paragraph()
        kp=doc.add_paragraph()
        _run(kp,f"◈ Krankmeldung: {', '.join(kl)}",size=10,color=ROT)

    doc.add_paragraph()
    bh4=doc.add_paragraph()
    _run(bh4,"MEDIZINPRODUKTE / VERBRAUCHSMATERIAL",bold=True,size=11,color=ROYAL_D)
    _bdr(bh4,_s(GOLD),"6")
    _mat_tbl(doc,ROYAL,HELLB,GOLD)

    pfad=os.path.join(ZIELORDNER,"D10_RoyalPrestige_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")


# ═══════════════════════════════════════════════════════════════════════════════
# D11 – NORDIC SLATE (Nordisch-minimalistisch: Dunkelgrau/Weiß/Orange-Akzent)
# Inspiriert by "Innovative & Audacious" – Portland Orange, jade, dark gray
# ═══════════════════════════════════════════════════════════════════════════════
def design11_nordic_slate(data, bul=4, einz=28, pat=5, pax=42500):
    SLATE_D  = c(0x2E,0x33,0x38)
    SLATE_M  = c(0x48,0x4E,0x54)
    SLATE_H  = c(0xF2,0xF3,0xF4)
    ORANGE   = c(0xFF,0x6B,0x35)
    JADE     = c(0x2A,0xB5,0x7D)
    WEISS_N  = c(0xFF,0xFF,0xFF)
    GRAU_N   = c(0x99,0x9A,0x9B)
    GRN      = JADE
    ORG      = ORANGE
    ROT      = c(0xFF,0x33,0x33)

    doc=Document(); _margins(doc,l=1.6,r=1.6); _logo_header(doc,ORANGE,GRAU_N); _footer(doc)

    # Minimalistischer Seitentitel
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _run(tp,"Stärkemeldung\nund Einsätze",bold=True,size=28,color=SLATE_D)
    _bdr(tp,_s(ORANGE),"14")
    dp=doc.add_paragraph()
    _run(dp,"25.03.2026  ·  07:45 Uhr  ·  DRK Köln  ·  Flughafen Köln/Bonn",size=9,color=GRAU_N)
    doc.add_paragraph()

    # Metriken – nur Zahl + Label
    mr=doc.add_table(rows=1,cols=4); mr.style='Table Grid'
    fc,label=_bulmor_col(bul,GRN,ORG,ROT)
    kz=[("Einsätze",str(einz),ORANGE),("Patienten",str(pat),JADE),
        ("PAX",f"{pax:,}".replace(",","."),GRAU_N),("Personal",str(_personal(data)),SLATE_M)]
    for i,(lbl,val,f) in enumerate(kz):
        c_=mr.rows[0].cells[i]; _shd(c_,SLATE_H)
        p_=c_.paragraphs[0]; p_.alignment=WD_ALIGN_PARAGRAPH.LEFT
        _run(p_,f"{val}",bold=True,size=30,color=f)
        p2_=c_.add_paragraph()
        _run(p2_,lbl,size=9,color=GRAU_N)
    doc.add_paragraph()

    # Bulmor – minimalistische Zeile
    bh=doc.add_paragraph()
    _run(bh,"Bulmor-Fahrzeuge",bold=True,size=12,color=SLATE_D)
    _bdr(bh,_s(ORANGE),"6")
    bt=doc.add_table(rows=1,cols=5); bt.style='Table Grid'
    for i in range(5):
        aktiv=(i+1)<=bul
        c_=bt.rows[0].cells[i]
        _shd(c_,fc if aktiv else SLATE_H)
        sym="■" if aktiv else "□"
        _cell(c_,f"{sym}  B{i+1}",bold=aktiv,size=11,
              color=WEISS_N if aktiv else GRAU_N,align=WD_ALIGN_PARAGRAPH.CENTER)
    sp=doc.add_paragraph()
    _run(sp,f"  {bul}/5 im Einsatz  —  {label}    Fahrer: {', '.join(_bulmor_fahrer(data))}",size=10,color=fc,bold=True)
    doc.add_paragraph()

    # Zwei-Spalten-Layout: Dispo | Betreuer
    lt=doc.add_table(rows=1,cols=2)
    lc=lt.rows[0].cells[0]; rc=lt.rows[0].cells[1]
    lp_=lc.paragraphs[0]; _run(lp_,"Disposition",bold=True,size=11,color=SLATE_D)
    _bdr(lp_,_s(ORANGE),"6")
    for key,namen in _gruppen_dispo(data).items():
        zp=lc.add_paragraph()
        _run(zp,f"{key}",bold=True,size=9,color=ORANGE)
        np_=lc.add_paragraph(); np_.paragraph_format.left_indent=Cm(0.3)
        _run(np_," / ".join(namen),size=9,color=SLATE_D)
    rp_=rc.paragraphs[0]; _run(rp_,"Betreuer",bold=True,size=11,color=SLATE_D)
    _bdr(rp_,_s(JADE),"6")
    for key,namen in _gruppen_bet(data).items():
        zp2=rc.add_paragraph(); _run(zp2,f"{key}",bold=True,size=9,color=JADE)
        np2_=rc.add_paragraph(); np2_.paragraph_format.left_indent=Cm(0.3)
        _run(np2_," / ".join(namen),size=9,color=SLATE_D)

    kl=_kranke(data)
    if kl:
        doc.add_paragraph()
        kp=doc.add_paragraph()
        _run(kp,f"Krank: {', '.join(kl)}",size=10,color=ROT,italic=True)

    doc.add_paragraph()
    bh2=doc.add_paragraph(); _run(bh2,"Medizinprodukte",bold=True,size=12,color=SLATE_D)
    _bdr(bh2,_s(ORANGE),"6")
    _mat_tbl(doc,SLATE_M,SLATE_H,ORANGE)

    pfad=os.path.join(ZIELORDNER,"D11_NordicSlate_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")


# ═══════════════════════════════════════════════════════════════════════════════
# D12 – EMERALD NIGHT (Smaragdgrün/Dunkel – Nacht-Einsatz-Optik)
# Inspiriert by: "Eye-Catching and Sleek" – Viridian Green + Dark + Telemagenta
# ═══════════════════════════════════════════════════════════════════════════════
def design12_emerald_night(data, bul=1, einz=28, pat=5, pax=42500):
    EMERALD  = c(0x00,0x7C,0x5F)
    EMERALD_H= c(0x00,0xAA,0x80)
    NACHT    = c(0x05,0x14,0x0F)
    DUNKEL   = c(0x0C,0x26,0x1F)
    MITTEL   = c(0x12,0x3D,0x2E)
    MINT_E   = c(0xCC,0xF7,0xEE)
    MAGENT   = c(0xCC,0x14,0x76)
    GOLD     = c(0xFF,0xD1,0x00)
    GRN      = EMERALD_H
    ORG      = c(0xFF,0x85,0x00)
    ROT      = MAGENT

    doc=Document(); _margins(doc); _logo_header(doc,EMERALD_H,EMERALD); _footer(doc)

    # Nachteinsatz-Banner
    tbl=doc.add_table(rows=1,cols=1)
    tc=tbl.rows[0].cells[0]; _shd(tc,NACHT)
    tp=tc.paragraphs[0]; tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(tp,"STÄRKEMELDUNG UND EINSÄTZE",bold=True,size=20,color=EMERALD_H)
    sp=tc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(sp,"DRK KÖLN  ·  FLUGHAFEN KÖLN/BONN  ·  25.03.2026  ·  07:45",size=9,color=GOLD)
    doc.add_paragraph()

    # Scorecard
    sc=doc.add_table(rows=2,cols=4); sc.style='Table Grid'
    fc,label=_bulmor_col(bul,GRN,ORG,ROT)
    kz=[("EINSÄTZE",str(einz),EMERALD_H),("PATIENTEN",str(pat),GOLD),
        ("PAX",f"{pax:,}".replace(",","."),c(0x88,0xCC,0xAA)),("PERSONAL",str(_personal(data)),EMERALD_H)]
    for i,(lbl,val,f) in enumerate(kz):
        h=sc.rows[0].cells[i]; v=sc.rows[1].cells[i]
        _shd(h,DUNKEL); _shd(v,MITTEL)
        _cell(h,lbl,bold=True,size=8,color=EMERALD_H,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(v,val,bold=True,size=26,color=f,align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Bulmor – KRITISCH-Warnung
    bh=doc.add_paragraph()
    _run(bh,"BULMOR – FAHRZEUGSTATUS",bold=True,size=12,color=EMERALD_H)
    _bdr(bh,_s(EMERALD),"8")
    bt=doc.add_table(rows=2,cols=5); bt.style='Table Grid'
    for i in range(5):
        aktiv=(i+1)<=bul
        r0=bt.rows[0].cells[i]; r1=bt.rows[1].cells[i]
        _shd(r0,fc if aktiv else DUNKEL)
        _shd(r1,MITTEL)
        sym="◉" if aktiv else "○"
        _cell(r0,f"{sym}  B{i+1}",bold=True,size=12,
              color=WEISS,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(r1,"EINSATZ" if aktiv else "STAND-BY",size=8,
              color=fc if aktiv else c(0x44,0x77,0x55),align=WD_ALIGN_PARAGRAPH.CENTER)
    sp2=doc.add_paragraph(); sp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _run(sp2,f"⚠  {bul} von 5 Bulmor im Einsatz  —  {label}  ⚠",bold=True,size=12,color=fc)
    fp2=doc.add_paragraph()
    _run(fp2,f"  Fahrer: {', '.join(_bulmor_fahrer(data))}",size=9,color=EMERALD_H)
    doc.add_paragraph()

    # Schichten
    bh2=doc.add_paragraph(); _run(bh2,"DISPOSITION",bold=True,size=11,color=EMERALD_H)
    _bdr(bh2,_s(EMERALD),"6")
    _schicht_tbl(doc,_gruppen_dispo(data),DUNKEL,MITTEL,EMERALD_H,txt=MINT_E)
    doc.add_paragraph()
    bh3=doc.add_paragraph(); _run(bh3,"BEHINDERTENBETREUER",bold=True,size=11,color=EMERALD_H)
    _bdr(bh3,_s(EMERALD),"6")
    _schicht_tbl(doc,_gruppen_bet(data),DUNKEL,MITTEL,EMERALD_H,txt=MINT_E)

    kl=_kranke(data)
    if kl:
        doc.add_paragraph()
        kp=doc.add_paragraph()
        _run(kp,f"KRANK: {', '.join(kl)}",size=10,color=MAGENT,bold=True)

    doc.add_paragraph()
    bh4=doc.add_paragraph(); _run(bh4,"MEDIZINPRODUKTE / VERBRAUCHSMATERIAL",bold=True,size=11,color=EMERALD_H)
    _bdr(bh4,_s(EMERALD),"6")
    _mat_tbl(doc,DUNKEL,MITTEL,EMERALD_H,txt=MINT_E)

    pfad=os.path.join(ZIELORDNER,"D12_EmeraldNight_25032026.docx")
    doc.save(pfad); print(f"[OK] {os.path.basename(pfad)}")


# ═══════════════════════════════════════════════════════════════════════════════
# HILFSFUNKTIONEN (geteilt)
# ═══════════════════════════════════════════════════════════════════════════════
def _schicht_tbl(doc, gruppen, hdr_bg, row_bg, acc, txt=None):
    txt_ = txt or SCHWARZ
    if not gruppen:
        doc.add_paragraph().add_run("(keine Einträge)").font.size=Pt(9)
        return
    tbl=doc.add_table(rows=1+len(gruppen),cols=3); tbl.style='Table Grid'
    for ci,h in enumerate(["Schicht","Mitarbeiter","Anz."]):
        hc=tbl.rows[0].cells[ci]; _shd(hc,hdr_bg)
        _cell(hc,h,bold=True,size=9,color=acc,align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri,(key,namen) in enumerate(gruppen.items(),1):
        row=tbl.rows[ri]
        bg=row_bg if ri%2==1 else c(max(0,row_bg[0]-15 if hasattr(row_bg,'__getitem__') else 0xFF),
                                     0xFF,0xFF)
        for ci in range(3): _shd(row.cells[ci],row_bg)
        _cell(row.cells[0],key,bold=True,size=9,color=acc)
        _cell(row.cells[1]," / ".join(namen),size=9,color=txt_)
        _cell(row.cells[2],str(len(namen)),bold=True,size=10,color=acc,align=WD_ALIGN_PARAGRAPH.CENTER)

def _schicht_tbl_v2(doc, gruppen, hdr_bg, row_bg_a, row_bg_b, acc, txt_col):
    if not gruppen: return
    tbl=doc.add_table(rows=1+len(gruppen),cols=3); tbl.style='Table Grid'
    for ci,h in enumerate(["Schicht","Mitarbeiter","#"]):
        hc=tbl.rows[0].cells[ci]; _shd(hc,hdr_bg)
        _cell(hc,h,bold=True,size=9,color=acc,align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri,(key,namen) in enumerate(gruppen.items(),1):
        row=tbl.rows[ri]
        bg=row_bg_a if ri%2==1 else row_bg_b
        for ci in range(3): _shd(row.cells[ci],bg)
        _cell(row.cells[0],key,bold=True,size=9,color=acc)
        _cell(row.cells[1]," / ".join(namen),size=9,color=txt_col)
        _cell(row.cells[2],str(len(namen)),bold=True,size=10,color=acc,align=WD_ALIGN_PARAGRAPH.CENTER)

def _mat_tbl(doc, hdr_bg, row_bg, acc, txt=None):
    txt_=txt or SCHWARZ
    tbl=doc.add_table(rows=1+len(MAT),cols=4); tbl.style='Table Grid'
    for ci,h in enumerate(["Material","Einheit","Soll","Mindest"]):
        hc=tbl.rows[0].cells[ci]; _shd(hc,hdr_bg)
        _cell(hc,h,bold=True,size=9,color=acc,align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri,(mat,eh,soll,mind) in enumerate(MAT,1):
        row=tbl.rows[ri]
        bg=row_bg if ri%2==1 else c(0xFF,0xFF,0xFF)
        for ci in range(4): _shd(row.cells[ci],bg)
        _cell(row.cells[0],mat,size=9,color=txt_)
        _cell(row.cells[1],eh,size=9,color=txt_,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[2],soll,bold=True,size=9,color=txt_,align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[3],mind,bold=True,size=9,color=ROT_A,align=WD_ALIGN_PARAGRAPH.CENTER)

def _alle_personen_sortiert(d):
    alle=[]
    for p in d.get('dispo',[]):
        if p.get('ist_krank') not in (True,'True'): alle.append(('DISPO',p))
    for p in d.get('betreuer',[]):
        if p.get('ist_krank') not in (True,'True'): alle.append(('BET.',p))
    alle.sort(key=lambda x:(x[1].get('start_zeit') or 'ZZZZ'))
    return alle

# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Lade Dienstplan ...")
    data=lade_dienstplan()
    bet=len(data.get('betreuer',[])); dis=len(data.get('dispo',[])); kra=len(data.get('kranke',[]))
    print(f"  Betreuer: {bet}  |  Dispo: {dis}  |  Krank: {kra}")
    print(f"\nZielordner: {ZIELORDNER}\n")

    # Verschiedene Bulmor-Szenarien pro Design (1=rot, 2=rot, 3=orange, 4=grün, 5=grün)
    print("Erstelle 6 neue Designs ...")
    design7_fiery_ocean(data,     bul=4, einz=28, pat=5, pax=42500)
    design8_gunmetal_futurist(data,bul=2,einz=28, pat=5, pax=42500)
    design9_teal_medical(data,    bul=5, einz=28, pat=5, pax=42500)
    design10_royal_prestige(data, bul=3, einz=28, pat=5, pax=42500)
    design11_nordic_slate(data,   bul=4, einz=28, pat=5, pax=42500)
    design12_emerald_night(data,  bul=1, einz=28, pat=5, pax=42500)

    print(f"\n✓ Alle 6 neuen Beispiele gespeichert in:\n  {ZIELORDNER}")
