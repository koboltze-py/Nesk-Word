# -*- coding: utf-8 -*-
"""
Erstellt 3 Design-Beispiele fuer die neue Word-Staerkemeldung
Zielordner: Desktop/bei
"""
import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ZIELORDNER = (
    r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband "
    r"Koeln e.V\Desktop\bei"
)
# Versuche den echten Pfad mit Sonderzeichen
import glob as _glob
_onedrive = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
if os.path.exists(_onedrive):
    ZIELORDNER = os.path.join(_onedrive, "Desktop", "bei")

LOGO_PFAD = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                         "Daten", "Email", "Logo.jpg")

# ── Farben ────────────────────────────────────────────────────────────────────
DRK_ROT        = RGBColor(0xBE, 0x00, 0x00)     # DRK-Rot
DRK_DUNKELROT  = RGBColor(0x8B, 0x00, 0x00)     # dunkleres Rot
WEISS          = RGBColor(0xFF, 0xFF, 0xFF)
GRAU_HELL      = RGBColor(0xF5, 0xF5, 0xF5)
GRAU_MITTEL    = RGBColor(0xD0, 0xD0, 0xD0)
GRAU_TEXT      = RGBColor(0x55, 0x55, 0x55)
SCHWARZ        = RGBColor(0x1A, 0x1A, 0x1A)

STATUS_GRÜN    = RGBColor(0x10, 0x7E, 0x3E)    # 4-5 Bulmor
STATUS_ORANGE  = RGBColor(0xFF, 0x89, 0x00)    # 3 Bulmor
STATUS_ROT     = RGBColor(0xBE, 0x00, 0x00)    # 1-2 Bulmor

BLAU_DUNKEL    = RGBColor(0x1F, 0x49, 0x7D)    # Akzentblau
BLAU_MITTEL    = RGBColor(0x27, 0x6F, 0xBF)
BLAU_HELL      = RGBColor(0xD6, 0xE4, 0xF5)

# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def _rgb_hex(rgb: RGBColor) -> str:
    """RGBColor -> 6-stelliger Hex-String. str(RGBColor) gibt direkt 'RRGGBB'."""
    return str(rgb)

def _set_cell_bg(cell, rgb: RGBColor):
    """Zellhintergrundfarbe setzen."""
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    shd   = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _rgb_hex(rgb))
    tcPr.append(shd)

def _set_para_border_bottom(para, color_hex="888888", sz="6"):
    """Untere Linie unter einem Paragraphen."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def _add_header_drk(doc):
    """Einheitliche DRK-Kopfzeile für alle Varianten."""
    section = doc.sections[0]
    header  = section.header
    htab    = header.add_table(rows=1, cols=2, width=Inches(6.5))
    htab.autofit = False

    logo_path = Path(LOGO_PFAD)
    if logo_path.exists():
        lp   = htab.rows[0].cells[0].paragraphs[0]
        lrun = lp.add_run()
        lrun.add_picture(str(logo_path), width=Inches(1.1))
    else:
        htab.rows[0].cells[0].paragraphs[0].add_run("DRK")

    rp = htab.rows[0].cells[1].paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = rp.add_run("Deutsches Rotes Kreuz Kreisverband Köln e.V.\n")
    r1.font.size  = Pt(9)
    r1.font.bold  = True
    r1.font.color.rgb = DRK_ROT
    r2 = rp.add_run("Unfallhilfsstelle · Erste-Hilfe-Station Flughafen Köln/Bonn")
    r2.font.size  = Pt(8)
    r2.font.color.rgb = GRAU_TEXT

    tp = header.add_paragraph()
    _set_para_border_bottom(tp, "BE0000", "8")

def _add_footer_drk(doc):
    """Einheitliche DRK-Fußzeile."""
    section  = doc.sections[0]
    footer   = section.footer
    fp       = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("☎ +49 2203 40-2323   |   ✉ flughafen@drk-koeln.de   |   Stationsleitung: Lars Peters")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAU_TEXT

def _bulmor_farbe(im_einsatz: int, gesamt: int = 5) -> RGBColor:
    if im_einsatz <= 2:
        return STATUS_ROT
    elif im_einsatz == 3:
        return STATUS_ORANGE
    else:
        return STATUS_GRÜN

def _bulmor_label(im_einsatz: int) -> str:
    if im_einsatz <= 2:
        return "KRITISCH"
    elif im_einsatz == 3:
        return "EINGESCHRÄNKT"
    else:
        return "VOLLSTÄNDIG"


# =============================================================================
# VARIANTE A – Kompakt & Tabellenfokus (blau-grau-professionell)
# =============================================================================
def erstelle_variante_a(zielordner: str):
    doc = Document()

    # Seitenränder
    for s in doc.sections:
        s.left_margin   = Cm(2)
        s.right_margin  = Cm(2)
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.0)

    _add_header_drk(doc)
    _add_footer_drk(doc)

    # ── Hauptüberschrift ──────────────────────────────────────────────────────
    titel = doc.add_paragraph()
    titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = titel.add_run("Stärkemeldung und Einsätze")
    tr.font.size  = Pt(20)
    tr.font.bold  = True
    tr.font.color.rgb = BLAU_DUNKEL

    # Datum & Zeitraum
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Zeitraum: 25.03.2026 bis 25.03.2026   |   erstellt: 25.03.2026, 07:45 Uhr")
    mr.font.size  = Pt(9)
    mr.font.color.rgb = GRAU_TEXT
    _set_para_border_bottom(meta, "1F497D", "6")

    doc.add_paragraph()  # Abstand

    # ── BLOCK 1: Bulmor-Status ────────────────────────────────────────────────
    bul_h = doc.add_paragraph()
    bhr   = bul_h.add_run("▪  Bulmor – Fahrzeugstatus")
    bhr.font.size  = Pt(12)
    bhr.font.bold  = True
    bhr.font.color.rgb = BLAU_DUNKEL

    # Bulmor-Tabelle: 5 Spalten (eine je Fahrzeug)
    bul_table = doc.add_table(rows=2, cols=6)
    bul_table.style = 'Table Grid'
    bul_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Beispiel: 4 von 5 im Einsatz
    im_einsatz = 4
    gesamt     = 5

    hdr_row = bul_table.rows[0]
    _set_cell_bg(hdr_row.cells[0], BLAU_DUNKEL)
    lbl = hdr_row.cells[0].paragraphs[0]
    lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = lbl.add_run("BULMOR")
    r.font.bold  = True
    r.font.size  = Pt(10)
    r.font.color.rgb = WEISS

    for i in range(1, 6):
        cell  = hdr_row.cells[i]
        aktiv = i <= im_einsatz
        bg    = _bulmor_farbe(im_einsatz) if aktiv else GRAU_MITTEL
        _set_cell_bg(cell, bg)
        p     = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = p.add_run(f"B{i}")
        cr.font.bold  = True
        cr.font.size  = Pt(10)
        cr.font.color.rgb = WEISS

    # Zweite Zeile: Status-Text
    status_row = bul_table.rows[1]
    farbe      = _bulmor_farbe(im_einsatz)
    label      = _bulmor_label(im_einsatz)

    stat_cell = status_row.cells[0]
    stat_cell.merge(status_row.cells[5])
    _set_cell_bg(stat_cell, GRAU_HELL)
    sp = stat_cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"Im Einsatz: {im_einsatz} von {gesamt}   –   Status: {label}")
    sr.font.size  = Pt(11)
    sr.font.bold  = True
    sr.font.color.rgb = farbe

    doc.add_paragraph()  # Abstand

    # ── BLOCK 2: Einsatzzahlen ────────────────────────────────────────────────
    eit_h = doc.add_paragraph()
    ehr   = eit_h.add_run("▪  Einsatzzahlen")
    ehr.font.size  = Pt(12)
    ehr.font.bold  = True
    ehr.font.color.rgb = BLAU_DUNKEL

    eins_table = doc.add_table(rows=2, cols=2)
    eins_table.style = 'Table Grid'
    eins_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ["Einsätze gesamt", "Patienten auf Station"]
    werte   = ["17", "3"]

    for i, (h, w) in enumerate(zip(headers, werte)):
        hc = eins_table.rows[0].cells[i]
        vc = eins_table.rows[1].cells[i]
        _set_cell_bg(hc, BLAU_DUNKEL)
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr_ = hc.paragraphs[0].add_run(h)
        hr_.font.bold = True; hr_.font.size = Pt(10); hr_.font.color.rgb = WEISS
        _set_cell_bg(vc, GRAU_HELL)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vc.paragraphs[0].add_run(w)
        vr.font.size = Pt(22); vr.font.bold = True; vr.font.color.rgb = BLAU_DUNKEL

    doc.add_paragraph()  # Abstand

    # ── BLOCK 3: Disposition ──────────────────────────────────────────────────
    disp_h = doc.add_paragraph()
    dhr    = disp_h.add_run("▪  Disposition")
    dhr.font.size  = Pt(12)
    dhr.font.bold  = True
    dhr.font.color.rgb = BLAU_DUNKEL

    disp_table = doc.add_table(rows=4, cols=3)
    disp_table.style = 'Table Grid'

    spalten = ["Schicht", "Namen", "Stärke"]
    for i, s in enumerate(spalten):
        c = disp_table.rows[0].cells[i]
        _set_cell_bg(c, BLAU_DUNKEL)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ = p.add_run(s)
        r_.font.bold = True; r_.font.size = Pt(10); r_.font.color.rgb = WEISS

    daten = [
        ("06:00 bis 14:00", "Müller / Schmidt / Meier / Klein", "4"),
        ("14:00 bis 22:00", "Wagner / Bauer / Fischer",         "3"),
        ("22:00 bis 06:00", "Hoffmann / Schäfer",               "2"),
    ]
    for ri, (schicht, namen, staerke) in enumerate(daten, 1):
        row = disp_table.rows[ri]
        bg  = WEISS if ri % 2 == 1 else GRAU_HELL
        for ci in range(3):
            _set_cell_bg(row.cells[ci], bg)
        row.cells[0].paragraphs[0].add_run(schicht).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(namen).font.size   = Pt(10)
        nrp = row.cells[2].paragraphs[0]
        nrp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nrp.add_run(staerke).font.size = Pt(10)

    doc.add_paragraph()

    # ── BLOCK 4: Behindertenbetreuer ──────────────────────────────────────────
    bet_h = doc.add_paragraph()
    btr   = bet_h.add_run("▪  Behindertenbetreuer")
    btr.font.size  = Pt(12)
    btr.font.bold  = True
    btr.font.color.rgb = BLAU_DUNKEL

    bet = doc.add_paragraph()
    bet.add_run("06:00 bis 14:00     Weber / Keller").font.size = Pt(10)

    # ── PAX ───────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    pax_p = doc.add_paragraph()
    pax_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = pax_p.add_run("PAX:  — 42.500 —")
    pr.font.size = Pt(13)
    pr.font.bold = True
    pr.font.color.rgb = BLAU_DUNKEL

    pfad = os.path.join(zielordner, "Variante_A_Profi-Tabelle.docx")
    doc.save(pfad)
    print(f"[OK] Variante A gespeichert: {pfad}")


# =============================================================================
# VARIANTE B – Modern & Visuell (DRK-Rot als Akzentfarbe)
# =============================================================================
def erstelle_variante_b(zielordner: str):
    doc = Document()

    for s in doc.sections:
        s.left_margin   = Cm(1.8)
        s.right_margin  = Cm(1.8)
        s.top_margin    = Cm(2.8)
        s.bottom_margin = Cm(2.0)

    _add_header_drk(doc)
    _add_footer_drk(doc)

    # ── Titel-Banner ──────────────────────────────────────────────────────────
    banner_tbl = doc.add_table(rows=1, cols=1)
    banner_cell = banner_tbl.rows[0].cells[0]
    _set_cell_bg(banner_cell, DRK_ROT)
    bp = banner_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("Stärkemeldung und Einsätze")
    br.font.size  = Pt(22)
    br.font.bold  = True
    br.font.color.rgb = WEISS

    sub_p = banner_cell.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("DRK Köln · Flughafen Köln/Bonn · 25.03.2026")
    sr.font.size  = Pt(10)
    sr.font.color.rgb = RGBColor(0xFF, 0xCC, 0xCC)

    doc.add_paragraph()  # Abstand

    # ── Kennzahlen-Zeile ──────────────────────────────────────────────────────
    kz = doc.add_table(rows=2, cols=3)
    kz.style = 'Table Grid'
    kz.alignment = WD_TABLE_ALIGNMENT.LEFT

    kz_labels = ["Einsätze gesamt", "Patienten auf Station", "Personal gesamt"]
    kz_werte  = ["17", "3", "9"]
    kz_farben = [BLAU_DUNKEL, DRK_ROT, STATUS_GRÜN]

    for i, (lbl, wert, farbe) in enumerate(zip(kz_labels, kz_werte, kz_farben)):
        hc = kz.rows[0].cells[i]
        vc = kz.rows[1].cells[i]
        _set_cell_bg(hc, farbe)
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr_ = hc.paragraphs[0].add_run(lbl)
        hr_.font.bold = True; hr_.font.size = Pt(9); hr_.font.color.rgb = WEISS
        _set_cell_bg(vc, GRAU_HELL)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vc.paragraphs[0].add_run(wert)
        vr.font.size = Pt(28); vr.font.bold = True; vr.font.color.rgb = farbe

    doc.add_paragraph()

    # ── Bulmor-Status ─────────────────────────────────────────────────────────
    bh = doc.add_paragraph()
    bhr = bh.add_run("BULMOR – FAHRZEUGSTATUS")
    bhr.font.size  = Pt(11)
    bhr.font.bold  = True
    bhr.font.color.rgb = DRK_ROT
    _set_para_border_bottom(bh, "BE0000", "6")

    # Beispiel: 2 von 5 (kritisch → rot)
    szenarien = [
        ("Szenario: 2 von 5 (Rot – kritisch)",   2),
        ("Szenario: 3 von 5 (Orange – eingeschränkt)", 3),
        ("Szenario: 4 von 5 (Grün – vollständig)",    4),
    ]

    for bez, im_einsatz in szenarien:
        # Beschriftung
        sp = doc.add_paragraph()
        sr_ = sp.add_run(f"  {bez}")
        sr_.font.size = Pt(9)
        sr_.font.color.rgb = GRAU_TEXT
        sr_.font.italic = True

        bt = doc.add_table(rows=1, cols=5)
        bt.style = 'Table Grid'
        farbe = _bulmor_farbe(im_einsatz)

        for i in range(5):
            c  = bt.rows[0].cells[i]
            bg = farbe if (i + 1) <= im_einsatz else GRAU_MITTEL
            _set_cell_bg(c, bg)
            p  = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sym  = "●" if (i + 1) <= im_einsatz else "○"
            cr_  = p.add_run(f"{sym}  B{i+1}")
            cr_.font.size = Pt(11)
            cr_.font.bold = True
            cr_.font.color.rgb = WEISS if (i + 1) <= im_einsatz else GRAU_TEXT

        stat_p = doc.add_paragraph()
        stat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st_r = stat_p.add_run(f"▶  {im_einsatz} / 5 im Einsatz  –  {_bulmor_label(im_einsatz)}")
        st_r.font.size = Pt(10); st_r.font.bold = True; st_r.font.color.rgb = farbe

        doc.add_paragraph()

    # ── Disposition ───────────────────────────────────────────────────────────
    dh = doc.add_paragraph()
    dhr_ = dh.add_run("DISPOSITION")
    dhr_.font.size = Pt(11); dhr_.font.bold = True; dhr_.font.color.rgb = DRK_ROT
    _set_para_border_bottom(dh, "BE0000", "6")

    dt = doc.add_table(rows=4, cols=2)
    dt.style = 'Table Grid'

    for i, (label, val) in enumerate(zip(["Schicht", "Namen"],["Schicht", "Namen"])):
        for ci, txt in enumerate(["Schicht", "Mitarbeiter"]):
            c = dt.rows[0].cells[ci]
            _set_cell_bg(c, DRK_ROT)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_ = p.add_run(txt)
            r_.font.bold = True; r_.font.size = Pt(10); r_.font.color.rgb = WEISS
        break

    for ri, (schicht, namen) in enumerate([
        ("06:00 – 14:00", "Müller / Schmidt / Meier / Klein"),
        ("14:00 – 22:00", "Wagner / Bauer / Fischer"),
        ("22:00 – 06:00", "Hoffmann / Schäfer"),
    ], 1):
        row = dt.rows[ri]
        bg  = WEISS if ri % 2 == 1 else GRAU_HELL
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)
        row.cells[0].paragraphs[0].add_run(schicht).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(namen).font.size   = Pt(10)

    doc.add_paragraph()
    pax = doc.add_paragraph()
    pax.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr_ = pax.add_run("✈  PAX:  42.500  ✈")
    pr_.font.size = Pt(14); pr_.font.bold = True; pr_.font.color.rgb = DRK_ROT

    pfad = os.path.join(zielordner, "Variante_B_DRK-Modern.docx")
    doc.save(pfad)
    print(f"[OK] Variante B gespeichert: {pfad}")


# =============================================================================
# VARIANTE C – Dashboard-Stil (dunkel & sehr visuell)
# =============================================================================
def erstelle_variante_c(zielordner: str):
    doc = Document()

    for s in doc.sections:
        s.left_margin   = Cm(1.5)
        s.right_margin  = Cm(1.5)
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.0)

    _add_header_drk(doc)
    _add_footer_drk(doc)

    DUNKEL  = RGBColor(0x1A, 0x1A, 0x2E)   # sehr dunkelblau
    MITTEL  = RGBColor(0x16, 0x21, 0x3E)
    AKZENT  = RGBColor(0x0F, 0x3A, 0x60)
    CYAN    = RGBColor(0x00, 0xB4, 0xD8)
    GELB    = RGBColor(0xFF, 0xD1, 0x66)

    # ── Titel ─────────────────────────────────────────────────────────────────
    tit_tbl = doc.add_table(rows=1, cols=1)
    tc      = tit_tbl.rows[0].cells[0]
    _set_cell_bg(tc, DUNKEL)
    tp_     = tc.paragraphs[0]
    tp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr_     = tp_.add_run("STÄRKEMELDUNG UND EINSÄTZE")
    tr_.font.size = Pt(18); tr_.font.bold = True; tr_.font.color.rgb = CYAN

    sub_    = tc.add_paragraph()
    sub_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr_     = sub_.add_run("DRK Köln  ·  Flughafen Köln/Bonn  ·  25.03.2026  ·  07:45 Uhr")
    sr_.font.size = Pt(9); sr_.font.color.rgb = GELB

    doc.add_paragraph()

    # ── Scorecard-Zeile ───────────────────────────────────────────────────────
    sc = doc.add_table(rows=2, cols=4)
    sc.style = 'Table Grid'

    items = [
        ("Einsätze",         "17",     CYAN),
        ("Pat. auf Station",  "3",     STATUS_GRÜN),
        ("PAX (Passagiere)", "42.500", GELB),
        ("Personal",          "9",     RGBColor(0xCC, 0x99, 0xFF)),
    ]

    for i, (lbl, val, farbe) in enumerate(items):
        hc = sc.rows[0].cells[i]; vc = sc.rows[1].cells[i]
        _set_cell_bg(hc, MITTEL); _set_cell_bg(vc, AKZENT)
        hp = hc.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr__ = hp.add_run(lbl)
        hr__.font.size = Pt(9); hr__.font.bold = True; hr__.font.color.rgb = farbe
        vp = vc.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr__ = vp.add_run(val)
        vr__.font.size = Pt(24); vr__.font.bold = True; vr__.font.color.rgb = farbe

    doc.add_paragraph()

    # ── Bulmor: visuelle Ampelleiste ──────────────────────────────────────────
    bh_ = doc.add_paragraph()
    bhr__ = bh_.add_run("BULMOR – SPEZIALFAHRZEUGE (EINSATZSTATUS)")
    bhr__.font.size = Pt(11); bhr__.font.bold = True; bhr__.font.color.rgb = CYAN
    _set_para_border_bottom(bh_, "00B4D8", "6")

    # Aktives Szenario: 3 von 5
    im_einsatz = 3
    farbe_bul  = _bulmor_farbe(im_einsatz)

    bt_ = doc.add_table(rows=2, cols=5)
    bt_.style = 'Table Grid'

    for i in range(5):
        aktiv = (i + 1) <= im_einsatz
        c_  = bt_.rows[0].cells[i]
        c2_ = bt_.rows[1].cells[i]
        bg  = farbe_bul if aktiv else RGBColor(0x44, 0x44, 0x55)
        _set_cell_bg(c_,  bg)
        _set_cell_bg(c2_, AKZENT if aktiv else MITTEL)
        p_  = c_.paragraphs[0]; p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sym_ = "◉" if aktiv else "○"
        cr2_ = p_.add_run(f"{sym_}")
        cr2_.font.size = Pt(16); cr2_.font.bold = True
        cr2_.font.color.rgb = WEISS if aktiv else RGBColor(0x88, 0x88, 0x99)
        p2_ = c2_.paragraphs[0]; p2_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr_ = p2_.add_run(f"Bulmor {i+1}")
        lr_.font.size = Pt(9); lr_.font.bold = True
        lr_.font.color.rgb = farbe_bul if aktiv else RGBColor(0x88, 0x88, 0x99)

    stp_ = doc.add_paragraph()
    stp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    str_ = stp_.add_run(f"Status: {im_einsatz} von 5 im Einsatz – {_bulmor_label(im_einsatz)}")
    str_.font.size = Pt(12); str_.font.bold = True; str_.font.color.rgb = farbe_bul

    doc.add_paragraph()

    # ── Dienstübersicht ───────────────────────────────────────────────────────
    dh_ = doc.add_paragraph()
    dhr2  = dh_.add_run("DIENSTÜBERSICHT")
    dhr2.font.size = Pt(11); dhr2.font.bold = True; dhr2.font.color.rgb = CYAN
    _set_para_border_bottom(dh_, "00B4D8", "6")

    dt_ = doc.add_table(rows=4, cols=3)
    dt_.style = 'Table Grid'

    for ci, txt in enumerate(["Schicht", "Mitarbeiter", "Anzahl"]):
        c = dt_.rows[0].cells[ci]
        _set_cell_bg(c, MITTEL)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ = p.add_run(txt)
        r_.font.bold = True; r_.font.size = Pt(10); r_.font.color.rgb = CYAN

    schichten = [
        ("06:00 – 14:00", "Müller / Schmidt / Meier / Klein", "4"),
        ("14:00 – 22:00", "Wagner / Bauer / Fischer",         "3"),
        ("22:00 – 06:00", "Hoffmann / Schäfer",               "2"),
    ]
    for ri, (sch, nam, anz) in enumerate(schichten, 1):
        row  = dt_.rows[ri]
        bg   = GRAU_HELL
        for ci in range(3): _set_cell_bg(row.cells[ci], bg)
        row.cells[0].paragraphs[0].add_run(sch).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(nam).font.size = Pt(10)
        azp = row.cells[2].paragraphs[0]
        azp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        azr = azp.add_run(anz)
        azr.font.size = Pt(10); azr.font.bold = True; azr.font.color.rgb = BLAU_DUNKEL

    pfad = os.path.join(zielordner, "Variante_C_Dashboard.docx")
    doc.save(pfad)
    print(f"[OK] Variante C gespeichert: {pfad}")


# =============================================================================
# MAIN
# =============================================================================
if __name__ == "__main__":
    os.makedirs(ZIELORDNER, exist_ok=True)

    print("Erstelle Beispiel-Dokumente ...")
    erstelle_variante_a(ZIELORDNER)
    erstelle_variante_b(ZIELORDNER)
    erstelle_variante_c(ZIELORDNER)
    print("\nFertig! Alle 3 Beispiele gespeichert in:")
    print(f"  {ZIELORDNER}")
