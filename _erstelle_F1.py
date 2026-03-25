# -*- coding: utf-8 -*-
"""
F1_25032026.docx – Stärkemeldung (finaler Entwurf)
Änderungen gegenüber W_S3_OzeanBlau:
  - Keine Bulmor-Fahrer-Liste
  - Telefonnummer unter Erste-Hilfe-Station (rechter Header)
  - Uhrzeit entfernt
  - Schichtleiter-Sektion unter Disposition
"""
import os, sys
from pathlib import Path
from collections import defaultdict

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Pfade ─────────────────────────────────────────────────────────────────────
_OD   = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
ZIEL  = os.path.join(_OD, "Desktop", "bei") if os.path.exists(_OD) else r"C:\Temp\bei"
os.makedirs(ZIEL, exist_ok=True)
LOGO  = Path(os.path.dirname(os.path.abspath(__file__))) / "Daten" / "Email" / "Logo.jpg"
EXCEL = (_OD + r"\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26"
         r"\04_Tagesdienstpläne\03_März\25.03.2026.xlsx")
DATUM   = "25.03.2026"
STATION = "Erste-Hilfe-Station · Flughafen Köln/Bonn"
TEL     = "+49 2203 40-2323"
MAIL    = "flughafen@drk-koeln.de"

# ── Farbschema: OzeanBlau ─────────────────────────────────────────────────────
BG  = "1A3460"   # linkes Panel
DK  = "0F1F3C"   # dunklerer Ton
HE  = "C8DAFF"   # heller Ton
AZ  = "00C8FF"   # Tag-Akzent
AZ2 = "00A878"   # Nacht-Akzent

# ── Daten laden ────────────────────────────────────────────────────────────────
def lade():
    try:
        from functions.dienstplan_parser import DienstplanParser
        r = DienstplanParser(EXCEL, alle_anzeigen=True).parse()
        if r.get("success") in (True, "True"): return r
    except Exception as e:
        print(f"  [Parser] {e}")
    return {"betreuer": [], "dispo": [], "kranke": []}

def _ist_tag(p):
    s = (p.get('start_zeit') or '00:00')[:5]
    try: return int(s.split(':')[0]) < 14
    except: return True

def _zeit(p):
    s = (p.get('start_zeit') or '')[:5]
    e = (p.get('end_zeit')   or '')[:5]
    return f"{s}–{e}" if s and e else ''

def _nach_schicht(d, typ='betreuer'):
    dt, dn = [], []
    for p in d.get(typ, []):
        if p.get('ist_krank') in (True, 'True'): continue
        name = p.get('anzeigename', '').strip()
        if not name: continue
        if _ist_tag(p): dt.append((name, _zeit(p)))
        else:           dn.append((name, _zeit(p)))
    dt.sort(key=lambda x: x[0]); dn.sort(key=lambda x: x[0])
    return {'DT': dt, 'DN': dn}

def _kranke(d):
    out = []
    for p in d.get('betreuer', []) + d.get('dispo', []):
        if p.get('ist_krank') in (True, 'True') and p.get('anzeigename', '').strip():
            out.append((p.get('anzeigename', '').strip(), _zeit(p)))
    return out

def _schichtleiter(d):
    """Sucht Schichtleiter im Daten-Dict, Fallback: erster Dispo-DT"""
    # Direkt im Dict (falls Parser dieses Feld liefert)
    if d.get('schichtleiter'):
        sl = d['schichtleiter']
        if isinstance(sl, dict):
            return sl.get('anzeigename', '').strip(), _zeit(sl)
        if isinstance(sl, str):
            return sl.strip(), ''
    # Fallback: erste Dispo-Person mit Tagdienst die als Schichtleiter markiert ist
    for p in d.get('dispo', []):
        if p.get('ist_krank') in (True, 'True'): continue
        if p.get('ist_schichtleiter') in (True, 'True'):
            return p.get('anzeigename', '').strip(), _zeit(p)
    # Letzter Fallback: ersten DT-Dispo nehmen
    for p in d.get('dispo', []):
        if p.get('ist_krank') in (True, 'True'): continue
        if _ist_tag(p):
            return p.get('anzeigename', '').strip(), _zeit(p)
    return '', ''

def _bulmor_anzahl(d):
    """Zählt verfügbare Bulmor (Fahrer die kein Krank sind)"""
    return len([p for p in d.get('betreuer', []) + d.get('dispo', [])
                if p.get('ist_bulmorfahrer') in (True, 'True')
                and p.get('ist_krank') not in (True, 'True')])

# ── Word-Hilfsfunktionen ───────────────────────────────────────────────────────
def _rgb(hx):
    hx = hx.lstrip('#')
    return RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

def _set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#').upper())
    tcPr.append(shd)

def _no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcB.append(b)
    tcPr.append(tcB)

def _p(cell, text, bold=False, size=9, fg="000000", align="left", sa=0, sb=0):
    p = cell.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                   else WD_ALIGN_PARAGRAPH.RIGHT if align == "right"
                   else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = _rgb(fg)
    return p

def _trennlinie(cell, color_hex, oben=False):
    """Absatz mit Unter- oder Oberrand als Trennlinie"""
    sep = cell.add_paragraph()
    sep.paragraph_format.space_before = Pt(3)
    sep.paragraph_format.space_after  = Pt(1)
    pPr = sep._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    edge = OxmlElement('w:top' if oben else 'w:bottom')
    edge.set(qn('w:val'), 'single'); edge.set(qn('w:sz'), '6')
    edge.set(qn('w:space'), '1'); edge.set(qn('w:color'), color_hex.upper())
    pBdr.append(edge); pPr.append(pBdr)

def _section_hdr(container, text, bg_hex, fg="FFFFFF", width=Cm(12.6), sa=0, sb=0):
    """Einfarbige Header-Leiste als 1×1 Tabelle"""
    t = container.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.cell(0, 0); _no_border(c); _set_bg(c, bg_hex); c.width = width
    _p(c, text, bold=True, size=8.5, fg=fg, sa=sa, sb=sb)
    return t

def _namen_tabelle(container, lst, width_name=Cm(4.0), width_zeit=Cm(2.3),
                   bg_a="F4F4F4", bg_b="FFFFFF", col_count=2):
    """
    Gibt Personen (name, zeit) in einer 2-spaltigen Namensliste aus.
    col_count=2: 2er-Spalten-Modus (zwei Namen nebeneinander)
    col_count=1: schlichte Einzelliste
    """
    if not lst: return
    if col_count == 2:
        half = (len(lst) + 1) // 2
        tbl = container.add_table(rows=half, cols=4); tbl.style = 'Table Grid'
        for i, (name, zeit) in enumerate(lst):
            row = i % half; col = (i // half) * 2
            nc = tbl.cell(row, col); zc = tbl.cell(row, col + 1)
            _no_border(nc); _no_border(zc)
            bg = bg_a if row % 2 == 0 else bg_b
            _set_bg(nc, bg); _set_bg(zc, bg)
            nc.width = width_name; zc.width = width_zeit
            pn = nc.paragraphs[0]
            pn.paragraph_format.space_before = Pt(1)
            pn.paragraph_format.space_after  = Pt(1)
            rn = pn.add_run(name)
            rn.font.size = Pt(8); rn.font.color.rgb = _rgb("111111")
            pz = zc.paragraphs[0]
            pz.paragraph_format.space_before = Pt(1)
            pz.paragraph_format.space_after  = Pt(1)
            rz = pz.add_run(zeit)
            rz.font.size = Pt(7.5); rz.font.color.rgb = _rgb("777777"); rz.bold = True
    else:
        for i, (name, zeit) in enumerate(lst):
            dt2 = container.add_table(rows=1, cols=2); dt2.style = 'Table Grid'
            nc = dt2.cell(0, 0); zc = dt2.cell(0, 1)
            _no_border(nc); _no_border(zc)
            bg = bg_a if i % 2 == 0 else bg_b
            _set_bg(nc, bg); _set_bg(zc, bg)
            nc.width = Cm(9.6); zc.width = Cm(3.0)
            pn = nc.paragraphs[0]
            pn.paragraph_format.space_before = Pt(1); pn.paragraph_format.space_after = Pt(1)
            rn = pn.add_run(name); rn.font.size = Pt(8.5)
            rn.font.color.rgb = _rgb("111111"); rn.bold = True
            pz = zc.paragraphs[0]; pz.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pz.paragraph_format.space_before = Pt(1); pz.paragraph_format.space_after = Pt(1)
            rz = pz.add_run(zeit); rz.font.size = Pt(8)
            rz.font.color.rgb = _rgb(AZ); rz.bold = True


# ═══════════════════════════════════════════════════════════════════════════════
# F1 – HAUPTFUNKTION
# ═══════════════════════════════════════════════════════════════════════════════
def erstelle_F1(data, bul=5, einz=28, pat=5, pax=42500,
                schichtleiter_name="", schichtleiter_zeit=""):

    # Status-Werte
    if bul <= 2:
        fc_hex, lbl_bul, fc_bg_h = "FF3333", "KRITISCH",       "3A0000"
    elif bul == 3:
        fc_hex, lbl_bul, fc_bg_h = "E07800", "EINGESCHRÄNKT",  "3A2000"
    else:
        fc_hex, lbl_bul, fc_bg_h = "10A050", "VOLLSTÄNDIG",    "003A18"

    betreuer_sch = _nach_schicht(data, 'betreuer')
    dispo_sch    = _nach_schicht(data, 'dispo')
    kranke_lst   = _kranke(data)
    bt_dt = betreuer_sch['DT']; bt_dn = betreuer_sch['DN']
    di_dt = dispo_sch['DT'];    di_dn = dispo_sch['DN']
    gesamt = len(bt_dt) + len(bt_dn) + len(di_dt) + len(di_dn)

    # Schichtleiter ermitteln
    if not schichtleiter_name:
        schichtleiter_name, schichtleiter_zeit = _schichtleiter(data)

    # ── Dokument ──────────────────────────────────────────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.page_width   = Cm(21.0); sec.page_height = Cm(29.7)
        sec.top_margin   = sec.bottom_margin = Cm(0.7)
        sec.left_margin  = sec.right_margin  = Cm(0.5)

    LEFT_W  = Cm(7.4)
    RIGHT_W = Cm(12.6)

    # Haupt-Tabelle: 2 Spalten
    main = doc.add_table(rows=1, cols=2)
    main.style = 'Table Grid'
    lc = main.cell(0, 0); rc = main.cell(0, 1)
    lc.width = LEFT_W; rc.width = RIGHT_W
    _no_border(lc); _no_border(rc)
    _set_bg(lc, BG); _set_bg(rc, "FFFFFF")
    lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ══════════════════════════════════════════════════════════════════════════
    # LINKE SPALTE
    # ══════════════════════════════════════════════════════════════════════════

    # Logo
    try:
        if LOGO.exists():
            lp = lc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp.paragraph_format.space_before = Pt(4)
            lp.add_run().add_picture(str(LOGO), width=Cm(2.8))
    except: pass

    _p(lc, "Deutsches Rotes Kreuz",  bold=True, size=10, fg="FFFFFF", align="center", sb=2)
    _p(lc, "Kreisverband Köln e.V.", size=7.5,  fg=HE,    align="center")
    # Telefonnummer direkt unter der Station-Zeile (kein Uhrzeit!)
    _p(lc, STATION,                  size=7,    fg="AAAAAA", align="center")
    _p(lc, TEL,                      bold=True,  size=8,  fg=HE,    align="center", sa=2)

    _trennlinie(lc, AZ)
    _p(lc, f"📅  {DATUM}", bold=True, size=10, fg="FFFFFF", align="center", sb=2, sa=4)

    # Kennzahlen
    for lbl, val, vc in [
        ("✦ Einsätze",  str(einz),                      AZ),
        ("✦ Patienten", str(pat),                        "AAFFCC"),
        ("✦ PAX",       f"{pax:,}".replace(",", "."),    HE),
        ("✦ Personal",  str(gesamt),                     HE),
    ]:
        p = lc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(2)
        r1 = p.add_run(f"{lbl}  "); r1.font.size = Pt(7); r1.font.color.rgb = _rgb(HE)
        r2 = p.add_run(val); r2.bold = True; r2.font.size = Pt(15); r2.font.color.rgb = _rgb(vc)

    # Bulmor-Status (OHNE Fahrerliste)
    _trennlinie(lc, AZ, oben=True)
    _p(lc, "BULMOR – FAHRZEUGSTATUS", bold=True, size=8, fg=AZ, align="center", sb=2)

    # Farbige Status-Box
    st = lc.add_table(rows=1, cols=1); st.style = 'Table Grid'
    stc = st.cell(0, 0); _no_border(stc); _set_bg(stc, fc_bg_h); stc.width = LEFT_W
    p_st = stc.paragraphs[0]; p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_st = p_st.add_run(f"{bul}/5 Bulmor  ·  {lbl_bul}")
    r_st.bold = True; r_st.font.size = Pt(10); r_st.font.color.rgb = _rgb(fc_hex)

    # Bulmor-Symbole (5 Kreise mit ✓/✕)
    p_sym = lc.add_paragraph()
    p_sym.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sym.paragraph_format.space_before = Pt(3)
    p_sym.paragraph_format.space_after  = Pt(2)
    for i in range(5):
        aktiv = (i + 1) <= bul
        rb = p_sym.add_run(f"B{i+1}")
        rb.bold = True; rb.font.size = Pt(8)
        rb.font.color.rgb = _rgb(fc_hex if aktiv else "777777")
        rc_sym = p_sym.add_run("✓ " if aktiv else "✕ ")
        rc_sym.font.size = Pt(8)
        rc_sym.font.color.rgb = _rgb(fc_hex if aktiv else "FF4444")

    # Kranke
    if kranke_lst:
        _trennlinie(lc, "FF4444", oben=True)
        _p(lc, "KRANKMELDUNG", bold=True, size=7.5, fg="FF4444", align="center", sb=2)
        for kname, kzeit in kranke_lst:
            _p(lc, f"{kname}  ({kzeit})", size=7.5, fg="FF8888", align="center")

    # ══════════════════════════════════════════════════════════════════════════
    # RECHTE SPALTE
    # ══════════════════════════════════════════════════════════════════════════

    # ── Org-Header (Station + Telefon + Datum) ────────────────────────────────
    hdr = rc.add_table(rows=1, cols=1); hdr.style = 'Table Grid'
    hc = hdr.cell(0, 0); _no_border(hc); _set_bg(hc, DK); hc.width = RIGHT_W
    _p(hc, STATION,        bold=True, size=9,   fg=HE,      sb=2)
    _p(hc, TEL + "  ·  " + MAIL, size=7.5, fg="AAAAAA")
    _p(hc, DATUM,          bold=True, size=9,   fg=AZ,      sa=2)

    # ── Tag / Nacht Legende ───────────────────────────────────────────────────
    leg = rc.add_table(rows=1, cols=3); leg.style = 'Table Grid'
    lc1 = leg.cell(0, 0); lc2 = leg.cell(0, 1); lc3 = leg.cell(0, 2)
    for c_, bg_, lbl_ in [(lc1, AZ, "■ TAGDIENST"), (lc2, AZ2, "■ NACHTDIENST"),
                           (lc3, "333344", f"■ GESAMT: {gesamt} Personen")]:
        _no_border(c_); _set_bg(c_, bg_); c_.width = Cm(4.2)
        _p(c_, lbl_, bold=True, size=7, fg="FFFFFF", align="center")
    sp = rc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

    # ── Betreuer Tagdienst ────────────────────────────────────────────────────
    _section_hdr(rc, f"BETREUER – TAGDIENST  ({len(bt_dt)} Personen)", DK, width=RIGHT_W)
    _namen_tabelle(rc, bt_dt, col_count=2)

    # ── Betreuer Nachtdienst ──────────────────────────────────────────────────
    if bt_dn:
        _section_hdr(rc, f"BETREUER – NACHTDIENST  ({len(bt_dn)} Personen)", "2A2A4A", width=RIGHT_W)
        _namen_tabelle(rc, bt_dn, bg_a="EDEDFF", col_count=2)

    # ── Disposition ───────────────────────────────────────────────────────────
    _section_hdr(rc, f"DISPOSITION  ({len(di_dt)+len(di_dn)} Personen)", DK, width=RIGHT_W)

    if di_dt:
        # Sub-Header Tagdienst
        sh1 = rc.add_table(rows=1, cols=1); sh1.style = 'Table Grid'
        sh1c = sh1.cell(0, 0); _no_border(sh1c); _set_bg(sh1c, AZ); sh1c.width = RIGHT_W
        _p(sh1c, "Tagdienst", bold=True, size=7.5, fg="FFFFFF", sb=1, sa=1)
        _namen_tabelle(rc, di_dt, col_count=1)

    if di_dn:
        sh2 = rc.add_table(rows=1, cols=1); sh2.style = 'Table Grid'
        sh2c = sh2.cell(0, 0); _no_border(sh2c); _set_bg(sh2c, AZ2); sh2c.width = RIGHT_W
        _p(sh2c, "Nachtdienst", bold=True, size=7.5, fg="FFFFFF", sb=1, sa=1)
        _namen_tabelle(rc, di_dn, bg_a="EDEDFF", col_count=1)

    # ── Schichtleiter ─────────────────────────────────────────────────────────
    _section_hdr(rc, "SCHICHTLEITER", "143060", width=RIGHT_W)

    sl_tbl = rc.add_table(rows=1, cols=2); sl_tbl.style = 'Table Grid'
    sl_nc = sl_tbl.cell(0, 0); sl_zc = sl_tbl.cell(0, 1)
    _no_border(sl_nc); _no_border(sl_zc)
    _set_bg(sl_nc, "EEF6FF"); _set_bg(sl_zc, "EEF6FF")
    sl_nc.width = Cm(9.6); sl_zc.width = Cm(3.0)

    pn = sl_nc.paragraphs[0]
    pn.paragraph_format.space_before = Pt(2); pn.paragraph_format.space_after = Pt(2)
    rn = pn.add_run(schichtleiter_name if schichtleiter_name else "—")
    rn.bold = True; rn.font.size = Pt(10); rn.font.color.rgb = _rgb("1A3460")

    pz = sl_zc.paragraphs[0]; pz.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pz.paragraph_format.space_before = Pt(2); pz.paragraph_format.space_after = Pt(2)
    rz = pz.add_run(schichtleiter_zeit if schichtleiter_zeit else "")
    rz.bold = True; rz.font.size = Pt(9); rz.font.color.rgb = _rgb(AZ)

    # ── Footer rechts ─────────────────────────────────────────────────────────
    sp2 = rc.add_paragraph(); sp2.paragraph_format.space_before = Pt(6)
    ft = rc.add_table(rows=1, cols=1); ft.style = 'Table Grid'
    ftc = ft.cell(0, 0); _no_border(ftc); _set_bg(ftc, DK); ftc.width = RIGHT_W
    _p(ftc, f"DRK Köln  ·  {TEL}  ·  {MAIL}",
       size=7.5, fg="FFFFFF", align="center", sb=3, sa=3)

    # ── Speichern ─────────────────────────────────────────────────────────────
    out = os.path.join(ZIEL, "F1_25032026.docx")
    doc.save(out)
    print(f"[OK]  F1_25032026.docx  →  {out}")


# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Lade Dienstplan-Daten ...")
    data = lade()
    bs = _nach_schicht(data, 'betreuer'); ds = _nach_schicht(data, 'dispo')
    sl_name, sl_zeit = _schichtleiter(data)
    print(f"  Betreuer  DT:{len(bs['DT'])}  DN:{len(bs['DN'])}")
    print(f"  Dispo     DT:{len(ds['DT'])}  DN:{len(ds['DN'])}")
    print(f"  Schichtleiter: {sl_name}  {sl_zeit}")
    print(f"  Krank:    {len(_kranke(data))}")
    print(f"  Bulmor verfügbar: {_bulmor_anzahl(data)}")
    print()

    erstelle_F1(
        data,
        bul   = _bulmor_anzahl(data),   # echte Anzahl aus Daten
        einz  = 28,
        pat   = 5,
        pax   = 42500,
        schichtleiter_name = sl_name,
        schichtleiter_zeit = sl_zeit,
    )
    print(f"\nGespeichert in: {ZIEL}")
