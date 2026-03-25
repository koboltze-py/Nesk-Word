# -*- coding: utf-8 -*-
"""
P9-Split-Stil v2 – Tag/Nacht-Trennung, Arbeitszeiten je Person, Bulmor-Statusbeispiele
"""
import os, sys
from pathlib import Path
from collections import defaultdict

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor, white, black
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W, H = A4

# ── Pfade ─────────────────────────────────────────────────────────────────────
_OD   = r"C:\Users\DRKairport\OneDrive - Deutsches Rotes Kreuz - Kreisverband Köln e.V"
ZIEL  = os.path.join(_OD, "Desktop", "bei") if os.path.exists(_OD) else r"C:\Temp\bei"
os.makedirs(ZIEL, exist_ok=True)
LOGO  = Path(os.path.dirname(os.path.abspath(__file__))) / "Daten" / "Email" / "Logo.jpg"
EXCEL = (_OD + r"\Dateien von Erste-Hilfe-Station-Flughafen - DRK Köln e.V_ - !Gemeinsam.26"
         r"\04_Tagesdienstpläne\03_März\25.03.2026.xlsx")
DATUM   = "25.03.2026"
UHRZEIT = "07:45 Uhr"
STATION = "Erste-Hilfe-Station · Flughafen Köln/Bonn"

def h(hex_): return HexColor(f"#{hex_.lstrip('#')}")
ROT_WARN = h("FF3333"); GRN_OK = h("10A050"); ORG_WARN = h("E07800")

def _bul_col(n):
    if n <= 2:   return h("FF3333"), "KRITISCH",      h("3A0000"), h("FF8888")
    elif n == 3: return h("E07800"), "EINGESCHRÄNKT", h("3A2000"), h("FFD09A")
    else:        return h("10A050"), "VOLLSTÄNDIG",   h("003A18"), h("A0FFC8")

# ── Daten ─────────────────────────────────────────────────────────────────────
def lade():
    try:
        from functions.dienstplan_parser import DienstplanParser
        r = DienstplanParser(EXCEL, alle_anzeigen=True).parse()
        if r.get("success") in (True, "True"): return r
    except Exception as e:
        print(f"  [Parser] {e}")
    return {"betreuer": [], "dispo": [], "kranke": []}

def _ist_tag(p):
    """Tagdienst = Beginn vor 14:00"""
    s = (p.get('start_zeit') or '00:00')[:5]
    try: return int(s.split(':')[0]) < 14
    except: return True

def _zeit(p):
    s = (p.get('start_zeit') or '')[:5]
    e = (p.get('end_zeit')   or '')[:5]
    return f"{s}–{e}" if s and e else ''

def _nach_schicht(d, typ='betreuer'):
    """Gibt {'DT': [(name, zeit)], 'DN': [(name, zeit)]} zurück"""
    dt, dn = [], []
    for p in d.get(typ, []):
        if p.get('ist_krank') in (True, 'True'): continue
        name = p.get('anzeigename', '').strip()
        if not name: continue
        z = _zeit(p)
        if _ist_tag(p): dt.append((name, z))
        else:           dn.append((name, z))
    dt.sort(key=lambda x: x[0])
    dn.sort(key=lambda x: x[0])
    return {'DT': dt, 'DN': dn}

def _kranke(d):
    out = []
    for p in d.get('betreuer', []) + d.get('dispo', []):
        if p.get('ist_krank') in (True, 'True') and p.get('anzeigename', '').strip():
            out.append((p.get('anzeigename', '').strip(), _zeit(p)))
    return out

def _bulfhr(d):
    return [p.get('anzeigename', '').strip()
            for p in d.get('betreuer', []) + d.get('dispo', [])
            if p.get('ist_bulmorfahrer') in (True, 'True') and p.get('anzeigename', '').strip()]

# ── PDF-Primitiven ─────────────────────────────────────────────────────────────
def rect(cv, x, yt, w_, ht, color):
    cv.setFillColor(color); cv.rect(x, H-yt-ht, w_, ht, fill=1, stroke=0)

def ln(cv, x1, y1, x2, y2, color, lw=0.8):
    cv.setStrokeColor(color); cv.setLineWidth(lw); cv.line(x1, H-y1, x2, H-y2)

def t(cv, x, yt, text, font="Helvetica", size=9, color=black, align="left"):
    cv.setFont(font, size); cv.setFillColor(color)
    if align == "center": cv.drawCentredString(x, H-yt, str(text))
    elif align == "right": cv.drawRightString(x, H-yt, str(text))
    else: cv.drawString(x, H-yt, str(text))

def logo_draw(cv, x, yt, w_=55, ht=52):
    if LOGO.exists():
        try: cv.drawImage(ImageReader(str(LOGO)), x, H-yt-ht, w_, ht,
                          mask='auto', preserveAspectRatio=True)
        except: pass

def kreis(cv, cx, cy_top, r, fill, stroke_c=None, lw=1.2):
    cv.setFillColor(fill)
    if stroke_c:
        cv.setStrokeColor(stroke_c); cv.setLineWidth(lw)
        cv.circle(cx, H-cy_top, r, fill=1, stroke=1)
    else:
        cv.circle(cx, H-cy_top, r, fill=1, stroke=0)

def section_hdr(cv, x, yt, w_, label, bg, fg=white, tag_label="", tag_bg=None, size=8.5):
    rect(cv, x, yt, w_, 17, bg)
    t(cv, x+6, yt+12, label, "Helvetica-Bold", size, fg)
    if tag_label and tag_bg:
        tw = len(tag_label)*6 + 8
        rect(cv, x+w_-tw-4, yt+2, tw, 13, tag_bg)
        t(cv, x+w_-5, yt+12, tag_label, "Helvetica-Bold", 7, white, "right")


# ═══════════════════════════════════════════════════════════════════════════════
# KERN: Zeichne Split-Layout MIT Tag/Nacht + Arbeitszeiten
# ═══════════════════════════════════════════════════════════════════════════════
def draw_split_v2(cv, BG, BG_D, HELL, AKZ, AKZ2,
                  data, bul, einz, pat, pax):
    """
    BG      = linkes Panel (Hauptton)
    BG_D    = dunklerer Ton (Kacheln, Section-Header rechts)
    HELL    = heller Ton (Subtitle-Text links)
    AKZ     = Akzentfarbe (Trennlinien, DT-Badge)
    AKZ2    = zweiter Akzent (Überschriften rechts, DN-Badge)
    """
    fc, lbl_bul, fc_bg, fc_txt = _bul_col(bul)
    betreuer_sch = _nach_schicht(data, 'betreuer')
    dispo_sch    = _nach_schicht(data, 'dispo')
    kranke_lst   = _kranke(data)
    bulfhr_lst   = _bulfhr(data)
    bt_dt = betreuer_sch['DT']; bt_dn = betreuer_sch['DN']
    di_dt = dispo_sch['DT'];    di_dn = dispo_sch['DN']
    gesamt = len(bt_dt) + len(bt_dn) + len(di_dt) + len(di_dn)

    WEISS = white; DUNKEL = h("111111"); GRAU = h("555555")
    LW = W * 0.41  # Breite linkes Panel

    rect(cv, 0, 0, W, H, WEISS)
    rect(cv, 0, 0, LW, H, BG)
    rect(cv, LW-3, 0, 3, H, BG_D)

    # ── LINKES PANEL ──────────────────────────────────────────────────────────
    logo_draw(cv, 14, 12, w_=52, ht=50)

    t(cv, LW/2, 74, "Stärkemeldung", "Helvetica-Bold", 17, WEISS, "center")
    t(cv, LW/2, 91, "und Einsätze", "Helvetica-Bold", 17, AKZ, "center")
    ln(cv, 14, 98, LW-14, 98, WEISS, 1)
    t(cv, LW/2, 111, DATUM, "Helvetica-Bold", 10, HELL, "center")
    t(cv, LW/2, 124, UHRZEIT, "Helvetica", 8.5, h("AAAAAA"), "center")

    # Kennzahlen-Kacheln
    kz = [
        ("Einsätze",  str(einz),                             AKZ),
        ("Patienten", str(pat),                              h("AAFFCC")),
        ("PAX",       f"{pax:,}".replace(",", "."),          HELL),
        ("Personal",  str(gesamt),                           HELL),
    ]
    for i, (lbl, val, vc) in enumerate(kz):
        ky = 133 + i * 48
        rect(cv, 12, ky, LW-24, 41, BG_D)
        ln(cv, 12, ky, 16, ky, AKZ, 4)
        t(cv, 22, ky+12, lbl, "Helvetica", 7, HELL)
        t(cv, 22, ky+36, val, "Helvetica-Bold", 19, vc)

    # ── BULMOR-STATUS ──────────────────────────────────────────────────────────
    by = 133 + 4*48 + 8
    ln(cv, 14, by, LW-14, by, AKZ, 0.8)
    t(cv, LW/2, by+13, "BULMOR – FAHRZEUGSTATUS", "Helvetica-Bold", 8.5, WEISS, "center")

    # Hintergrund für Status-Box (farbig je Status!)
    rect(cv, 12, by+17, LW-24, 18, fc_bg)
    t(cv, LW/2, by+30, f"{bul}/5  ·  {lbl_bul}", "Helvetica-Bold", 9.5, fc, "center")

    # 5 Bulmor-Kreise
    step = (LW-40)/4
    for i in range(5):
        aktiv = (i+1) <= bul
        cx_ = 20 + i*step
        kreis(cv, cx_, by+54, 14, fc if aktiv else BG_D, WEISS if aktiv else h("555555"), 1)
        t(cv, cx_, by+58, f"B{i+1}", "Helvetica-Bold", 7,
          WEISS if aktiv else h("888888"), "center")
        # X wenn nicht aktiv
        if not aktiv:
            t(cv, cx_, by+48, "✕", "Helvetica-Bold", 9, h("FF4444"), "center")
        else:
            t(cv, cx_, by+48, "✓", "Helvetica-Bold", 9, GRN_OK, "center")

    # Bulmor-Fahrer (kompakt)
    if bulfhr_lst:
        fy = by+72
        t(cv, LW/2, fy, "Fahrer:", "Helvetica-Bold", 7, HELL, "center")
        for j, fn in enumerate(bulfhr_lst[:5]):
            t(cv, LW/2, fy+10+j*11, fn, "Helvetica", 7.5, h("CCCCCC"), "center")
        fy += 10 + len(bulfhr_lst[:5])*11
    else:
        fy = by+76

    # Kranke (links unten)
    if kranke_lst:
        ln(cv, 14, fy+6, LW-14, fy+6, h("FF4444"), 0.7)
        t(cv, LW/2, fy+18, "KRANKMELDUNG", "Helvetica-Bold", 7.5, ROT_WARN, "center")
        for j, (kname, kzeit) in enumerate(kranke_lst[:4]):
            t(cv, LW/2, fy+28+j*11, f"{kname}  ({kzeit})", "Helvetica", 7, h("FFAAAA"), "center")

    # Footer links
    t(cv, LW/2, H-8, "+49 2203 40-2323", "Helvetica", 6.5, h("888888"), "center")

    # ── RECHTES PANEL ─────────────────────────────────────────────────────────
    RX = LW + 10
    RW = W - RX - 10

    # Org-Header
    rect(cv, LW, 0, W-LW, 62, BG_D)
    ln(cv, LW, 62, W, 62, AKZ, 2)
    t(cv, RX, 18, STATION, "Helvetica-Bold", 9, HELL)
    t(cv, RX, 31, "DRK Kreisverband Köln e.V.", "Helvetica", 8, h("AAAAAA"))
    t(cv, W-12, 18, DATUM,   "Helvetica-Bold", 10, AKZ, "right")
    t(cv, W-12, 31, UHRZEIT, "Helvetica", 8.5, HELL, "right")
    # Tag/Nacht-Legende
    rect(cv, RX, 40, 50, 14, AKZ);  t(cv, RX+25, 50, "TAGDIENST",  "Helvetica-Bold", 6.5, white, "center")
    rect(cv, RX+56, 40, 60, 14, AKZ2); t(cv, RX+86, 50, "NACHTDIENST","Helvetica-Bold", 6.5, white, "center")

    cy = 70

    # ── BETREUER: TAGDIENST ───────────────────────────────────────────────────
    section_hdr(cv, LW, cy, W-LW, f"BETREUER – TAGDIENST  ({len(bt_dt)} Pers.)", BG_D, tag_label="DT", tag_bg=AKZ)
    cy += 19
    half = (len(bt_dt)+1)//2
    for i, (name, zeit) in enumerate(bt_dt):
        col = i // half; row = i % half
        bx = LW + col*(RW/2); by2 = cy + row*13
        if by2 > H-100: break
        bg = h("F4F4F4") if row%2==0 else WEISS
        rect(cv, bx, by2, RW/2-1, 12, bg)
        t(cv, bx+5, by2+9, name, "Helvetica", 7.5, DUNKEL)
        if zeit:
            t(cv, bx+RW/2-6, by2+9, zeit, "Helvetica", 6.5, h("888888"), "right")
    cy += max(half, 1)*13 + 4

    # ── BETREUER: NACHTDIENST ─────────────────────────────────────────────────
    if bt_dn and cy < H-80:
        section_hdr(cv, LW, cy, W-LW, f"BETREUER – NACHTDIENST  ({len(bt_dn)} Pers.)", h("2A2A4A"), tag_label="DN", tag_bg=AKZ2)
        cy += 19
        nhalf = (len(bt_dn)+1)//2
        for i, (name, zeit) in enumerate(bt_dn):
            col = i // nhalf; row = i % nhalf
            bx = LW + col*(RW/2); by3 = cy + row*13
            if by3 > H-60: break
            bg = h("F0F0FF") if row%2==0 else WEISS
            rect(cv, bx, by3, RW/2-1, 12, bg)
            t(cv, bx+5, by3+9, name, "Helvetica", 7.5, DUNKEL)
            if zeit:
                t(cv, bx+RW/2-6, by3+9, zeit, "Helvetica", 6.5, h("888888"), "right")
        cy += max(nhalf, 1)*13 + 4

    # ── DISPOSITION ───────────────────────────────────────────────────────────
    if cy < H-60:
        # DT + DN in einer Tabelle (Seite-an-Seite)
        section_hdr(cv, LW, cy, W-LW,
                    f"DISPOSITION  ({len(di_dt)+len(di_dn)} Pers.)", BG_D)
        cy += 19
        # Tagdienst
        if di_dt:
            t(cv, RX, cy+8, "Tagdienst", "Helvetica-Bold", 7.5, AKZ)
            for i, (name, zeit) in enumerate(di_dt):
                row_y = cy+10+i*13
                if row_y > H-24: break
                bg = h("F4F4F4") if i%2==0 else WEISS
                rect(cv, LW, row_y, W-LW, 12, bg)
                t(cv, RX+14, row_y+9, name, "Helvetica", 8, DUNKEL)
                if zeit: t(cv, W-14, row_y+9, zeit, "Helvetica-Bold", 7.5, AKZ, "right")
            cy += 10 + len(di_dt)*13 + 2
        # Nachtdienst
        if di_dn and cy < H-30:
            t(cv, RX, cy+8, "Nachtdienst", "Helvetica-Bold", 7.5, AKZ2)
            for i, (name, zeit) in enumerate(di_dn):
                row_y = cy+10+i*13
                if row_y > H-24: break
                bg = h("F0F0FF") if i%2==0 else WEISS
                rect(cv, LW, row_y, W-LW, 12, bg)
                t(cv, RX+14, row_y+9, name, "Helvetica", 8, DUNKEL)
                if zeit: t(cv, W-14, row_y+9, zeit, "Helvetica-Bold", 7.5, AKZ2, "right")
            cy += 10 + len(di_dn)*13

    # ── FOOTER ────────────────────────────────────────────────────────────────
    rect(cv, 0, H-18, W, 18, BG_D)
    ln(cv, 0, H-18, W, H-18, AKZ, 1.5)
    t(cv, W/2, H-6, "DRK Köln  ·  +49 2203 40-2323  ·  flughafen@drk-koeln.de",
      "Helvetica", 7.5, WEISS, "center")


# ═══════════════════════════════════════════════════════════════════════════════
# FARBSCHEMATA (bg, dunkel, hell, akz_tag, akz_nacht)
# ═══════════════════════════════════════════════════════════════════════════════
SCHEMES = {
    "Smaragd":   ("1B6B45", "0F4D32", "D4F5E2", "F5C518", "0097A7"),
    "Burgund":   ("7B1A3A", "590F28", "FFD6E0", "FFB347", "A86AA0"),
    "OzeanBlau": ("1A3460", "0F1F3C", "C8DAFF", "00C8FF", "00A878"),
    "Schiefer":  ("2D4A6B", "1A2D45", "D0E8FF", "E8603A", "5B9BD5"),
    "Violett":   ("4A1E7E", "32145A", "EDD6FF", "00DDB8", "E040FB"),
    "Anthrazit": ("2C3440", "1A2028", "D8E8F0", "A8D040", "40C4FF"),
}


# ═══════════════════════════════════════════════════════════════════════════════
# PDF-GENERIERUNG: 6 Farbvarianten (Vollbesetzung)
# ═══════════════════════════════════════════════════════════════════════════════
def erstelle_farb_pdfs(data, bul=5, einz=28, pat=5, pax=42500):
    for sname, (bg, dk, he, az, az2) in SCHEMES.items():
        out = os.path.join(ZIEL, f"V2_{sname}_25032026.pdf")
        cv_ = canvas.Canvas(out, pagesize=A4)
        draw_split_v2(cv_, h(bg), h(dk), h(he), h(az), h(az2),
                      data, bul=bul, einz=einz, pat=pat, pax=pax)
        cv_.save()
        print(f"[OK] V2_{sname}_25032026.pdf")


# ═══════════════════════════════════════════════════════════════════════════════
# PDF-GENERIERUNG: Bulmor-Status-Beispiele (bul=1…5, immer Ozeanblau)
# ═══════════════════════════════════════════════════════════════════════════════
BUL_BEISPIELE = [
    (1, "Bul1_KRITISCH_25032026.pdf",      "Nur 1 Bulmor – Ausfall"),
    (2, "Bul2_KRITISCH_25032026.pdf",      "2 Bulmor – KRITISCH"),
    (3, "Bul3_EINGESCHRAENKT_25032026.pdf","3 Bulmor – EINGESCHRÄNKT"),
    (4, "Bul4_VOLLSTAENDIG_25032026.pdf",  "4 Bulmor – ausreichend"),
    (5, "Bul5_VOLLSTAENDIG_25032026.pdf",  "5 Bulmor – Vollbesetzung"),
]

def erstelle_bulmor_pdfs(data, einz=28, pat=5, pax=42500):
    # Ozeanblau als Basis
    bg, dk, he, az, az2 = SCHEMES["OzeanBlau"]
    print("\nBulmor-Status-Beispiele:")
    for bul, fname, desc in BUL_BEISPIELE:
        out = os.path.join(ZIEL, fname)
        cv_ = canvas.Canvas(out, pagesize=A4)
        draw_split_v2(cv_, h(bg), h(dk), h(he), h(az), h(az2),
                      data, bul=bul, einz=einz, pat=pat, pax=pax)
        fc_, lbl_, _, _ = _bul_col(bul)
        cv_.save()
        print(f"  [OK] {fname}  ({desc})")


# ═══════════════════════════════════════════════════════════════════════════════
# WORD-VERSION mit Tag/Nacht + Arbeitszeiten
# ═══════════════════════════════════════════════════════════════════════════════
def _rgb(hex_str):
    hx = hex_str.lstrip("#")
    return RGBColor(int(hx[0:2],16), int(hx[2:4],16), int(hx[4:6],16))

def _set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    for k, v in [('w:val','clear'),('w:color','auto'),('w:fill', hex_color.lstrip('#').upper())]:
        shd.set(qn(k), v)
    tcPr.append(shd)

def _no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcB.append(b)
    tcPr.append(tcB)

def _par(cell, text, bold=False, size=9, color="000000", align="left",
         sa=0, sb=0):
    p = cell.add_paragraph()
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = _rgb(color)
    return p

def _add_hdr_row(tbl, cols, texts_bolds_sizes_fgs_bgs, widths):
    """Fügt eine Header-Zeile in eine Tabelle ein"""
    row = tbl.add_row()
    for i, (txt, bold, sz, fg, bg) in enumerate(texts_bolds_sizes_fgs_bgs):
        c = row.cells[i]; _no_border(c); _set_cell_bg(c, bg)
        c.width = Cm(widths[i])
        p = c.paragraphs[0]; r = p.add_run(str(txt))
        r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = _rgb(fg)
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    return row

def erstelle_word_v2(data, scheme_name="OzeanBlau", bul=5, einz=28, pat=5, pax=42500):
    bg, dk, he, az, az2 = SCHEMES[scheme_name]
    if bul <= 2:
        fc_hex, lbl_bul, fc_bg_h, fc_txt_h = "FF3333", "KRITISCH",       "3A0000", "FF8888"
    elif bul == 3:
        fc_hex, lbl_bul, fc_bg_h, fc_txt_h = "E07800", "EINGESCHRÄNKT",  "3A2000", "FFD09A"
    else:
        fc_hex, lbl_bul, fc_bg_h, fc_txt_h = "10A050", "VOLLSTÄNDIG",    "003A18", "A0FFC8"

    betreuer_sch = _nach_schicht(data, 'betreuer')
    dispo_sch    = _nach_schicht(data, 'dispo')
    kranke_lst   = _kranke(data)
    bulfhr_lst   = _bulfhr(data)
    bt_dt = betreuer_sch['DT']; bt_dn = betreuer_sch['DN']
    di_dt = dispo_sch['DT'];    di_dn = dispo_sch['DN']
    gesamt = len(bt_dt)+len(bt_dn)+len(di_dt)+len(di_dn)

    doc = Document()
    for sec in doc.sections:
        sec.page_width  = Cm(21.0); sec.page_height = Cm(29.7)
        sec.top_margin  = sec.bottom_margin = Cm(0.7)
        sec.left_margin = sec.right_margin  = Cm(0.5)

    # Haupt-2-Spalten-Tabelle
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    LEFT_W = Cm(7.4); RIGHT_W = Cm(12.6)
    lc = tbl.cell(0, 0); rc = tbl.cell(0, 1)
    lc.width = LEFT_W; rc.width = RIGHT_W
    _no_border(lc); _no_border(rc)
    _set_cell_bg(lc, bg); _set_cell_bg(rc, "FFFFFF")
    lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ── LINKE ZELLE ──────────────────────────────────────────────────────────
    try:
        if LOGO.exists():
            lp = lc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp.paragraph_format.space_before = Pt(4)
            lp.add_run().add_picture(str(LOGO), width=Cm(2.8))
    except: pass

    _par(lc, "Deutsches Rotes Kreuz", bold=True, size=10, color="FFFFFF", align="center", sb=2)
    _par(lc, "Kreisverband Köln e.V.", size=7.5, color=he, align="center")
    _par(lc, STATION, size=7, color="AAAAAA", align="center", sa=3)

    # Trennlinie
    sep = lc.add_paragraph(); sep.paragraph_format.space_after = Pt(1)
    pPr = sep._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), az.upper())
    pBdr.append(bot); pPr.append(pBdr)

    _par(lc, f"📅  {DATUM}", bold=True, size=10, color="FFFFFF", align="center", sb=2)
    _par(lc, f"🕖  {UHRZEIT}", size=8.5, color=he, align="center", sa=4)

    # Kennzahlen
    for lbl, val, vc in [("✦ Einsätze", str(einz), az), ("✦ Patienten", str(pat), "AAFFCC"),
                          ("✦ PAX", f"{pax:,}".replace(",","."), he), ("✦ Personal", str(gesamt), he)]:
        p = lc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(2)
        r1 = p.add_run(f"{lbl}  "); r1.font.size = Pt(7); r1.font.color.rgb = _rgb(he)
        r2 = p.add_run(val); r2.bold = True; r2.font.size = Pt(15); r2.font.color.rgb = _rgb(vc)

    # Bulmor
    sep2 = lc.add_paragraph(); sep2.paragraph_format.space_before = Pt(4)
    pPr2 = sep2._p.get_or_add_pPr(); pBdr2 = OxmlElement('w:pBdr')
    t2 = OxmlElement('w:top')
    t2.set(qn('w:val'),'single'); t2.set(qn('w:sz'),'4')
    t2.set(qn('w:color'),az.upper()); t2.set(qn('w:space'),'1')
    pBdr2.append(t2); pPr2.append(pBdr2)

    _par(lc, "BULMOR – FAHRZEUGSTATUS", bold=True, size=8, color=az, align="center", sb=2)

    # Status-Box farbig
    st_tbl = lc.add_table(rows=1, cols=1); st_tbl.style = 'Table Grid'
    stc = st_tbl.cell(0,0); _no_border(stc); _set_cell_bg(stc, fc_bg_h)
    stc.width = LEFT_W
    p_st = stc.paragraphs[0]; p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_st = p_st.add_run(f"{bul}/5 Bulmor  ·  {lbl_bul}")
    r_st.bold = True; r_st.font.size = Pt(10); r_st.font.color.rgb = _rgb(fc_hex)

    # Bulmor-Symbole
    aktiv_sym  = " ✓ "
    inaktiv_sym= " ✕ "
    p_sym = lc.add_paragraph(); p_sym.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sym.paragraph_format.space_after = Pt(1); p_sym.paragraph_format.space_before = Pt(3)
    for i in range(5):
        rs = p_sym.add_run("B" + str(i+1))
        rs.bold = True; rs.font.size = Pt(8)
        rs.font.color.rgb = _rgb(fc_hex if (i+1)<=bul else "888888")
        chk = p_sym.add_run("✓ " if (i+1)<=bul else "✕ ")
        chk.font.size = Pt(8)
        chk.font.color.rgb = _rgb(fc_hex if (i+1)<=bul else "FF4444")

    if bulfhr_lst:
        _par(lc, "Fahrer:", bold=True, size=7, color=he, align="center", sb=3)
        for fn in bulfhr_lst:
            _par(lc, fn, size=7.5, color="CCCCCC", align="center")

    if kranke_lst:
        _par(lc, "── KRANK ──", bold=True, size=7.5, color="FF4444", align="center", sb=4)
        for kn, kz in kranke_lst:
            _par(lc, f"{kn}  ({kz})", size=7.5, color="FF8888", align="center")

    _par(lc, "+49 2203 40-2323", size=7, color="888888", align="center", sb=8)

    # ── RECHTE ZELLE ─────────────────────────────────────────────────────────
    # Org-Header
    hdr = rc.add_table(rows=1, cols=1); hdr.style = 'Table Grid'
    hc = hdr.cell(0,0); _no_border(hc); _set_cell_bg(hc, dk); hc.width = RIGHT_W
    _par(hc, f"{STATION}  ·  {DATUM}  ·  {UHRZEIT}", bold=True, size=8.5, color=he, sb=1)

    # Legende Tag/Nacht
    leg = rc.add_table(rows=1, cols=2); leg.style = 'Table Grid'
    lc1 = leg.cell(0,0); lc2 = leg.cell(0,1)
    _no_border(lc1); _no_border(lc2)
    _set_cell_bg(lc1, az); _set_cell_bg(lc2, az2)
    lc1.width = Cm(3); lc2.width = Cm(3)
    _par(lc1, "■ TAGDIENST",   bold=True, size=7, color="FFFFFF", align="center")
    _par(lc2, "■ NACHTDIENST", bold=True, size=7, color="FFFFFF", align="center")
    rc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Betreuer Tagdienst ────────────────────────────────────────────────────
    bth1 = rc.add_table(rows=1, cols=1); bth1.style = 'Table Grid'
    bh1c = bth1.cell(0,0); _no_border(bh1c); _set_cell_bg(bh1c, dk); bh1c.width = RIGHT_W
    _par(bh1c, f"BETREUER – TAGDIENST  ({len(bt_dt)} Personen)", bold=True, size=8.5, color="FFFFFF")

    if bt_dt:
        half = (len(bt_dt)+1)//2
        bt_tbl = rc.add_table(rows=half, cols=4); bt_tbl.style='Table Grid'
        NW = Cm(4.0); ZW = Cm(2.3)
        for i, (name, zeit) in enumerate(bt_dt):
            row = i % half; col = (i // half) * 2
            nc = bt_tbl.cell(row, col); _no_border(nc)
            _set_cell_bg(nc, "F4F4F4" if row%2==0 else "FFFFFF")
            zc = bt_tbl.cell(row, col+1); _no_border(zc)
            _set_cell_bg(zc, "F4F4F4" if row%2==0 else "FFFFFF")
            nc.width = NW; zc.width = ZW
            pn = nc.paragraphs[0]; rn = pn.add_run(name)
            rn.font.size = Pt(8); rn.font.color.rgb = _rgb("111111")
            pn.paragraph_format.space_before = Pt(1); pn.paragraph_format.space_after = Pt(1)
            pz = zc.paragraphs[0]; rz = pz.add_run(zeit)
            rz.font.size = Pt(7.5); rz.font.color.rgb = _rgb("888888"); rz.bold = True
            pz.paragraph_format.space_before = Pt(1); pz.paragraph_format.space_after = Pt(1)

    # ── Betreuer Nachtdienst ─────────────────────────────────────────────────
    if bt_dn:
        bth2 = rc.add_table(rows=1, cols=1); bth2.style = 'Table Grid'
        bh2c = bth2.cell(0,0); _no_border(bh2c); _set_cell_bg(bh2c, "2A2A4A"); bh2c.width = RIGHT_W
        _par(bh2c, f"BETREUER – NACHTDIENST  ({len(bt_dn)} Personen)", bold=True, size=8.5, color="FFFFFF")

        nhalf = (len(bt_dn)+1)//2
        bn_tbl = rc.add_table(rows=nhalf, cols=4); bn_tbl.style='Table Grid'
        NW2 = Cm(4.0); ZW2 = Cm(2.3)
        for i, (name, zeit) in enumerate(bt_dn):
            row = i % nhalf; col = (i // nhalf) * 2
            nc = bn_tbl.cell(row, col); _no_border(nc)
            _set_cell_bg(nc, "EDEDFF" if row%2==0 else "FFFFFF")
            zc = bn_tbl.cell(row, col+1); _no_border(zc)
            _set_cell_bg(zc, "EDEDFF" if row%2==0 else "FFFFFF")
            nc.width = NW2; zc.width = ZW2
            pn = nc.paragraphs[0]; rn = pn.add_run(name)
            rn.font.size = Pt(8); rn.font.color.rgb = _rgb("111111")
            pn.paragraph_format.space_before = Pt(1); pn.paragraph_format.space_after = Pt(1)
            pz = zc.paragraphs[0]; rz = pz.add_run(zeit)
            rz.font.size = Pt(7.5); rz.font.color.rgb = _rgb("888888"); rz.bold = True
            pz.paragraph_format.space_before = Pt(1); pz.paragraph_format.space_after = Pt(1)

    # ── Disposition ──────────────────────────────────────────────────────────
    dh = rc.add_table(rows=1, cols=1); dh.style = 'Table Grid'
    dhc = dh.cell(0,0); _no_border(dhc); _set_cell_bg(dhc, dk); dhc.width = RIGHT_W
    _par(dhc, f"DISPOSITION  ({len(di_dt)+len(di_dn)} Personen)", bold=True, size=8.5, color="FFFFFF")

    def _dispo_block(lst, label, hc_):
        if not lst: return
        hdr_t = rc.add_table(rows=1, cols=1); hdr_t.style='Table Grid'
        h_cell = hdr_t.cell(0,0); _no_border(h_cell); _set_cell_bg(h_cell, hc_); h_cell.width = RIGHT_W
        _par(h_cell, label, bold=True, size=7.5, color="FFFFFF")
        for i, (name, zeit) in enumerate(lst):
            d_tbl = rc.add_table(rows=1, cols=2); d_tbl.style='Table Grid'
            nc = d_tbl.cell(0,0); zc = d_tbl.cell(0,1)
            _no_border(nc); _no_border(zc)
            bg_ = "F4F4F4" if i%2==0 else "FFFFFF"
            _set_cell_bg(nc, bg_); _set_cell_bg(zc, bg_)
            nc.width = Cm(9.6); zc.width = Cm(3.0)
            pn = nc.paragraphs[0]; rn = pn.add_run(name)
            rn.font.size = Pt(8.5); rn.font.color.rgb = _rgb("111111"); rn.bold = True
            pn.paragraph_format.space_before = Pt(1); pn.paragraph_format.space_after = Pt(1)
            pz = zc.paragraphs[0]; pz.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rz = pz.add_run(zeit); rz.font.size = Pt(8); rz.font.color.rgb = _rgb(az); rz.bold = True
            pz.paragraph_format.space_before = Pt(1); pz.paragraph_format.space_after = Pt(1)

    _dispo_block(di_dt, "Tagdienst", az)
    _dispo_block(di_dn, "Nachtdienst", az2)

    # Footer
    ft = rc.add_table(rows=1, cols=1); ft.style='Table Grid'
    ftc = ft.cell(0,0); _no_border(ftc); _set_cell_bg(ftc, dk); ftc.width = RIGHT_W
    _par(ftc, "DRK Köln  ·  +49 2203 40-2323  ·  flughafen@drk-koeln.de",
         size=7.5, color="FFFFFF", align="center", sb=4)

    suffix = f"Bul{bul}" if bul < 5 else ""
    out = os.path.join(ZIEL, f"W_V2_{scheme_name}{suffix}_25032026.docx")
    doc.save(out)
    print(f"  [OK] W_V2_{scheme_name}{suffix}_25032026.docx")


# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Lade Dienstplan-Daten ...")
    data = lade()
    bs = _nach_schicht(data,'betreuer'); ds = _nach_schicht(data,'dispo')
    print(f"  Betreuer  DT:{len(bs['DT'])}  DN:{len(bs['DN'])}")
    print(f"  Dispo     DT:{len(ds['DT'])}  DN:{len(ds['DN'])}")
    print(f"  Krank:    {len(_kranke(data))}")
    print(f"\nZielordner: {ZIEL}\n")

    print("── 6 Farbvarianten (PDF) ──────────────────────────────────────────")
    erstelle_farb_pdfs(data, bul=5, einz=28, pat=5, pax=42500)

    print("\n── Bulmor-Status-Beispiele  bul=1..5 (PDF) ───────────────────────")
    erstelle_bulmor_pdfs(data, einz=28, pat=5, pax=42500)

    print("\n── Word-Versionen ─────────────────────────────────────────────────")
    # 3 Farbvarianten als Word
    for sn in ("OzeanBlau", "Smaragd", "Burgund"):
        erstelle_word_v2(data, scheme_name=sn, bul=5,  einz=28, pat=5, pax=42500)
    # Bulmor-Status-Beispiele als Word (OzeanBlau)
    for bul_n in (1, 2, 3, 4):
        erstelle_word_v2(data, scheme_name="OzeanBlau", bul=bul_n, einz=28, pat=5, pax=42500)

    print(f"\n✓ Alle Dateien gespeichert in:\n  {ZIEL}")
